from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from uuid import uuid4
import json
import os
import httpx
import hashlib
import shutil
from pathlib import Path
from typing import Optional, Dict, Any, List
import logging
import sys

# Загрузка переменных окружения из .env файла
try:
    from dotenv import load_dotenv
    # Загружаем .env из директории сервиса
    env_path = Path(__file__).resolve().parent / ".env"
    if env_path.exists():
        load_dotenv(env_path)
        logging.getLogger(__name__).info(f"✅ Загружен .env файл: {env_path}")
    else:
        # Пробуем загрузить из корня проекта
        project_env = Path(__file__).resolve().parent.parent.parent.parent / ".env"
        if project_env.exists():
            load_dotenv(project_env)
            logging.getLogger(__name__).info(f"✅ Загружен .env файл: {project_env}")
except ImportError:
    # python-dotenv не установлен, используем только системные переменные
    pass
from file_parser import parse_file
import database
from openpyxl import load_workbook
from settings.excel_semantic_settings import get_excel_semantic_mode
from ai.ai_excel_semantic_parser import CanonicalSourceData
from utils.canonical_collector import collect_canonical_from_workbook
from utils.canonical_collector import analyze_workbook_sheets
from utils.canonical_to_passport import canonical_to_passport_payload
from domain.passport_requirements import (
    evaluate_generation_readiness,
    GenerationReadinessResult,
)
from utils.energy_aggregator import (
    aggregate_energy_data,
    should_aggregate_file,
    write_aggregation_json,
    aggregate_from_db_json,
    aggregate_usage_categories,
    distribute_categories_by_quarter,
)
from utils.equipment_parser import (
    is_equipment_file,
    parse_equipment_workbook,
    write_equipment_json,
)
from utils.building_envelope_parser import (
    is_envelope_file,
    parse_building_envelope,
    write_envelope_json,
)
from utils.nodes_parser import (
    is_nodes_file,
    parse_nodes_workbook,
    write_nodes_json,
    load_nodes_from_json,
)
from utils.balance_sheet_detector import is_balance_sheet_file
from utils.balance_sheet_node_extractor import extract_node_consumption_from_balance_sheet
from utils.aggregation_log import log_aggregation_event
from utils.intelligent_router import IntelligentRouter
from utils.progress_tracker import (
    FileType,
    ProcessingStage,
    create_progress_tracker,
    get_progress_tracker,
    remove_progress_tracker,
)
from utils.readiness_validator import (
    validate_generation_readiness,
    get_upload_checklist,
)
from utils.data_validator import (
    validate_data_for_template,
)
from models.schemas import ValidateRequest, EnterpriseCreate, EditablePayload

# Настройка логирования ДО использования logger
# Уровень логирования можно изменить через переменную окружения LOG_LEVEL
log_level = os.getenv("LOG_LEVEL", "DEBUG").upper()
logging.basicConfig(
    level=getattr(logging, log_level, logging.DEBUG),
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)
logger.info(f"🔧 Уровень логирования установлен: {log_level}")

# Переключатель режимов работы системы
SYSTEM_MODE = os.getenv("SYSTEM_MODE", "debug").lower()  # По умолчанию debug для разработки
if SYSTEM_MODE not in ["debug", "production"]:
    SYSTEM_MODE = "debug"
    logger.warning("⚠️ Неверный SYSTEM_MODE, установлен режим 'debug'")
logger.info(f"🔧 Режим работы системы: {SYSTEM_MODE.upper()} (для изменения установите SYSTEM_MODE=debug или SYSTEM_MODE=production, или используйте переключатель в веб-интерфейсе)")

# Добавляем tools в путь для импорта генератора
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent.parent / "tools"))
try:
    from pkm690_excel_generator import PKM690ExcelGenerator

    HAS_GENERATOR = True
except ImportError as e:
    logger.warning(f"PKM690 генератор не доступен: {e}")
    HAS_GENERATOR = False

try:
    from fill_energy_passport import (
        fill_struktura_pr2,
        fill_nodes_sheet,
        load_default_nodes,
        fill_building_envelope_sheet,
        fill_equipment_sheet,
        fill_balans_sheet,
        fill_dinamika_sheet,
        fill_fuel_dynamics_sheet,
        fill_specific_consumption_sheet,
        fill_meropriyatiya_sheet,
        fill_monthly_sheet,
    )

    HAS_FILLER = True
except ImportError as e:
    logger.warning(f"fill_energy_passport недоступен: {e}")
    HAS_FILLER = False

# Импорт AI восстановителя формул
try:
    from utils.ai_formula_restorer import AIFormulaRestorer

    HAS_FORMULA_RESTORER = True
except ImportError as e:
    logger.warning(f"ai_formula_restorer недоступен: {e}")
    HAS_FORMULA_RESTORER = False

try:
    from utils.word_report_generator import WordReportGenerator

    HAS_WORD_GENERATOR = True
except ImportError as e:
    logger.warning(f"WordReportGenerator недоступен: {e}")
    HAS_WORD_GENERATOR = False

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent.parent

WEB_DIR = Path(__file__).resolve().parent / "web"
WEB_DIR.mkdir(parents=True, exist_ok=True)

database.init_db()

app = FastAPI(title="EAIP ingest", version="0.1.0")

# Inbox directory for uploaded files
INBOX_DIR = os.getenv("INBOX_DIR", "/data/inbox")
os.makedirs(INBOX_DIR, exist_ok=True)

AGGREGATED_DIR = Path(
    os.getenv("AGGREGATED_DIR", os.path.join(INBOX_DIR, "aggregated"))
)
AGGREGATED_DIR.mkdir(parents=True, exist_ok=True)
# Временная директория для обработки файлов (для Word валидации и т.д.)
DATA_DIR = Path(os.getenv("DATA_DIR", os.path.join(INBOX_DIR, "temp")))
DATA_DIR.mkdir(parents=True, exist_ok=True)

# Константы согласно ТЗ (раздел 4.1, 4.2)
ALLOWED_EXTENSIONS = {".xlsx", ".xlsm", ".docx", ".pdf", ".jpg", ".jpeg", ".png"}
ALLOWED_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",  # XLSX
    "application/vnd.ms-excel.sheet.macroEnabled.12",  # XLSM
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # DOCX
    "application/pdf",
    "image/jpeg",
    "image/jpg",
    "image/png",
}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 МБ согласно ТЗ (раздел 4.2)
FILE_TYPE_LABELS = {
    ".xlsx": "Excel (XLSX)",
    ".xlsm": "Excel с макросами (XLSM)",
    ".docx": "Word (DOCX)",
    ".pdf": "PDF",
    ".jpg": "Изображение (JPG)",
    ".jpeg": "Изображение (JPEG)",
    ".png": "Изображение (PNG)",
}


@app.get("/api/batches/{batch_id}/generation-readiness")
def api_generation_readiness(batch_id: str):
    """
    Оценка готовности к генерации паспорта на основе CanonicalSourceData.
    """
    upload = database.get_upload_by_batch(batch_id)
    if not upload:
        raise HTTPException(status_code=404, detail=f"Загрузка {batch_id} не найдена")

    raw_json = upload.get("raw_json") or {}
    canonical = None
    if isinstance(raw_json, dict):
        canonical_dict = raw_json.get("canonical_source")
        if isinstance(canonical_dict, dict):
            try:
                canonical = CanonicalSourceData.parse_obj(canonical_dict)
            except Exception:
                canonical = None
    if canonical is None:
        # Try to reconstruct from the original file if name is known
        try:
            filename = upload.get("filename")
            if filename:
                canonical = collect_canonical_from_workbook(filename)
        except Exception:
            canonical = None

    result = evaluate_generation_readiness(canonical)
    return {
        "batch_id": batch_id,
        "overall_status": result.overall_status,
        "missing_required": [rf.__dict__ for rf in result.missing_required],
        "missing_optional": [rf.__dict__ for rf in result.missing_optional],
        "notes": result.notes,
        "mode": get_excel_semantic_mode(),
    }


@app.get("/api/batches/{batch_id}/canonical-debug")
def api_canonical_debug(batch_id: str):
    """
    Внутренний отладочный эндпоинт для просмотра CanonicalSourceData и вклада AI.
    Не выполняет генерацию и не меняет состояние.
    """
    mode = get_excel_semantic_mode()
    logger.info("Canonical debug requested for batch_id=%s mode=%s", batch_id, mode)

    upload = database.get_upload_by_batch(batch_id)
    if not upload:
        raise HTTPException(status_code=404, detail=f"Загрузка {batch_id} не найдена")

    # 1) Пытаемся извлечь canonical_source из сохранённых raw_json
    raw_json = upload.get("raw_json") or {}
    canonical = None
    if isinstance(raw_json, dict):
        canonical_dict = raw_json.get("canonical_source")
        if isinstance(canonical_dict, dict):
            try:
                canonical = CanonicalSourceData.parse_obj(canonical_dict)
            except Exception:
                canonical = None

    # 2) Если нет — пытаемся реконструировать из исходного файла (best-effort)
    if canonical is None:
        try:
            filename = upload.get("filename")
            if filename:
                canonical = collect_canonical_from_workbook(filename)
        except Exception as exc:
            logger.warning(
                "Failed to reconstruct canonical for batch_id=%s: %s", batch_id, exc
            )
            canonical = None

    # Формируем ответ
    response = {
        "batch_id": batch_id,
        "mode": mode,
        "canonical_source": canonical.dict() if canonical else None,
        "provenance": (canonical.provenance if canonical else {}) or {},
        "sheets": [],
    }
    # Попытка добавить sheet-level детали (on-the-fly анализ исходного файла)
    try:
        filename = upload.get("filename")
        if filename:
            response["sheets"] = analyze_workbook_sheets(filename)
    except Exception as exc:
        logger.warning("Sheet-level debug analysis failed for %s: %s", batch_id, exc)
    return response


# Mapping of resource type codes to human-friendly labels
RESOURCE_LABELS: Dict[str, str] = {
    "electricity": "Электроэнергия",
    "gas": "Газ",
    "heat": "Тепловая энергия",
    "water": "Вода",
    "fuel": "Топливо и ГСМ",
    "equipment": "Оборудование",
    "envelope": "Расчет теплопотерь по зданиям",
    "nodes": "Узлы учета",
    "other": "Прочее",
}

# Хранилище результатов парсинга (в продакшене использовать Redis/PostgreSQL)
# Для совместимости оставили кеш, но первичное хранилище — SQLite.
parsing_results_cache: Dict[str, Dict[str, Any]] = {}

# Enable CORS for web interface
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify exact origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def render_html_page(filename: str) -> HTMLResponse:
    file_path = WEB_DIR / filename
    if not file_path.exists():
        logger.error("HTML page %s not found in %s", filename, WEB_DIR)
        raise HTTPException(status_code=500, detail="Страница временно недоступна")
    return HTMLResponse(content=file_path.read_text(encoding="utf-8"))


def normalize_file_type(parsing_result: Optional[Dict[str, Any]], file_ext: str) -> str:
    if parsing_result and parsing_result.get("file_type"):
        return str(parsing_result.get("file_type")).lower()
    ext = file_ext.lower()
    if ext in {".xlsx", ".xlsm"}:
        return "excel"
    if ext in {".docx"}:
        return "docx"
    if ext in {".pdf"}:
        return "pdf"
    return "unknown"


def get_file_type_enum(file_ext: str) -> FileType:
    """Преобразует расширение файла в FileType enum"""
    ext = file_ext.lower()
    if ext in {".xlsx", ".xlsm", ".xls"}:
        return FileType.EXCEL
    elif ext in {".docx"}:
        return FileType.WORD
    elif ext in {".pdf"}:
        return FileType.PDF
    elif ext in {".jpg", ".jpeg", ".png"}:
        return FileType.IMAGE
    else:
        return FileType.UNKNOWN


def build_parsing_summary(
    parsing_result: Optional[Dict[str, Any]], file_ext: str
) -> Optional[Dict[str, Any]]:
    if not parsing_result or not parsing_result.get("parsed"):
        return None

    data = parsing_result.get("data") or {}
    file_type = normalize_file_type(parsing_result, file_ext)

    if file_type == "pdf":
        metadata = data.get("metadata", {})
        return {
            "pages": metadata.get("num_pages", 0),
            "characters": data.get("total_characters", 0),
            "tables": data.get("total_tables", 0),
        }
    if file_type == "excel":
        sheets = data.get("sheets", [])
        return {
            "sheets": len(sheets),
            "total_rows": sum(sheet.get("max_row", 0) for sheet in sheets),
        }
    if file_type == "docx":
        return {
            "paragraphs": len(data.get("paragraphs", [])),
            "tables": len(data.get("tables", [])),
        }
    return None


def build_editable_text(parsing_result: Optional[Dict[str, Any]]) -> str:
    if not parsing_result or not parsing_result.get("parsed"):
        return ""

    data = parsing_result.get("data") or {}
    file_type = normalize_file_type(parsing_result, "")

    if file_type == "pdf":
        text = data.get("text")
        return text if isinstance(text, str) else ""

    if file_type == "excel":
        lines: List[str] = []
        sheets = data.get("sheets", [])
        for sheet in sheets:
            name = sheet.get("name") or "Без названия"
            lines.append(f"=== Лист: {name} ===")
            rows = sheet.get("rows") or []
            if not rows:
                lines.append("(Нет данных)")
            else:
                for row in rows[:100]:
                    if isinstance(row, list):
                        values = ["" if value is None else str(value) for value in row]
                    elif isinstance(row, dict):
                        values = [
                            "" if value is None else str(value)
                            for value in row.values()
                        ]
                    else:
                        values = ["" if row is None else str(row)]
                    lines.append("\t".join(values))
            lines.append("")
        return "\n".join(lines).strip()

    if file_type == "docx":
        paragraphs = data.get("paragraphs") or []
        return "\n".join(
            paragraph.get("text", "")
            for paragraph in paragraphs
            if paragraph.get("text")
        )

    return json.dumps(data, ensure_ascii=False, indent=2)


def ensure_parsing_cached(batch_id: str) -> Optional[Dict[str, Any]]:
    """
    Обеспечивает наличие результатов парсинга в кэше.
    Если данных нет в кэше, загружает их из БД.
    """
    if batch_id not in parsing_results_cache:
        record = database.get_upload_by_batch(batch_id)
        if record:
            # Формируем структуру данных для кэша из записи БД
            raw_json = record.get("raw_json")
            if raw_json:
                # Если raw_json уже является полной структурой (с batch_id, parsing и т.д.)
                if isinstance(raw_json, dict) and "batch_id" in raw_json:
                    parsing_results_cache[batch_id] = raw_json
                else:
                    # Если raw_json - это только данные парсинга, формируем полную структуру
                    parsing_results_cache[batch_id] = {
                        "batch_id": batch_id,
                        "filename": record.get("filename"),
                        "file_path": None,  # Путь не сохраняется в БД
                        "file_type": record.get("file_type"),
                        "file_size": record.get("file_size"),
                        "parsing": raw_json if isinstance(raw_json, dict) else {},
                        "status": record.get("status", "unknown"),
                        "resource_type": record.get("parsing_summary", {}).get("resource_type") if record.get("parsing_summary") else None,
                    }
            else:
                # Если raw_json отсутствует, создаем минимальную структуру
                logger.warning(f"raw_json отсутствует для batch_id={batch_id}, создана минимальная структура")
                parsing_results_cache[batch_id] = {
                    "batch_id": batch_id,
                    "filename": record.get("filename"),
                    "file_path": None,
                    "file_type": record.get("file_type"),
                    "file_size": record.get("file_size"),
                    "parsing": None,
                    "status": record.get("status", "error"),
                    "error": "Данные парсинга не найдены в БД",
                }
    return parsing_results_cache.get(batch_id)


@app.get("/health")
def health():
    return {"service": "ingest", "status": "ok"}


@app.get("/test-xlsm")
def test_xlsm():
    """Простой тестовый endpoint для проверки .xlsm поддержки"""
    return {
        "xlsm_in_allowed": ".xlsm" in ALLOWED_EXTENSIONS,
        "allowed_extensions": sorted(ALLOWED_EXTENSIONS),
        "message": "Тест поддержки .xlsm",
    }


@app.get("/api/debug/extensions")
@app.get("/debug/extensions")
def debug_extensions():
    """Диагностический endpoint для проверки поддерживаемых форматов"""
    from datetime import datetime

    try:
        return {
            "allowed_extensions": sorted(ALLOWED_EXTENSIONS),
            "xlsm_supported": ".xlsm" in ALLOWED_EXTENSIONS,
            "allowed_mime_types": list(ALLOWED_MIME_TYPES),
            "file_type_labels": FILE_TYPE_LABELS,
            "timestamp": datetime.now().isoformat(),
            "code_version": "2025-01-16-xlsm-support",
            "server_status": "running",
        }
    except Exception as e:
        return {
            "error": str(e),
            "allowed_extensions": list(ALLOWED_EXTENSIONS)
            if "ALLOWED_EXTENSIONS" in globals()
            else "not found",
        }


@app.get("/web/upload")
def upload_page():
    return render_html_page("upload.html")


@app.get("/web/files")
def files_page():
    """Веб-интерфейс для просмотра загруженных файлов"""
    return render_html_page("files.html")


@app.get("/web/results")
def results_page():
    return render_html_page("results.html")


@app.get("/web/normative")
def normative_upload_page():
    """Веб-интерфейс для загрузки нормативных документов"""
    return render_html_page("normative_upload.html")


@app.get("/web/normative/upload")
def normative_upload_page_alt():
    """Веб-интерфейс для загрузки нормативных документов (альтернативный URL)"""
    return render_html_page("normative_upload.html")


@app.get("/web/normative/dashboard")
def web_normative_dashboard():
    """Дашборд со статистикой нормативов"""
    return render_html_page("normative_dashboard.html")


@app.get("/api/normative/test")
def test_normative_endpoint():
    """Тестовый endpoint для проверки доступности API нормативных документов"""
    return {
        "status": "ok",
        "message": "API нормативных документов доступен",
        "endpoints": {
            "upload": "/api/normative/upload (POST)",
            "documents": "/api/normative/documents (GET)",
            "rules": "/api/normative/rules/{rule_type} (GET)",
            "ai_status": "/api/normative/ai-status (GET)",
        },
    }


@app.get("/api/normative/ai-status")
def get_ai_status():
    """Получить статус AI для нормативных документов"""
    try:
        from settings.ai_settings import get_ai_status, get_ai_settings
        from domain.normative_importer import get_normative_importer

        # Получаем статус AI из единого модуля настроек
        ai_status = get_ai_status()
        get_ai_settings()

        # Проверяем доступность импортера
        importer = get_normative_importer()
        importer_available = importer is not None
        ai_parser_available = (
            importer and importer.ai_parser is not None and importer.ai_parser.enabled
            if importer
            else False
        )

        return {
            **ai_status,
            "importer_available": importer_available,
            "ai_parser_available": ai_parser_available,
        }
    except Exception as e:
        logger.exception("Ошибка проверки статуса AI")
        return {
            "ai_enabled": False,
            "ai_provider": "unknown",
            "has_api_key": False,
            "has_valid_config": False,
            "importer_available": False,
            "ai_parser_available": False,
            "error": str(e),
            "message": f"❌ Ошибка проверки статуса AI: {e}",
        }


@app.get("/api/debug/extensions")
def api_debug_extensions():
    """Диагностический endpoint для проверки поддерживаемых форматов (API версия)"""
    from datetime import datetime

    try:
        return {
            "allowed_extensions": sorted(ALLOWED_EXTENSIONS),
            "xlsm_supported": ".xlsm" in ALLOWED_EXTENSIONS,
            "allowed_mime_types": list(ALLOWED_MIME_TYPES),
            "file_type_labels": FILE_TYPE_LABELS,
            "timestamp": datetime.now().isoformat(),
            "code_version": "2025-01-16-xlsm-support",
            "server_status": "running",
        }
    except Exception as e:
        return {
            "error": str(e),
            "allowed_extensions": list(ALLOWED_EXTENSIONS)
            if "ALLOWED_EXTENSIONS" in globals()
            else "not found",
        }


@app.get("/api/enterprises")
def api_list_enterprises():
    return {"items": database.list_enterprises()}


@app.post("/api/enterprises")
def api_create_enterprise(payload: EnterpriseCreate):
    enterprise = database.get_or_create_enterprise(payload.name)
    return enterprise


@app.get("/api/enterprises/{enterprise_id}/uploads")
def api_enterprise_history(enterprise_id: int):
    enterprise = database.get_enterprise_by_id(enterprise_id)
    if not enterprise:
        raise HTTPException(status_code=404, detail="Предприятие не найдено")
    history = database.list_uploads_for_enterprise(enterprise_id)
    return {"enterprise": enterprise, "uploads": history}


@app.get("/api/enterprises/{enterprise_id}/upload-checklist")
def api_get_upload_checklist(enterprise_id: int):
    """
    Возвращает чек-лист требуемых файлов для предприятия.

    Returns:
        Словарь с чек-листом требуемых и опциональных файлов.
    """
    enterprise = database.get_enterprise_by_id(enterprise_id)
    if not enterprise:
        raise HTTPException(status_code=404, detail="Предприятие не найдено")

    checklist = get_upload_checklist(enterprise_id)
    return {
        "enterprise_id": enterprise_id,
        "enterprise_name": enterprise.get("name"),
        **checklist,
    }


@app.get("/api/enterprises/{enterprise_id}/generation-readiness")
def api_get_generation_readiness(enterprise_id: int):
    """
    Возвращает статус готовности данных для генерации энергетического паспорта.

    Returns:
        Словарь с результатами проверки готовности:
        - ready: готовность к генерации
        - completeness_score: показатель готовности (0.0-1.0)
        - missing_resources: список недостающих ресурсов
        - missing_files: список недостающих файлов
        - available_resources: список доступных ресурсов
        - warnings: список предупреждений
        - progress_percentage: процент готовности
    """
    enterprise = database.get_enterprise_by_id(enterprise_id)
    if not enterprise:
        raise HTTPException(status_code=404, detail="Предприятие не найдено")

    readiness = validate_generation_readiness(enterprise_id)
    return {
        "enterprise_id": enterprise_id,
        "enterprise_name": enterprise.get("name"),
        **readiness,
    }


@app.get("/api/uploads/{batch_id}")
def api_get_upload(batch_id: str):
    record = database.get_upload_by_batch(batch_id)
    if not record:
        raise HTTPException(status_code=404, detail="Загрузка не найдена")
    return record


@app.get("/api/uploads/{batch_id}/editable")
def api_get_editable(batch_id: str):
    record = database.get_upload_by_batch(batch_id)
    if not record:
        raise HTTPException(status_code=404, detail="Загрузка не найдена")
    return {
        "batch_id": batch_id,
        "editable_text": record.get("editable_text") or "",
        "updated_at": record.get("parsed_updated_at"),
    }


@app.post("/api/uploads/{batch_id}/editable")
def api_update_editable(batch_id: str, payload: EditablePayload):
    ensure_parsing_cached(batch_id)
    database.update_editable_text(batch_id, payload.text)
    return {"batch_id": batch_id, "status": "saved"}


@app.get("/ingest/parse/{batch_id}")
async def get_parsing_results(batch_id: str):
    data = ensure_parsing_cached(batch_id)
    if not data:
        raise HTTPException(
            status_code=404,
            detail=f"Результаты парсинга для batch_id {batch_id} не найдены",
        )
    return JSONResponse(content=data)


@app.get("/api/progress/{batch_id}")
async def get_progress(batch_id: str):
    """
    Получение текущего прогресса обработки файла

    Returns:
        Статус прогресса с детальной информацией по этапам
    """
    tracker = get_progress_tracker(batch_id)
    if not tracker:
        # Если трекер не найден, возможно обработка еще не началась или уже завершена
        # Проверяем в БД
        upload = database.get_upload_by_batch(batch_id)
        if upload:
            # Возвращаем финальный статус из БД
            return {
                "batch_id": batch_id,
                "file_type": "unknown",
                "overall_progress": 100,
                "current_stage": "completed",
                "stages": {},
                "is_completed": True,
                "has_error": upload.get("status") == "error",
                "error": None
                if upload.get("status") != "error"
                else "Обработка завершена с ошибкой",
                "message": "Обработка завершена. Трекер удален из памяти.",
            }
        raise HTTPException(
            status_code=404,
            detail=f"Трекер прогресса для batch_id {batch_id} не найден",
        )

    return tracker.get_status()


@app.post("/web/upload/{batch_id}/cancel")
async def cancel_upload(batch_id: str):
    """
    Отмена обработки загруженного файла

    Args:
        batch_id: Уникальный идентификатор загрузки

    Returns:
        Статус отмены
    """
    tracker = get_progress_tracker(batch_id)
    if not tracker:
        # Проверяем в БД, возможно обработка уже завершена
        upload = database.get_upload_by_batch(batch_id)
        if upload:
            return {
                "batch_id": batch_id,
                "cancelled": False,
                "message": "Обработка уже завершена, отмена невозможна",
            }
        raise HTTPException(
            status_code=404,
            detail=f"Трекер прогресса для batch_id {batch_id} не найден",
        )

    if tracker.is_cancelled():
        return {
            "batch_id": batch_id,
            "cancelled": True,
            "message": "Обработка уже была отменена ранее",
        }

    if tracker.completed_at:
        return {
            "batch_id": batch_id,
            "cancelled": False,
            "message": "Обработка уже завершена, отмена невозможна",
        }

    tracker.cancel()
    logger.info(f"Обработка файла отменена пользователем: batch_id={batch_id}")

    return {
        "batch_id": batch_id,
        "cancelled": True,
        "message": "Обработка файла отменена",
        "status": tracker.get_status(),
    }


@app.get("/api/diagnose/pdf")
async def diagnose_pdf_endpoint(
    file_path: Optional[str] = None, batch_id: Optional[str] = None
):
    """
    Диагностика PDF файла

    Args:
        file_path: Прямой путь к PDF файлу (для тестирования)
        batch_id: ID загрузки для диагностики загруженного файла

    Returns:
        Диагностический отчет
    """
    from utils.pdf_diagnostics import diagnose_pdf

    if batch_id:
        # Получаем путь к файлу из загрузки
        upload = database.get_upload_by_batch(batch_id)
        if not upload:
            raise HTTPException(
                status_code=404, detail=f"Загрузка {batch_id} не найдена"
            )

        file_path = upload.get("file_path")
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(
                status_code=404, detail=f"Файл для batch_id {batch_id} не найден"
            )

    if not file_path:
        raise HTTPException(status_code=400, detail="Укажите file_path или batch_id")

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail=f"Файл {file_path} не найден")

    try:
        report = diagnose_pdf(file_path)
        return report
    except Exception as e:
        logger.error(f"Ошибка диагностики PDF {file_path}: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка диагностики: {e}")


@app.get("/ingest/parse/{batch_id}/summary")
async def get_parsing_summary(batch_id: str):
    data = ensure_parsing_cached(batch_id)
    if not data:
        raise HTTPException(
            status_code=404,
            detail=f"Результаты парсинга для batch_id {batch_id} не найдены",
        )

    parsing_data = data.get("parsing") or {}
    if not parsing_data or not parsing_data.get("parsed"):
        return {
            "batch_id": batch_id,
            "status": data.get("status", "error"),
            "message": "Файл не был распознан или произошла ошибка",
        }

    file_path = data.get("file_path")
    file_ext = Path(file_path).suffix.lower() if file_path else ""
    summary = build_parsing_summary(parsing_data, file_ext) or {}
    summary.update(
        {
            "batch_id": batch_id,
            "filename": data.get("filename"),
            "file_type": data.get("file_type"),
            "status": data.get("status", "success"),
            "parsed": True,
        }
    )
    return summary


@app.post("/ingest/validate")
async def proxy_validate(req: ValidateRequest):
    """Прокси для вызова validate сервиса из веб-интерфейса"""
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                "http://validate:8002/validate/run", json={"batchId": req.batchId}
            )
            resp.raise_for_status()
            return resp.json()
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Validate service timeout")
    except httpx.HTTPStatusError as exc:
        raise HTTPException(
            status_code=exc.response.status_code,
            detail=f"Validate service error: {exc.response.text}",
        )
    except httpx.RequestError as exc:
        raise HTTPException(
            status_code=503, detail=f"Validate service unavailable: {exc}"
        )


def validate_file(file: UploadFile):
    """Валидация файла согласно ТЗ (раздел 4.1) с улучшенным логированием"""
    logger.info("=" * 70)
    logger.info(f"🔍 [VALIDATE] НАЧАЛО ВАЛИДАЦИИ файла: {file.filename}")
    logger.info("=" * 70)
    
    # Проверка 1: Имя файла
    if not file.filename:
        logger.error("❌ [VALIDATE] ОШИБКА: Имя файла отсутствует")
        return False, "Имя файла обязательно"

    file_ext = Path(file.filename).suffix.lower()
    logger.info(f"📋 [VALIDATE] Расширение файла: {file_ext}")
    logger.info(f"📋 [VALIDATE] Разрешенные расширения: {sorted(ALLOWED_EXTENSIONS)}")

    # Проверка 2: Расширение файла
    if file_ext not in ALLOWED_EXTENSIONS:
        allowed_str = ", ".join(sorted(ALLOWED_EXTENSIONS))
        error_msg = f"Неподдерживаемый формат файла: {file_ext}. Разрешены: {allowed_str}"
        logger.error("=" * 70)
        logger.error(f"❌ [VALIDATE] ОШИБКА ВАЛИДАЦИИ: {error_msg}")
        logger.error(f"   Файл: {file.filename}")
        logger.error(f"   Расширение: {file_ext}")
        logger.error(f"   Разрешенные: {allowed_str}")
        logger.error("=" * 70)
        return False, error_msg

    logger.info(f"✅ [VALIDATE] Расширение {file_ext} разрешено")

    # Проверка 3: MIME type
    content_type = getattr(file, "content_type", None)
    logger.info(f"📋 [VALIDATE] MIME type файла: {content_type or 'не указан'}")

    if content_type:
        # Специальная обработка для Word файлов (аналогично .xlsm)
        if file_ext == ".docx":
            word_mime_types = [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",  # Старые версии Word
                "application/octet-stream",  # Некоторые браузеры
            ]
            if content_type in word_mime_types or content_type in ALLOWED_MIME_TYPES:
                logger.info(f"✅ [VALIDATE] MIME type для Word файла принят: {content_type}")
            else:
                logger.warning(
                    f"⚠️ [VALIDATE] Неожиданный MIME type для .docx файла: {content_type}, "
                    f"но разрешаем загрузку (расширение корректно)"
                )
        # Для Excel файлов с макросами может быть разный MIME type
        elif file_ext == ".xlsm":
            # Принимаем как стандартный XLSX MIME type, так и специальный для макросов
            xlsm_mime_types = [
                "application/vnd.ms-excel.sheet.macroEnabled.12",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "application/octet-stream",  # Некоторые браузеры отправляют так
            ]
            if (
                content_type not in ALLOWED_MIME_TYPES
                and content_type not in xlsm_mime_types
            ):
                logger.warning(
                    f"⚠️ [VALIDATE] Неожиданный MIME type для .xlsm файла: {content_type}, но разрешаем загрузку"
                )
            else:
                logger.info(
                    f"✅ [VALIDATE] MIME type для .xlsm файла принят: {content_type}"
                )
        elif content_type not in ALLOWED_MIME_TYPES:
            error_msg = f"Неподдерживаемый тип файла: {content_type}. Разрешены: {', '.join(sorted(ALLOWED_MIME_TYPES))}"
            logger.error("=" * 70)
            logger.error(f"❌ [VALIDATE] ОШИБКА ВАЛИДАЦИИ: {error_msg}")
            logger.error(f"   Файл: {file.filename}")
            logger.error(f"   Расширение: {file_ext}")
            logger.error(f"   MIME type: {content_type}")
            logger.error(f"   Разрешенные MIME types: {sorted(ALLOWED_MIME_TYPES)}")
            logger.error("=" * 70)
            return False, error_msg
        else:
            logger.info(f"✅ [VALIDATE] MIME type {content_type} разрешен")
    else:
        logger.warning(
            f"⚠️ [VALIDATE] MIME type не указан для файла {file.filename}, "
            f"но расширение {file_ext} корректно - разрешаем загрузку"
        )

    logger.info("=" * 70)
    logger.info(f"✅ [VALIDATE] ВАЛИДАЦИЯ ПРОЙДЕНА: {file.filename}")
    logger.info("=" * 70)
    return True, None


async def validate_file_size(file: UploadFile):
    content = await file.read()
    await file.seek(0)

    if len(content) > MAX_FILE_SIZE:
        max_mb = MAX_FILE_SIZE / (1024 * 1024)
        file_mb = len(content) / (1024 * 1024)
        return (
            False,
            f"Размер файла ({file_mb:.2f} МБ) превышает максимально допустимый ({max_mb} МБ)",
        )

    if len(content) == 0:
        return False, "Файл пустой"

    return True, None


@app.post("/web/upload")
async def upload_file(
    file: UploadFile = File(...),
    enterprise_id: Optional[int] = Form(None),
    enterprise_name: Optional[str] = Form(None),
    resource_type: Optional[str] = Form(None),
    system_mode: Optional[str] = Form(None),
):
    """Приём файла через веб-интерфейс с привязкой к предприятию"""
    if file.filename:
        file_ext = Path(file.filename).suffix.lower()
        logger.info(f"📤 [UPLOAD] Начало загрузки файла: {file.filename}")
        logger.info(f"📤 [UPLOAD] Расширение файла: {file_ext}")
        logger.info(f"📤 [UPLOAD] Разрешенные расширения: {sorted(ALLOWED_EXTENSIONS)}")
        logger.info(f"📤 [UPLOAD] MIME type: {getattr(file, 'content_type', 'не указан')}")

    if enterprise_name and enterprise_name.strip():
        enterprise = database.get_or_create_enterprise(enterprise_name)
    elif enterprise_id is not None:
        enterprise = database.get_enterprise_by_id(int(enterprise_id))
        if not enterprise:
            raise HTTPException(status_code=400, detail="Предприятие не найдено")
    else:
        raise HTTPException(
            status_code=400, detail="Укажите предприятие или введите новое"
        )

    # Сохраняем тип, указанный пользователем (если есть)
    user_provided_type = (
        resource_type if resource_type and resource_type in RESOURCE_LABELS else None
    )

    # Временная инициализация resource_type для сохранения в кеш
    # Будет переопределена после парсинга на основе содержимого
    if not resource_type or resource_type not in RESOURCE_LABELS:
        resource_type = (
            "other"  # Временное значение, будет переопределено после парсинга
        )

    is_valid, error_msg = validate_file(file)
    if not is_valid:
        logger.error(
            f"❌ [UPLOAD] Валидация не пройдена для {file.filename}: {error_msg}"
        )
        raise HTTPException(status_code=400, detail=error_msg)
    logger.info(f"✅ [UPLOAD] Валидация пройдена для {file.filename}")

    # Проверка размера файла для всех типов
    size_valid, size_error = await validate_file_size(file)
    if not size_valid:
        raise HTTPException(status_code=400, detail=size_error)

    os.makedirs(INBOX_DIR, exist_ok=True)
    batch_id = str(uuid4())
    safe_filename = os.path.basename(file.filename)
    file_ext = Path(file.filename).suffix.lower()
    dst = os.path.join(INBOX_DIR, f"{batch_id}__{safe_filename}")

    # Создаем трекер прогресса
    file_type_enum = get_file_type_enum(file_ext)
    tracker = create_progress_tracker(batch_id, file_type_enum)
    tracker.update_stage(
        ProcessingStage.UPLOAD, progress=0, message="Начало загрузки файла"
    )

    try:
        file_hash = hashlib.sha1()
        file_size_total = 0
        # Получаем размер файла для расчета прогресса загрузки
        file.file.seek(0, 2)  # Переходим в конец
        total_size = file.file.tell()
        file.file.seek(0)  # Возвращаемся в начало

        with open(dst, "wb") as output_file:
            while True:
                # Проверяем, не была ли отменена обработка
                if tracker.is_cancelled():
                    os.remove(dst)  # Удаляем частично загруженный файл
                    raise HTTPException(
                        status_code=499, detail="Обработка файла отменена пользователем"
                    )

                chunk = file.file.read(8192)
                if not chunk:
                    break
                output_file.write(chunk)
                file_hash.update(chunk)
                file_size_total += len(chunk)
                # Обновляем прогресс загрузки
                if total_size > 0:
                    upload_progress = int((file_size_total / total_size) * 100)
                    tracker.update_stage(
                        ProcessingStage.UPLOAD,
                        progress=upload_progress,
                        message=f"Загрузка файла: {file_size_total / 1024:.1f} КБ из {total_size / 1024:.1f} КБ",
                    )

        file_digest = file_hash.hexdigest()
        file_size = os.path.getsize(dst)
        file_mtime = os.path.getmtime(dst)
        tracker.complete_stage(ProcessingStage.UPLOAD, "Файл загружен успешно")
    except Exception as exc:  # pragma: no cover - unexpected IO failure
        logger.exception("Failed to write uploaded file")
        tracker.set_error(ProcessingStage.UPLOAD, f"Ошибка загрузки: {exc}")
        raise HTTPException(
            status_code=500, detail=f"Не удалось сохранить файл: {exc}"
        ) from exc

    file_type_label = FILE_TYPE_LABELS.get(file_ext, "Неизвестный тип")

    # Валидация
    tracker.update_stage(
        ProcessingStage.VALIDATION, progress=50, message="Проверка файла..."
    )

    existing_upload = database.find_duplicate_upload(
        enterprise_id=enterprise["id"],
        filename=safe_filename,
        file_size=file_size,
        file_hash=file_digest,
    )

    logger.info(
        f"🔍 Проверка дубликатов: enterprise_id={enterprise['id']}, "
        f"filename={safe_filename}, size={file_size}, hash={file_digest[:16]}..."
    )
    if existing_upload:
        logger.info(f"✅ Дубликат найден: batch_id={existing_upload.get('batch_id')}")
    else:
        logger.info("ℹ️ Дубликат не найден, создаем новую загрузку")

    if existing_upload:
        existing_batch_id = existing_upload["batch_id"]
        
        # 🔧 ПЕРЕКЛЮЧАТЕЛЬ РЕЖИМОВ: debug / production
        # Приоритет: параметр из формы > переменная окружения > production
        current_mode = system_mode.lower() if system_mode else SYSTEM_MODE
        if current_mode not in ["debug", "production"]:
            current_mode = SYSTEM_MODE
        
        if current_mode == "debug":
            # РЕЖИМ ОТЛАДКИ: всегда переобрабатывать
            logger.info(
                f"🔧 [DEBUG MODE] Дубликат найден (batch_id={existing_batch_id}), "
                f"переобрабатываем файл (режим отладки, установлен через {'веб-интерфейс' if system_mode else 'переменную окружения'})"
            )
            deleted = database.delete_upload_by_batch_id(existing_batch_id)
            if deleted:
                logger.info("✅ Старая запись удалена, обрабатываем файл заново")
            # Продолжаем обработку
        else:
            # РЕЖИМ РАБОТЫ: проверяем изменения по hash
            existing_hash = existing_upload.get("file_hash", "")
            
            if existing_hash == file_digest:
                # Hash совпадает - файл не изменился, пропускаем
                logger.info(
                    f"✅ Найден дубликат {safe_filename} для предприятия {enterprise['name']} "
                    f"(batch_id={existing_batch_id}). Файл не изменился (hash совпадает), пропускаем обработку"
                )
                os.remove(dst)
                tracker.complete_stage(
                    ProcessingStage.VALIDATION, "Файл уже был загружен ранее (без изменений)"
                )
                tracker.complete()
                
                # Очищаем трекер для дубликата
                import threading
                def cleanup_tracker():
                    import time
                    time.sleep(60)
                    remove_progress_tracker(batch_id)
                threading.Thread(target=cleanup_tracker, daemon=True).start()
                
                return {
                    "batch_id": existing_batch_id,
                    "saved": existing_upload["filename"],
                    "file_type": existing_upload["file_type"],
                    "file_size": existing_upload["file_size"],
                    "parsing_status": existing_upload["status"],
                    "enterprise": {"id": enterprise["id"], "name": enterprise["name"]},
                    "parsing_summary": existing_upload.get("parsing_summary"),
                    "duplicate": True,
                    "skipped": True,
                    "reason": "Файл не изменился (hash совпадает)"
                }
            else:
                # Hash отличается - файл обновлен, переобрабатываем
                logger.info(
                    f"🔄 Дубликат найден, но файл обновлен (hash изменился). "
                    f"Переобрабатываем файл (batch_id={existing_batch_id})"
                )
                deleted = database.delete_upload_by_batch_id(existing_batch_id)
                if deleted:
                    logger.info("✅ Старая запись удалена, обрабатываем обновленный файл")
                # Продолжаем обработку

    tracker.complete_stage(ProcessingStage.VALIDATION, "Валидация пройдена")

    parsing_result: Optional[Dict[str, Any]] = None
    parsing_error: Optional[str] = None
    try:
        # Проверяем, не была ли отменена обработка перед началом парсинга
        if tracker.is_cancelled():
            os.remove(dst)  # Удаляем файл при отмене
            raise HTTPException(
                status_code=499, detail="Обработка файла отменена пользователем"
            )

        logger.info("Начинаю парсинг файла: %s", dst)
        tracker.update_stage(
            ProcessingStage.PARSING, progress=10, message="Начало парсинга файла..."
        )

        # Для больших файлов обновляем прогресс парсинга
        if file_type_enum == FileType.EXCEL:
            tracker.update_stage(
                ProcessingStage.PARSING, progress=30, message="Чтение листов Excel..."
            )
        elif file_type_enum == FileType.PDF:
            tracker.update_stage(
                ProcessingStage.PARSING,
                progress=30,
                message="Извлечение текста из PDF...",
            )
            # Для PDF может потребоваться OCR, поэтому заранее инициализируем этап
            tracker.update_stage(
                ProcessingStage.OCR, progress=0, message="Проверка необходимости OCR..."
            )
        elif file_type_enum == FileType.IMAGE:
            tracker.update_stage(
                ProcessingStage.PARSING, progress=50, message="Загрузка изображения..."
            )

        # Передаем batch_id в parse_file для проверки отмены
        try:
            parsing_result = parse_file(dst, batch_id=batch_id)
        except InterruptedError as cancel_exc:
            # Обработка была отменена пользователем
            logger.info(f"Парсинг файла отменен: {cancel_exc}")
            tracker.set_error(
                ProcessingStage.PARSING, "Обработка отменена пользователем"
            )
            os.remove(dst)  # Удаляем файл при отмене
            raise HTTPException(
                status_code=499, detail="Обработка файла отменена пользователем"
            )

        # Проверяем, использовался ли OCR для PDF
        pdf_data = parsing_result.get("data", {})
        if file_type_enum == FileType.PDF:
            is_scanned = pdf_data.get("is_scanned", False)
            ocr_attempted = pdf_data.get("ocr_attempted", False)
            ocr_success = pdf_data.get("ocr_success", False)
            ocr_error = pdf_data.get("ocr_error")

            if is_scanned:
                if ocr_attempted:
                    if ocr_success:
                        # OCR успешно применен
                        char_count = pdf_data.get("total_characters", 0)
                        tracker.update_stage(
                            ProcessingStage.OCR,
                            progress=100,
                            message=f"OCR распознавание завершено ({char_count} символов)",
                            metadata={
                                "char_count": char_count,
                                "pages": pdf_data.get("metadata", {}).get(
                                    "num_pages", 0
                                ),
                            },
                        )
                        tracker.complete_stage(
                            ProcessingStage.OCR, "OCR успешно применен"
                        )
                    else:
                        # OCR не удался (poppler не установлен или другая ошибка)
                        if ocr_error == "poppler_not_installed":
                            tracker.update_stage(
                                ProcessingStage.OCR,
                                progress=0,
                                message="OCR недоступен: poppler не установлен. Файл обработан без OCR.",
                                metadata={
                                    "error": "poppler_not_installed",
                                    "is_scanned": True,
                                },
                            )
                        else:
                            tracker.update_stage(
                                ProcessingStage.OCR,
                                progress=0,
                                message="OCR не извлек текст. Файл обработан без OCR.",
                                metadata={"error": "ocr_failed", "is_scanned": True},
                            )
                else:
                    # Сканированный документ, но OCR не был попытка (HAS_OCR = False)
                    tracker.update_stage(
                        ProcessingStage.OCR,
                        progress=0,
                        message="Обнаружен сканированный PDF, но OCR библиотеки не установлены.",
                        metadata={"is_scanned": True, "error": "ocr_not_available"},
                    )

        tracker.complete_stage(ProcessingStage.PARSING, "Парсинг завершен")
        status_value = "success" if parsing_result.get("parsed") else "partial"

        # Добавляем информацию о AI-анализе в результаты
        parsing_result_data = parsing_result.copy()
        if "ai_analysis" in parsing_result:
            # Добавляем краткую сводку AI-анализа
            ai_analysis = parsing_result.get("ai_analysis", {})
            parsing_result_data["ai_analysis_summary"] = {
                "confidence_score": ai_analysis.get("confidence_score", 0.0),
                "is_valid": ai_analysis.get("summary", {}).get("is_valid", True),
                "has_anomalies": ai_analysis.get("summary", {}).get(
                    "has_anomalies", False
                ),
                "is_compliant": ai_analysis.get("summary", {}).get(
                    "is_compliant", True
                ),
                "anomaly_count": ai_analysis.get("anomalies", {}).get(
                    "anomaly_count", 0
                ),
                "efficiency_class": ai_analysis.get("summary", {}).get(
                    "efficiency_class", "N/A"
                ),
            }

        # КРИТИЧЕСКОЕ ИЗМЕНЕНИЕ: Определяем тип ресурса ПОСЛЕ парсинга
        # на основе содержимого файла
        from utils.resource_classifier import ResourceClassifier

        # Формируем raw_json для анализа
        raw_json_for_analysis = {
            "file_type": file_type_label.lower(),
            "parsing": parsing_result_data,
        }

        # Определяем тип ресурса с приоритетом содержимого
        detected_resource_type = ResourceClassifier.classify(
            filename=safe_filename,
            raw_json=raw_json_for_analysis if parsing_result.get("parsed") else None,
            user_provided_type=user_provided_type,
        )

        # Обновляем resource_type на основе анализа содержимого
        resource_type = detected_resource_type
        logger.info(
            f"Тип ресурса определен после парсинга для {safe_filename}: {resource_type} "
            f"(пользователь указал: {user_provided_type or 'не указан'})"
        )

        # 🧠 ИНТЕЛЛЕКТУАЛЬНЫЙ МАРШРУТИЗАТОР: Анализ файла и определение оптимального пути обработки
        routing_map: Optional[Dict[str, Any]] = None
        try:
            router = IntelligentRouter()
            raw_json_for_routing = {
                "file_type": file_type_label.lower(),
                "parsing": parsing_result_data if parsing_result.get("parsed") else None,
            }
            routing_map = router.analyze_file(
                file_path=dst,
                filename=safe_filename,
                raw_json=raw_json_for_routing if parsing_result.get("parsed") else None,
                fast_mode=True
            )
            
            # Если уверенность низкая, выполняем глубокий анализ
            if routing_map.get("analysis", {}).get("confidence", 0.0) < 0.7:
                logger.info(f"⚠️ Низкая уверенность анализа ({routing_map.get('analysis', {}).get('confidence', 0.0):.2f}), выполняем глубокий анализ")
                routing_map = router.analyze_file(
                    file_path=dst,
                    filename=safe_filename,
                    raw_json=raw_json_for_routing if parsing_result.get("parsed") else None,
                    fast_mode=False
                )
            
            logger.info(
                f"🧠 Intelligent Router: document_type={routing_map.get('analysis', {}).get('document_type', 'unknown')}, "
                f"resource_type={routing_map.get('analysis', {}).get('resource_type', 'unknown')}, "
                f"data_type={routing_map.get('analysis', {}).get('data_type', 'unknown')}, "
                f"confidence={routing_map.get('analysis', {}).get('confidence', 0.0):.2f}, "
                f"primary_module={routing_map.get('routing', {}).get('primary_module', 'unknown')}"
            )
            
            # Сохраняем routing_map в parsing_result для дальнейшего использования
            if parsing_result:
                parsing_result["routing_map"] = routing_map
            
        except Exception as router_exc:
            logger.warning(f"⚠️ Ошибка при работе Intelligent Router: {router_exc}", exc_info=True)
            # Продолжаем обработку без routing_map

        parsing_results_cache[batch_id] = {
            "batch_id": batch_id,
            "filename": safe_filename,
            "file_path": dst,
            "file_type": file_type_label,
            "file_size": file_size,
            "parsing": parsing_result_data,
            "status": status_value,
        }
        if resource_type:
            parsing_results_cache[batch_id]["resource_type"] = resource_type
        logger.info(
            "Парсинг завершен для batch_id=%s, parsed=%s, resource_type=%s",
            batch_id,
            parsing_result.get("parsed"),
            resource_type,
        )
    except Exception as exc:  # pragma: no cover - parse failure path
        parsing_error = str(exc)
        logger.exception("Ошибка при парсинге файла %s", dst)
        tracker.set_error(ProcessingStage.PARSING, f"Ошибка парсинга: {exc}")
        parsing_results_cache[batch_id] = {
            "batch_id": batch_id,
            "filename": safe_filename,
            "file_path": dst,
            "file_type": file_type_label,
            "file_size": file_size,
            "parsing": None,
            "error": parsing_error,
            "status": "error",
        }
        if resource_type:
            parsing_results_cache[batch_id]["resource_type"] = resource_type

    parsing_summary = build_parsing_summary(parsing_result, file_ext)
    parsing_status = (
        "success"
        if parsing_result and parsing_result.get("parsed")
        else ("error" if parsing_error else "pending")
    )

    response_data: Dict[str, Any] = {
        "batch_id": batch_id,
        "saved": os.path.basename(dst),
        "file_type": file_type_label,
        "file_size": file_size,
        "parsing_status": parsing_status,
        "enterprise": {"id": enterprise["id"], "name": enterprise["name"]},
    }
    if resource_type:
        response_data["resource_type"] = resource_type
    if parsing_summary:
        response_data["parsing_summary"] = parsing_summary
    if routing_map:
        response_data["routing_map"] = routing_map

    try:
        summary_payload = dict(parsing_summary or {})
        if resource_type:
            summary_payload["resource_type"] = resource_type
            summary_payload["resource_type_label"] = RESOURCE_LABELS.get(
                resource_type, resource_type
            )

        # Добавляем информацию о готовом Word-отчете
        if parsing_result and file_type_enum == FileType.WORD:
            word_data = parsing_result.get("data", {})
            if word_data.get("is_ready_report", False):
                summary_payload["is_ready_report"] = True
                summary_payload["table_count"] = word_data.get("table_count", 0)
                summary_payload["report_type"] = "ready_word_report"

        # Добавляем информацию об использовании OCR для PDF
        if parsing_result and file_type_enum == FileType.PDF:
            # Используем pdf_data, так как она определена выше
            summary_payload["ocr_used"] = pdf_data.get("ocr_used", False)
            summary_payload["is_scanned"] = pdf_data.get("is_scanned", False)
            summary_payload["ocr_success"] = pdf_data.get("ocr_success", False)

        # Добавляем routing_map в summary_payload для сохранения в БД
        if routing_map:
            summary_payload["routing_map"] = {
                "document_type": routing_map.get("analysis", {}).get("document_type"),
                "resource_type": routing_map.get("analysis", {}).get("resource_type"),
                "data_type": routing_map.get("analysis", {}).get("data_type"),
                "period": routing_map.get("analysis", {}).get("period"),
                "confidence": routing_map.get("analysis", {}).get("confidence"),
                "primary_module": routing_map.get("routing", {}).get("primary_module"),
                "target_tables": routing_map.get("routing", {}).get("target_tables", []),
            }

        if summary_payload:
            response_data["parsing_summary"] = summary_payload

        database.create_upload(
            batch_id=batch_id,
            enterprise_id=enterprise["id"],
            filename=safe_filename,
            file_type=file_type_label,
            file_size=file_size,
            status=parsing_status,
            parsing_summary=summary_payload if summary_payload else None,
            file_hash=file_digest,
            file_mtime=file_mtime,
        )

        # Проверяем, является ли файл готовым Word-отчетом
        # Для готовых отчетов пропускаем агрегацию и специализированный парсинг
        is_ready_word_report = False
        if parsing_result and file_type_enum == FileType.WORD:
            is_ready_word_report = parsing_result.get("data", {}).get(
                "is_ready_report", False
            )
            if is_ready_word_report:
                logger.info(
                    f"Обнаружен готовый Word отчет, пропускаем агрегацию и специализированный парсинг (таблиц: {parsing_result.get('data', {}).get('table_count', 0)})"
                )

        aggregation_info: Optional[Dict[str, Any]] = None
        equipment_info: Optional[Dict[str, Any]] = None
        envelope_info: Optional[Dict[str, Any]] = None
        nodes_info: Optional[Dict[str, Any]] = None

        # Используем raw_json для улучшенного определения типа (нужно ПЕРЕД проверкой специализированных файлов)
        raw_json_for_check = (
            {"file_type": file_type_label.lower(), "parsing": parsing_result_data}
            if parsing_result.get("parsed")
            else None
        )

        # Для готовых Word-отчетов пропускаем агрегацию
        # ВАЖНО: Не агрегируем файлы узлов учёта, оборудования и других специализированных типов
        is_specialized_file = (
            is_nodes_file(safe_filename, raw_json_for_check)
            or is_equipment_file(safe_filename, raw_json_for_check)
            or is_envelope_file(safe_filename, raw_json_for_check)
        )
        
        if is_specialized_file:
            logger.info(
                f"📋 Файл {safe_filename} определен как специализированный (nodes/equipment/envelope), "
                f"пропускаем агрегацию энергоресурсов"
            )
        
        if should_aggregate_file(safe_filename) and not is_ready_word_report and not is_specialized_file:
            try:
                # Проверяем, не была ли отменена обработка
                if tracker.is_cancelled():
                    raise InterruptedError("Обработка файла отменена пользователем")

                tracker.update_stage(
                    ProcessingStage.AGGREGATION,
                    progress=10,
                    message="Агрегация данных...",
                )
                aggregation_data = aggregate_energy_data(dst)

                # Если агрегация из файла не удалась, пробуем агрегировать из распарсенных данных (БД)
                # НО только если это не специализированный файл
                if (
                    not aggregation_data
                    and parsing_result
                    and parsing_result.get("parsed")
                    and not is_specialized_file
                ):
                    logger.info(
                        "⚠️ [DIAG] Агрегация из файла не удалась, пробую агрегировать из распарсенных данных"
                    )
                    raw_json_for_aggregation = {
                        "batch_id": batch_id,
                        "filename": safe_filename,
                        "file_type": file_type_label.lower(),
                        "parsing": parsing_result_data,
                    }
                    aggregation_data = aggregate_from_db_json(raw_json_for_aggregation)
                    if aggregation_data:
                        logger.info(
                            "✅ [DIAG] Агрегация из распарсенных данных успешна"
                        )
                        logger.info(
                            f"📊 [DIAG] Структура агрегированных данных: resources={list(aggregation_data.get('resources', {}).keys())}"
                        )
                    else:
                        logger.warning(
                            "⚠️ [DIAG] Агрегация из распарсенных данных также не удалась"
                        )

                if aggregation_data:
                    # Агрегируем данные по категориям использования (by_usage)
                    tracker.update_stage(
                        ProcessingStage.AGGREGATION,
                        progress=50,
                        message="Обработка категорий использования...",
                    )

                    # Пытаемся загрузить данные оборудования для сверки
                    equipment_data_for_categories = None
                    equipment_json_path = AGGREGATED_DIR / f"{batch_id}_equipment.json"
                    if equipment_json_path.exists():
                        try:
                            equipment_data_for_categories = json.loads(
                                equipment_json_path.read_text(encoding="utf-8")
                            )
                            logger.info(
                                f"Загружены данные оборудования для сверки категорий: {equipment_json_path}"
                            )
                        except Exception as eq_exc:
                            logger.warning(
                                f"Не удалось загрузить данные оборудования для сверки: {eq_exc}"
                            )

                    usage_categories = aggregate_usage_categories(
                        dst, equipment_data=equipment_data_for_categories
                    )
                    if usage_categories:
                        # Сохраняем usage_categories.json (ожидается валидатором)
                        usage_path = AGGREGATED_DIR / "usage_categories.json"
                        try:
                            usage_path.write_text(
                                json.dumps(
                                    usage_categories, ensure_ascii=False, indent=2
                                ),
                                encoding="utf-8",
                            )
                            logger.info(
                                f"Сохранены категории использования: {usage_path}"
                            )
                        except Exception as usage_save_exc:
                            logger.warning(
                                f"Не удалось сохранить usage_categories.json: {usage_save_exc}"
                            )

                        # Распределяем категории по кварталам и добавляем в aggregation_data
                        logger.info("📊 [DIAG] Данные агрегации ДО распределения по категориям: "
                                    f"keys={list(aggregation_data.get('resources', {}).keys())}")
                        
                        aggregation_data = distribute_categories_by_quarter(
                            aggregation_data, usage_categories
                        )
                        
                        # Ключевая проверка после распределения
                        electricity_data_after = aggregation_data.get("resources", {}).get("electricity")
                        if electricity_data_after:
                            logger.info("✅ [DIAG] Данные по электроэнергии ПРИСУТСТВУЮТ после распределения по категориям. "
                                        f"Кол-во записей: {len(electricity_data_after)}")
                        else:
                            logger.warning("⚠️ [DIAG] Данные по электроэнергии ОТСУТСТВУЮТ после распределения по категориям.")
                        
                        logger.info("Категории использования распределены по кварталам")

                    aggregated_file = write_aggregation_json(
                        batch_id, aggregation_data, AGGREGATED_DIR
                    )
                    log_aggregation_event(
                        batch_id=batch_id,
                        source_file=dst,
                        output_file=aggregated_file,
                        status="success",
                    )
                    
                    # Импорт всех ресурсов в БД
                    logger.info("=" * 70)
                    logger.info(f"🚀 НАЧАЛО ИМПОРТА В БД для batch_id={batch_id}")
                    logger.info("=" * 70)
                    try:
                        resources = aggregation_data.get("resources", {})
                        logger.info(f"📦 Найдено ресурсов для импорта: {list(resources.keys())}")
                        imported_total = 0
                        
                        # Импортируем все доступные ресурсы
                        for resource_type, resource_data in resources.items():
                            if resource_data:
                                logger.info(f"📥 Импортирую ресурс: {resource_type}")
                                logger.info(f"   Количество периодов: {len(resource_data)}")
                                
                                imported_records = database.import_resource_to_db(
                                    enterprise_id=enterprise["id"],
                                    batch_id=batch_id,
                                    resource_type=resource_type,
                                    resource_data=resource_data,
                                )
                                imported_total += len(imported_records)
                                logger.info(
                                    f"✅ Импортировано {len(imported_records)} записей {resource_type} в БД для batch_id={batch_id}"
                                )
                                # Детальная информация о каждой записи
                                for record in imported_records[:5]:  # Показываем первые 5
                                    logger.info(f"   → Запись: {record.get('resource_type')} / {record.get('period')} (ID: {record.get('id')})")
                                if len(imported_records) > 5:
                                    logger.info(f"   ... и еще {len(imported_records) - 5} записей")
                            else:
                                logger.warning(f"⚠️ Ресурс {resource_type} пустой, пропускаю")
                        
                        if imported_total > 0:
                            logger.info("=" * 70)
                            logger.info(
                                f"✅ ИМПОРТ ЗАВЕРШЕН: Всего импортировано {imported_total} записей ресурсов в БД для batch_id={batch_id}"
                            )
                            logger.info("=" * 70)
                        else:
                            logger.warning(
                                f"⚠️ Данные ресурсов отсутствуют в агрегированных данных для batch_id={batch_id}"
                            )
                    except Exception as import_exc:
                        logger.error("=" * 70)
                        logger.error(
                            f"❌ ОШИБКА ПРИ ИМПОРТЕ РЕСУРСОВ В БД для batch_id={batch_id}: {import_exc}",
                            exc_info=True,
                        )
                        logger.error("=" * 70)
                    
                    aggregation_info = {
                        "output_file": aggregated_file.name,
                        "missing_sheets": aggregation_data.get("missing_sheets", []),
                        "resource_quarters": {
                            resource: len(entries)
                            for resource, entries in aggregation_data[
                                "resources"
                            ].items()
                        },
                    }
                    tracker.complete_stage(
                        ProcessingStage.AGGREGATION, "Агрегация данных завершена"
                    )
                else:
                    log_aggregation_event(
                        batch_id=batch_id,
                        source_file=dst,
                        output_file=None,
                        status="skipped",
                        message="Aggregator returned no data",
                    )
            except InterruptedError:
                # Обработка была отменена
                logger.info("Агрегация данных отменена пользователем")
                raise
            except Exception as agg_exc:  # pragma: no cover - aggregation failure path
                logger.exception("Aggregation failed for %s", dst)
                log_aggregation_event(
                    batch_id=batch_id,
                    source_file=dst,
                    output_file=None,
                    status="error",
                    message=str(agg_exc),
                )
        # Для готовых Word-отчетов пропускаем специализированный парсинг
        # raw_json_for_check уже определен выше

        if (
            is_equipment_file(safe_filename, raw_json_for_check)
            and not is_ready_word_report
        ):
            try:
                # Проверяем, не была ли отменена обработка
                if tracker.is_cancelled():
                    raise InterruptedError("Обработка файла отменена пользователем")

                tracker.update_stage(
                    ProcessingStage.SPECIALIZED_PARSING,
                    progress=10,
                    message="Парсинг оборудования...",
                )
                equipment_data = parse_equipment_workbook(dst)
                if equipment_data:
                    equipment_file = write_equipment_json(
                        batch_id, equipment_data, AGGREGATED_DIR
                    )
                    equipment_info = {
                        "output_file": equipment_file.name,
                        "sections": sum(
                            len(sheet["sections"])
                            for sheet in equipment_data.get("sheets", [])
                        ),
                        "total_items": equipment_data.get("summary", {}).get(
                            "total_items"
                        ),
                        "total_power_kw": equipment_data.get("summary", {}).get(
                            "total_power_kw"
                        ),
                    }
                    tracker.complete_stage(
                        ProcessingStage.SPECIALIZED_PARSING,
                        "Парсинг оборудования завершен",
                    )
            except InterruptedError:
                # Обработка была отменена
                logger.info("Парсинг оборудования отменен пользователем")
                raise
            except (
                Exception
            ) as equipment_exc:  # pragma: no cover - equipment parsing failure
                logger.exception("Equipment parsing failed for %s", dst)
                equipment_info = {
                    "error": str(equipment_exc),
                }
        # Для готовых Word-отчетов пропускаем специализированный парсинг
        # Парсинг файла расчета теплопотерь по зданиям
        if (
            is_envelope_file(safe_filename, raw_json_for_check)
            and not is_ready_word_report
        ):
            try:
                # Проверяем, не была ли отменена обработка
                if tracker.is_cancelled():
                    raise InterruptedError("Обработка файла отменена пользователем")

                envelope_data = parse_building_envelope(dst)
                if envelope_data:
                    envelope_file = write_envelope_json(
                        batch_id, envelope_data, AGGREGATED_DIR
                    )
                    envelope_info = {
                        "output_file": envelope_file.name,
                        "sections": envelope_data.get("summary", {}).get(
                            "total_sections"
                        ),
                        "total_area_m2": envelope_data.get("summary", {}).get(
                            "total_area_m2"
                        ),
                        "total_heat_loss": envelope_data.get("summary", {}).get(
                            "total_heat_loss"
                        ),
                    }
            except InterruptedError:
                # Обработка была отменена
                logger.info(
                    "Парсинг расчета теплопотерь по зданиям отменен пользователем"
                )
                raise
            except (
                Exception
            ) as envelope_exc:  # pragma: no cover - envelope parsing failure
                logger.exception(
                    "Ошибка парсинга расчета теплопотерь по зданиям для %s", dst
                )
                envelope_info = {
                    "error": str(envelope_exc),
                }
        # Для готовых Word-отчетов пропускаем специализированный парсинг
        if (
            is_nodes_file(safe_filename, raw_json_for_check)
            and not is_ready_word_report
        ):
            try:
                # Проверяем, не была ли отменена обработка
                if tracker.is_cancelled():
                    raise InterruptedError("Обработка файла отменена пользователем")

                nodes_data = parse_nodes_workbook(dst)
                if nodes_data:
                    nodes_file = write_nodes_json(batch_id, nodes_data, AGGREGATED_DIR)
                    nodes_info = {
                        "output_file": nodes_file.name,
                        "total_nodes": nodes_data.get("summary", {}).get("total_nodes"),
                    }
                    
                    # Импорт данных потребления по узлам учёта в БД
                    try:
                        # Преобразуем данные узлов в формат для импорта
                        # Примечание: в текущих данных узлов нет информации о периодах,
                        # поэтому создаем базовые записи. Данные потребления по периодам
                        # нужно будет получать из актов балансов или других источников
                        node_consumption_records = []
                        nodes_list = nodes_data.get("nodes", [])
                        
                        for node in nodes_list:
                            node_name = node.get("name")
                            if not node_name:
                                continue
                            
                            # Если есть агрегированные данные электроэнергии, пытаемся связать
                            # Пока создаем запись с базовыми данными узла
                            # TODO: Добавить логику для извлечения данных потребления по периодам
                            # из актов балансов или других источников
                            
                            node_consumption_records.append({
                                "node_name": node_name,
                                "period": "unknown",  # Будет обновлено при получении данных по периодам
                                "active_energy_kwh": node.get("active_energy_p"),
                                "reactive_energy_kvarh": node.get("reactive_energy_q"),
                                "cost_sum": None,  # Нужно получить из других источников
                                "data_json": {
                                    "tt": node.get("tt"),
                                    "coefficient": node.get("coefficient"),
                                    "seal_date": node.get("seal_date"),
                                    "supplier_seal_date": node.get("supplier_seal_date"),
                                    "note": node.get("note"),
                                }
                            })
                        
                        if node_consumption_records:
                            imported_nodes = database.import_node_consumption_to_db(
                                enterprise_id=enterprise["id"],
                                batch_id=batch_id,
                                node_consumption_data=node_consumption_records,
                            )
                            logger.info(
                                f"✅ Импортировано {len(imported_nodes)} записей потребления по узлам учёта "
                                f"в БД для batch_id={batch_id}"
                            )
                            nodes_info["imported_to_db"] = len(imported_nodes)
                    except Exception as import_exc:
                        logger.warning(
                            f"⚠️ Не удалось импортировать данные потребления по узлам в БД: {import_exc}"
                        )
                        # Не прерываем обработку, если импорт не удался
            except InterruptedError:
                # Обработка была отменена
                logger.info("Парсинг узлов учета отменен пользователем")
                raise
            except Exception as nodes_exc:  # pragma: no cover - nodes parsing failure
                logger.exception("Nodes parsing failed for %s", dst)
                nodes_info = {
                    "error": str(nodes_exc),
                }
        
        # Обработка актов балансов для извлечения данных потребления по узлам учёта
        # Согласно рекомендациям экспертов:
        # - Software Engineer: "Использовать OCR для извлечения из PDF актов"
        # - ML Engineer: "Использовать Gemini Vision (95% confidence)"
        # - QA Engineer: "Тестировать на нескольких файлах сначала"
        logger.info(f"🔍 Проверка файла на акт баланса: {safe_filename}")
        is_balance_sheet = is_balance_sheet_file(safe_filename, raw_json_for_check)
        if is_balance_sheet:
            logger.info(f"✅ Файл определен как акт баланса: {safe_filename}")
        else:
            logger.info(f"❌ Файл НЕ определен как акт баланса: {safe_filename} (проверка завершена)")
        
        if is_balance_sheet and not is_ready_word_report:
            try:
                # Проверяем, не была ли отменена обработка
                if tracker.is_cancelled():
                    raise InterruptedError("Обработка файла отменена пользователем")
                
                logger.info(f"📋 Обнаружен акт баланса: {safe_filename}")
                
                # Извлекаем данные по узлам учёта из акта баланса
                # (может быть потребление, производство или реализация электроэнергии)
                node_consumption_data = extract_node_consumption_from_balance_sheet(
                    file_path=dst,
                    batch_id=batch_id,
                    enterprise_id=enterprise["id"],
                    raw_json=raw_json_for_check
                )
                
                if node_consumption_data:
                    # Определяем тип данных для логирования
                    data_types = set(record.get("data_type", "consumption") for record in node_consumption_data)
                    data_type_label = {
                        "consumption": "потребления",
                        "production": "производства",
                        "realization": "реализации (произведенной и проданной)"
                    }
                    type_labels = [data_type_label.get(dt, dt) for dt in data_types]
                    
                    # Импортируем данные в БД
                    imported_nodes = database.import_node_consumption_to_db(
                        enterprise_id=enterprise["id"],
                        batch_id=batch_id,
                        node_consumption_data=node_consumption_data,
                    )
                    
                    logger.info(
                        f"✅ Импортировано {len(imported_nodes)} записей данных {', '.join(type_labels)} по узлам учёта "
                        f"из акта баланса {safe_filename} в БД для batch_id={batch_id}"
                    )
                    
                    # Добавляем информацию в ответ
                    if "balance_sheet" not in response_data:
                        response_data["balance_sheet"] = {}
                    response_data["balance_sheet"]["nodes_imported"] = len(imported_nodes)
                    response_data["balance_sheet"]["file"] = safe_filename
                else:
                    logger.info(
                        f"ℹ️ Данные потребления по узлам не найдены в акте баланса: {safe_filename}"
                    )
                    
            except InterruptedError:
                # Обработка была отменена
                logger.info("Обработка акта баланса отменена пользователем")
                raise
            except Exception as balance_exc:
                logger.warning(
                    f"⚠️ Ошибка обработки акта баланса {safe_filename}: {balance_exc}"
                )
                # Не прерываем обработку, если обработка акта баланса не удалась
        
        if aggregation_info:
            response_data["aggregation"] = aggregation_info
        if equipment_info:
            response_data["equipment"] = equipment_info
        if envelope_info:
            response_data["envelope"] = envelope_info
        if nodes_info:
            response_data["nodes"] = nodes_info

        # Сохранение данных
        # Проверяем, не была ли отменена обработка перед сохранением
        if tracker.is_cancelled():
            os.remove(dst)  # Удаляем файл при отмене
            raise HTTPException(
                status_code=499, detail="Обработка файла отменена пользователем"
            )

        tracker.update_stage(
            ProcessingStage.SAVING, progress=50, message="Сохранение данных..."
        )
        editable_text = build_editable_text(parsing_result)
        
        # Проверяем, что данные есть в кэше перед сохранением
        if batch_id not in parsing_results_cache:
            logger.error(f"⚠️ Данные парсинга отсутствуют в кэше для batch_id={batch_id}")
            # Создаем минимальную структуру из доступных данных
            parsing_results_cache[batch_id] = {
                "batch_id": batch_id,
                "filename": safe_filename,
                "file_path": dst,
                "file_type": file_type_label,
                "file_size": file_size,
                "parsing": parsing_result_data if parsing_result else None,
                "status": parsing_status,
                "error": "Данные парсинга не были сохранены в кэш",
            }
        
        try:
            database.save_parsed_content(
                batch_id,
                raw_json=parsing_results_cache[batch_id],
                editable_text=editable_text,
            )
            logger.info(f"✅ Данные парсинга успешно сохранены в БД для batch_id={batch_id}")
            tracker.complete_stage(ProcessingStage.SAVING, "Данные сохранены")
        except Exception as save_exc:
            logger.error(
                f"❌ Ошибка сохранения данных парсинга в БД для batch_id={batch_id}: {save_exc}",
                exc_info=True
            )
            tracker.set_error(ProcessingStage.SAVING, f"Ошибка сохранения: {save_exc}")
            # Не прерываем обработку, но логируем ошибку
        history = database.list_uploads_for_enterprise(enterprise["id"])
        response_data["history"] = history

        # Завершаем обработку
        tracker.complete()

        # Очищаем трекер через 5 минут (даем время для опроса прогресса)
        # В продакшене лучше использовать фоновую задачу
        import threading

        def cleanup_tracker():
            import time

            time.sleep(300)  # 5 минут
            remove_progress_tracker(batch_id)

        threading.Thread(target=cleanup_tracker, daemon=True).start()

    except InterruptedError:
        # Обработка была отменена пользователем
        logger.info("Обработка файла отменена пользователем")
        if os.path.exists(dst):
            try:
                os.remove(dst)  # Удаляем файл при отмене
            except Exception as e:
                logger.warning(f"Не удалось удалить файл при отмене: {e}")
        raise HTTPException(
            status_code=499, detail="Обработка файла отменена пользователем"
        )
    except Exception as exc:  # pragma: no cover - DB failure path
        logger.exception(
            f"Ошибка при сохранении метаданных загрузки: {exc}",
            exc_info=True
        )
        # Логируем детали для отладки
        logger.error(f"  batch_id: {batch_id}")
        logger.error(f"  filename: {safe_filename if 'safe_filename' in locals() else 'unknown'}")
        logger.error(f"  enterprise_id: {enterprise['id'] if 'enterprise' in locals() else 'unknown'}")
        logger.error(f"  file_type: {file_type_label if 'file_type_label' in locals() else 'unknown'}")
        
        if batch_id:
            tracker = get_progress_tracker(batch_id)
            if tracker:
                tracker.set_error(ProcessingStage.SAVING, f"Ошибка сохранения: {exc}")
        
        # Возвращаем более детальное сообщение об ошибке
        error_detail = f"Ошибка сохранения загрузки: {str(exc)}"
        if len(error_detail) > 200:  # Ограничиваем длину для безопасности
            error_detail = error_detail[:200] + "..."
        
        raise HTTPException(
            status_code=500, detail=error_detail
        ) from exc

    return response_data


@app.post("/api/generate-passport/{batch_id}")
async def generate_energy_passport(
    batch_id: str,
    template_name: str = Query(
        default="",
        description="Имя шаблона из templates_config (например, 'new_energy_passport', 'metin', 'default')",
    ),
    skip_readiness_check: str = Query(
        default="false",
        description="Пропустить проверку готовности (только для тестирования)",
    ),
):
    """
    Генерация энергетического паспорта по ПКМ №690 из данных в БД.

    Args:
        batch_id: ID загрузки
        template_name: Имя шаблона из templates_config (например, "new_energy_passport", "metin", "default")
                      Если не указано, используется дефолтный шаблон
        skip_readiness_check: Пропустить проверку готовности (только для тестирования)

    Поток:
    1. Проверка готовности данных (если не пропущена)
    2. Получить данные из БД (parsed_data.raw_json)
    3. Агрегировать поквартально
    4. Заполнить Excel шаблон
    5. Вернуть файл для скачивания
    """
    if not HAS_GENERATOR:
        raise HTTPException(
            status_code=503,
            detail="Генератор паспортов недоступен. Проверьте установку зависимостей.",
        )

    # 1. Получаем данные из БД
    upload = database.get_upload_by_batch(batch_id)
    if not upload:
        raise HTTPException(status_code=404, detail=f"Загрузка {batch_id} не найдена")

    # 2. Проверка готовности данных (если не пропущена)
    # Преобразуем строку в bool
    skip_check = skip_readiness_check.lower() in ("true", "1", "yes", "on")
    # Если template_name пустая строка или None, используем "metin" по умолчанию
    template_name_final = template_name if template_name else "metin"

    if not skip_check:
        # Feature-flag gate via EXCEL_SEMANTIC_AI_MODE
        excel_ai_mode = get_excel_semantic_mode()
        if excel_ai_mode in ("assist", "strict"):
            # Try to load canonical from stored raw_json or reconstruct
            canonical = None
            raw_json = upload.get("raw_json") or {}
            if isinstance(raw_json, dict):
                canonical_dict = raw_json.get("canonical_source")
                if isinstance(canonical_dict, dict):
                    try:
                        canonical = CanonicalSourceData.parse_obj(canonical_dict)
                    except Exception:
                        canonical = None
            if canonical is None:
                # Attempt reconstruction from file path (if known) or skip
                try:
                    filename = upload.get("filename")
                    if filename:
                        canonical = collect_canonical_from_workbook(filename)
                except Exception:
                    canonical = None
            readiness_result: GenerationReadinessResult = evaluate_generation_readiness(
                canonical
            )
            if readiness_result.overall_status == "blocked":
                raise HTTPException(
                    status_code=400,
                    detail={
                        "message": "Canonical data not ready for generation",
                        "overall_status": readiness_result.overall_status,
                        "missing_required": [
                            rf.__dict__ for rf in readiness_result.missing_required
                        ],
                        "missing_optional": [
                            rf.__dict__ for rf in readiness_result.missing_optional
                        ],
                        "notes": readiness_result.notes,
                        "mode": excel_ai_mode,
                    },
                )

        enterprise_id = upload.get("enterprise_id")
        if enterprise_id:
            readiness = validate_generation_readiness(enterprise_id)

            if not readiness["ready"]:
                logger.warning(
                    f"Попытка генерации паспорта для предприятия {enterprise_id} "
                    f"при неготовности данных: {readiness.get('missing_resources', [])}"
                )
                detail = {
                    "message": "Данные не готовы для генерации энергетического паспорта",
                    "missing_resources": readiness.get("missing_resources", []),
                    "missing_files": readiness.get("missing_files", []),
                    "missing_sheet_data": readiness.get("missing_sheet_data", []),
                    "sheet_validation": readiness.get("sheet_validation", {}),
                    "completeness_score": readiness.get("completeness_score", 0.0),
                    "warnings": readiness.get("warnings", []),
                    "progress_percentage": readiness.get("progress_percentage", 0),
                    "available_resources": readiness.get("available_resources", []),
                }
                raise HTTPException(status_code=400, detail=detail)

            logger.info(
                f"Проверка готовности пройдена для предприятия {enterprise_id}: "
                f"completeness={readiness['completeness_score']:.2f}"
            )

    # 2. Агрегируем данные из всех загрузок предприятия (не только из одного batch_id)
    enterprise_id = upload.get("enterprise_id")
    logger.info(
        f"Агрегация данных для предприятия {enterprise_id} (batch_id: {batch_id})"
    )

    # Агрегируем данные из всех загрузок предприятия
    aggregated = None
    from utils.readiness_validator import _get_aggregated_data_for_enterprise

    aggregated = _get_aggregated_data_for_enterprise(enterprise_id)

    # Если не получилось агрегировать из всех загрузок, пробуем из текущего batch_id
    if not aggregated:
        raw_json = upload.get("raw_json")
        if not raw_json:
            raise HTTPException(status_code=400, detail="Данные не распарсены")

        logger.info(f"Агрегация данных из batch_id: {batch_id}")
        logger.info(f"Структура raw_json: {list(raw_json.keys())}")
        logger.info(f"file_type: {raw_json.get('file_type')}")
        logger.info(f"parsing keys: {list(raw_json.get('parsing', {}).keys())}")

        # Логируем структуру parsing.data для отладки
        parsing_data = raw_json.get("parsing", {}).get("data", {})
        if parsing_data:
            logger.info(f"parsing.data keys: {list(parsing_data.keys())}")
            if "sheets" in parsing_data:
                sheets = parsing_data.get("sheets", [])
                logger.info(f"Количество листов: {len(sheets)}")
                for sheet in sheets[:5]:  # Первые 5 листов
                    sheet_name = sheet.get("name", "Без имени")
                    rows_count = len(sheet.get("rows", []))
                    logger.info(f"  Лист '{sheet_name}': {rows_count} строк")

        aggregated = aggregate_from_db_json(raw_json)

    # Дополнительное логирование результата агрегации
    if aggregated:
        logger.info(
            f"Агрегация успешна. Структура aggregated: {list(aggregated.keys())}"
        )
        if "resources" in aggregated:
            resources = aggregated["resources"]
            for resource_type, resource_data in resources.items():
                if (
                    resource_data
                    and isinstance(resource_data, dict)
                    and len(resource_data) > 0
                ):
                    quarters = list(resource_data.keys())
                    logger.info(
                        f"  Ресурс {resource_type}: {len(quarters)} кварталов - {quarters[:3]}..."
                    )
                    # Логируем детали по первому кварталу для отладки
                    if quarters:
                        first_quarter = quarters[0]
                        quarter_data = resource_data[first_quarter]
                        totals = quarter_data.get("quarter_totals", {})
                        logger.info(
                            f"    Квартал {first_quarter}: totals={list(totals.keys())}, values={list(totals.values())[:3]}"
                        )
                else:
                    logger.info(f"  Ресурс {resource_type}: пустой")

        # Валидация агрегированных данных перед заполнением шаблона
        logger.info("Валидация агрегированных данных перед заполнением шаблона...")
        try:
            is_valid, errors, warnings = validate_data_for_template(
                aggregated, raise_on_error=False
            )
            if not is_valid:
                logger.warning(f"Обнаружены ошибки валидации данных: {errors}")
                # Логируем предупреждения отдельно
                if warnings:
                    logger.info(f"Предупреждения валидации: {warnings}")
            else:
                logger.info("Валидация данных пройдена успешно")
                if warnings:
                    logger.info(f"Предупреждения валидации: {warnings}")
        except Exception as validation_exc:
            logger.error(
                f"Ошибка при валидации данных: {validation_exc}", exc_info=True
            )
            # Не прерываем генерацию из-за ошибок валидации, но логируем
    else:
        logger.error(f"Агрегация вернула None для предприятия {enterprise_id}")
        if "raw_json" in locals():
            logger.error(
                f"raw_json.parsing.parsed: {raw_json.get('parsing', {}).get('parsed')}"
            )
            logger.error(
                f"raw_json.parsing.sheets count: {len(raw_json.get('parsing', {}).get('sheets', []))}"
            )
        raise HTTPException(
            status_code=400,
            detail="Не удалось агрегировать данные. Проверьте структуру файла.",
        )

    # 3. Подготовка данных предприятия
    enterprise_data = {
        "id": upload.get("enterprise_id"),
        "name": upload.get("enterprise_name", "Неизвестное предприятие"),
        "inn": None,  # TODO: добавить в БД
        "address": None,
        "director_name": None,
        "industry": None,
        "reporting_year": 2024,
    }

    # 4. Генерируем паспорт
    output_dir = Path("/tmp/passports")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / f"{batch_id}_energy_passport.xlsx"

    # Определение пути к шаблону
    template_path = None
    if template_name_final:
        # Использование templates_config для выбора шаблона по имени
        try:
            templates_config_path = PROJECT_ROOT / "templates" / "pcm690"
            import sys

            if str(templates_config_path) not in sys.path:
                sys.path.insert(0, str(templates_config_path))
            from templates_config import get_template_path

            logger.info("🔍 Запрос шаблона по имени: '%s'", template_name_final)
            template_path = get_template_path(template_name_final)
            logger.info(
                "✅ Используется шаблон по имени '%s': %s",
                template_name_final,
                template_path,
            )
            if not template_path.exists():
                logger.error("❌ Файл шаблона не существует: %s", template_path)
                template_path = None
        except (ImportError, ValueError, FileNotFoundError) as e:
            logger.warning(
                "⚠️ Не удалось загрузить шаблон по имени '%s': %s. Используется дефолтный.",
                template_name_final,
                e,
            )
            logger.exception("Детали ошибки загрузки шаблона:")
            template_path = None

    # Если шаблон не выбран по имени, используем дефолтные кандидаты
    if not template_path:
        template_candidates = [
            PROJECT_ROOT
            / "data"
            / "source_files"
            / "audit_sinergys"
            / "EnergyPassport_PKM690_filled.xlsx",
            PROJECT_ROOT / "templates" / "pcm690" / "energy_passport_template.xlsx",
        ]
        template_path = next(
            (path for path in template_candidates if path.exists()), None
        )
        if not template_path:
            raise FileNotFoundError("Шаблон энергопаспорта не найден в ожидаемых путях")
        logger.info("Используется дефолтный шаблон энергопаспорта: %s", template_path)

    # Если указан template_name, ОБЯЗАТЕЛЬНО используем fill_energy_passport
    # (генератор PKM690ExcelGenerator не использует шаблоны)
    if template_name_final and not HAS_FILLER:
        raise HTTPException(
            status_code=503,
            detail=(
                f"Для использования шаблона '{template_name_final}' требуется fill_energy_passport. "
                "Модуль недоступен. Проверьте установку зависимостей."
            ),
        )

    if HAS_FILLER:
        try:
            # Копируем шаблон во временный файл для работы (защита оригинального шаблона)
            temp_template_path = output_dir / f"{batch_id}_passport_template_copy.xlsx"
            shutil.copyfile(template_path, temp_template_path)
            logger.info("Шаблон скопирован во временный файл: %s", temp_template_path)

            workbook = load_workbook(temp_template_path, data_only=False)

            resources_data = aggregated.get("resources") or aggregated
            if resources_data is None:
                resources_data = {}
            # Инициализируем все типы ресурсов
            for key in (
                "electricity",
                "gas",
                "water",
                "fuel",
                "coal",
                "heat",
                "production",
            ):
                resources_data.setdefault(key, {})

            # Логируем доступные ресурсы для отладки
            available_resources = [
                k
                for k, v in resources_data.items()
                if v and isinstance(v, dict) and len(v) > 0
            ]
            logger.info(f"Доступные ресурсы для заполнения: {available_resources}")
            logger.info(f"Всего ресурсов в данных: {list(resources_data.keys())}")

            # Логируем детали данных для каждого ресурса
            for resource_type, resource_data in resources_data.items():
                if (
                    resource_data
                    and isinstance(resource_data, dict)
                    and len(resource_data) > 0
                ):
                    logger.info(
                        f"  Ресурс {resource_type}: {len(resource_data)} кварталов"
                    )
                    for quarter, quarter_data in list(resource_data.items())[
                        :2
                    ]:  # Первые 2 квартала
                        totals = (
                            quarter_data.get("quarter_totals", {})
                            if isinstance(quarter_data, dict)
                            else {}
                        )
                        logger.info(
                            f"    {quarter}: totals keys={list(totals.keys())}, sample values={dict(list(totals.items())[:3])}"
                        )

            # (Опционально) Canonical payload for nodes/equipment when mode != off
            canonical_payload = None
            excel_ai_mode_runtime = get_excel_semantic_mode()
            if excel_ai_mode_runtime in ("assist", "strict"):
                try:
                    c = None
                    raw_json = upload.get("raw_json") or {}
                    if isinstance(raw_json, dict):
                        cdict = raw_json.get("canonical_source")
                        if isinstance(cdict, dict):
                            try:
                                c = CanonicalSourceData.parse_obj(cdict)
                            except Exception:
                                c = None
                    if c is None and upload.get("filename"):
                        c = collect_canonical_from_workbook(upload.get("filename"))
                    if c:
                        canonical_payload = canonical_to_passport_payload(c)
                        logger.info(
                            "Canonical payload prepared for nodes/equipment (mode=%s)",
                            excel_ai_mode_runtime,
                        )
                except Exception as e:
                    logger.warning("Failed to prepare canonical payload: %s", e)

            # Заполняем лист "Структура пр 2"
            struktura_sheet_names = [
                "Структура пр 2",
                "Структура пр 2 ",
                "Struktura pr2",
                "02_Структура",
            ]
            struktura_sheet = None
            # Сначала ищем точное совпадение (с учетом пробелов)
            for sheet_name in struktura_sheet_names:
                if sheet_name in workbook.sheetnames:
                    struktura_sheet = workbook[sheet_name]
                    break
            # Если не нашли, ищем с учетом пробелов в конце
            if not struktura_sheet:
                for ws_name in workbook.sheetnames:
                    ws_name_stripped = ws_name.strip()
                    for target_name in struktura_sheet_names:
                        if ws_name_stripped == target_name.strip():
                            struktura_sheet = workbook[ws_name]
                            break
                    if struktura_sheet:
                        break
            # Если не нашли точное совпадение, ищем по частичному
            if not struktura_sheet:
                for ws_name in workbook.sheetnames:
                    if "структура" in ws_name.lower() or "struktura" in ws_name.lower():
                        struktura_sheet = workbook[ws_name]
                        break

            # Создаем лист, если его нет
            if not struktura_sheet:
                logger.info("Лист 'Структура пр 2' не найден, создаем новый")
                struktura_sheet = workbook.create_sheet(title="02_Структура")

            if struktura_sheet:
                logger.info(
                    f"Заполнение листа '{struktura_sheet.title}' с данными: {len(available_resources)} ресурсов"
                )
                fill_struktura_pr2(
                    struktura_sheet,
                    resources_data,
                    loss_active_month=0.0,
                    loss_reactive_month=0.0,
                )
                logger.info(f"Лист '{struktura_sheet.title}' заполнен")

            # Заполняем лист "Узлы учета"
            nodes_json_path = AGGREGATED_DIR / f"{batch_id}_nodes.json"
            # Ищем nodes JSON во всех загрузках предприятия
            if not nodes_json_path.exists() and enterprise_id:
                uploads = database.list_uploads_for_enterprise(enterprise_id)
                for upload_item in uploads:
                    upload_batch_id = upload_item.get("batch_id")
                    if upload_batch_id:
                        candidate_path = (
                            AGGREGATED_DIR / f"{upload_batch_id}_nodes.json"
                        )
                        if candidate_path.exists():
                            nodes_json_path = candidate_path
                            logger.info(
                                f"Найден nodes JSON из другой загрузки: {nodes_json_path}"
                            )
                            break

            if canonical_payload and canonical_payload.get("nodes"):
                nodes_data = canonical_payload["nodes"]
                logger.info(
                    "Используются canonical узлы учета (%d шт.)", len(nodes_data)
                )
            else:
                if nodes_json_path.exists():
                    logger.info("Используется nodes JSON: %s", nodes_json_path)
                    nodes_data = load_nodes_from_json(nodes_json_path)
                else:
                    nodes_data = load_default_nodes()
                    logger.info("Используются узлы учета по умолчанию")

            nodes_sheet_names = [
                "01_Узлы учета",
                "Узел учета",
                "Узел учета ",
                "Узлы учета",
                "Nodes",
            ]
            nodes_sheet = None
            # Сначала ищем точное совпадение (с учетом пробелов)
            for sheet_name in nodes_sheet_names:
                if sheet_name in workbook.sheetnames:
                    nodes_sheet = workbook[sheet_name]
                    break
            # Если не нашли, ищем с учетом пробелов в конце
            if not nodes_sheet:
                for ws_name in workbook.sheetnames:
                    ws_name_stripped = ws_name.strip()
                    for target_name in nodes_sheet_names:
                        if ws_name_stripped == target_name.strip():
                            nodes_sheet = workbook[ws_name]
                            break
                    if nodes_sheet:
                        break
            # Если не нашли точное совпадение, ищем по частичному
            if not nodes_sheet:
                for ws_name in workbook.sheetnames:
                    if "узел" in ws_name.lower() or "nodes" in ws_name.lower():
                        nodes_sheet = workbook[ws_name]
                        break

            # Создаем лист, если его нет
            if not nodes_sheet:
                logger.info("Лист 'Узлы учета' не найден, создаем новый")
                nodes_sheet = workbook.create_sheet(title="01_Узлы учета")

            if nodes_sheet:
                logger.info(
                    f"Заполнение листа '{nodes_sheet.title}' с {len(nodes_data)} узлами"
                )
                fill_nodes_sheet(nodes_sheet, nodes_data)
                logger.info(f"Лист '{nodes_sheet.title}' заполнен")

            # Заполняем лист "Оборудование"
            equipment_json_path = AGGREGATED_DIR / f"{batch_id}_equipment.json"
            if not equipment_json_path.exists() and enterprise_id:
                uploads = database.list_uploads_for_enterprise(enterprise_id)
                for upload_item in uploads:
                    upload_batch_id = upload_item.get("batch_id")
                    if upload_batch_id:
                        candidate_path = (
                            AGGREGATED_DIR / f"{upload_batch_id}_equipment.json"
                        )
                        if candidate_path.exists():
                            equipment_json_path = candidate_path
                            break

            if canonical_payload and canonical_payload.get("equipment"):
                equipment_data = canonical_payload["equipment"]
                logger.info(
                    "Используются canonical данные оборудования (sections=%d)",
                    len(equipment_data.get("sheets", [])),
                )
                try:
                    # Ищем лист оборудования и заполняем
                    equipment_sheet_names = [
                        "Equipment",
                        "АНАЛИЗ ОБОРУДОВАНИЯ",
                        "Анализ оборудования",
                        "Оборудование",
                        "03_Оборудование",
                        "оборудование",
                    ]
                    equipment_sheet_found = False
                    for sheet_name in equipment_sheet_names:
                        if sheet_name in workbook.sheetnames:
                            fill_equipment_sheet(
                                workbook, equipment_data, sheet_name=sheet_name
                            )
                            equipment_sheet_found = True
                            break
                    if not equipment_sheet_found:
                        for sheet_name in workbook.sheetnames:
                            sheet_lower = sheet_name.lower()
                            if any(
                                keyword in sheet_lower
                                for keyword in ["equipment", "оборудование", "анализ"]
                            ):
                                fill_equipment_sheet(
                                    workbook, equipment_data, sheet_name=sheet_name
                                )
                                equipment_sheet_found = True
                                break
                    if not equipment_sheet_found:
                        logger.warning(
                            "Лист оборудования не найден для canonical payload"
                        )
                except Exception as equipment_exc:
                    logger.exception(
                        f"Ошибка при заполнении листа оборудования (canonical): {equipment_exc}"
                    )
            else:
                if equipment_json_path.exists():
                    logger.info("Используется equipment JSON: %s", equipment_json_path)
                    equipment_data = None
                    try:
                        import json as json_module

                        equipment_data = json_module.loads(
                            equipment_json_path.read_text(encoding="utf-8")
                        )
                        logger.info(
                            f"Загружены данные оборудования: sheets={len(equipment_data.get('sheets', []))}, summary={equipment_data.get('summary', {})}"
                        )
                    except Exception as equipment_exc:
                        logger.exception(
                            f"Ошибка при чтении JSON оборудования: {equipment_exc}"
                        )
                        equipment_data = None
                    if equipment_data:
                        # Ищем лист "Equipment" или другие варианты названий (сначала ищем точное совпадение)
                        equipment_sheet_names = [
                            "Equipment",
                            "АНАЛИЗ ОБОРУДОВАНИЯ",
                            "Анализ оборудования",
                            "Оборудование",
                            "03_Оборудование",
                            "оборудование",  # С маленькой буквы
                            "Sheet1",  # Возможное название в new_energy_passport
                        ]
                        equipment_sheet_found = False

                        # Сначала пытаемся найти точное совпадение
                        for sheet_name in equipment_sheet_names:
                            if sheet_name in workbook.sheetnames:
                                logger.info(f"Найден лист оборудования: '{sheet_name}'")
                                fill_equipment_sheet(
                                    workbook, equipment_data, sheet_name=sheet_name
                                )
                                logger.info(f"Лист '{sheet_name}' заполнен")
                                equipment_sheet_found = True
                                break

                        # Если точное совпадение не найдено, ищем частичное совпадение
                        if not equipment_sheet_found:
                            logger.info(
                                "Точное совпадение не найдено, ищем частичное совпадение..."
                            )
                            for sheet_name in workbook.sheetnames:
                                sheet_lower = sheet_name.lower()
                                if any(
                                    keyword in sheet_lower
                                    for keyword in [
                                        "equipment",
                                        "оборудование",
                                        "анализ",
                                    ]
                                ):
                                    logger.info(
                                        f"Найден лист по частичному совпадению: '{sheet_name}'"
                                    )
                                    fill_equipment_sheet(
                                        workbook, equipment_data, sheet_name=sheet_name
                                    )
                                    logger.info(f"Лист '{sheet_name}' заполнен")
                                    equipment_sheet_found = True
                                    break

                        if not equipment_sheet_found:
                            # Создаем лист Equipment, если его нет
                            logger.info("Лист оборудования не найден, создаем новый")
                            equipment_sheet = workbook.create_sheet(
                                title="03_Оборудование"
                            )
                            fill_equipment_sheet(
                                workbook, equipment_data, sheet_name="03_Оборудование"
                            )
                            equipment_sheet_found = True
                            logger.info("Лист '03_Оборудование' создан и заполнен")
                else:
                    logger.warning(f"Equipment JSON не найден: {equipment_json_path}")

            # Если есть годовые итоги из Canonical, добавляем строку 'ANNUAL' для ресурсов в баланс
            try:
                excel_ai_mode_runtime = get_excel_semantic_mode()
                if (
                    excel_ai_mode_runtime in ("assist", "strict")
                    and canonical_payload
                    and canonical_payload.get("balance", {}).get("annual_totals")
                ):
                    annual_totals = canonical_payload["balance"]["annual_totals"]
                    # Карта ключей для итогов по ресурсам (Balance используется для сравнения; кварталы не придумываем)
                    total_key_map = {
                        "electricity": "active_kwh",
                        "gas": "volume_m3",
                        "water": "volume_m3",
                        "heat": "energy_gcal",
                        "fuel": "volume_ton",
                        "coal": "volume_ton",
                    }
                    for resource_name, annual_value in annual_totals.items():
                        if annual_value is None:
                            continue
                        if resource_name not in resources_data:
                            # Пропускаем создание сложной структуры с нуля для отсутствующих ресурсов
                            continue
                        key = total_key_map.get(resource_name, "annual_total")
                        resources_data.setdefault(resource_name, {})
                        # Специальный ключ для годового итога — только в целях сравнения legacy vs canonical
                        resources_data[resource_name]["ANNUAL"] = {
                            "year": None,
                            "quarter": None,
                            "quarter_totals": {key: float(annual_value)},
                            "by_usage": {},
                        }
                        # Если есть canonical by_usage для ресурса — добавляем в ANNUAL
                        try:
                            byu = (
                                canonical_payload.get("balance", {}).get("by_usage", {})
                                or {}
                            ).get(resource_name)
                            if isinstance(byu, dict) and byu:
                                resources_data[resource_name]["ANNUAL"]["by_usage"] = (
                                    byu
                                )
                                logger.info(
                                    "Добавлен canonical %s.by_usage в ANNUAL: keys=%s",
                                    resource_name,
                                    list(byu.keys()),
                                )

                                # КРИТИЧНО: Распределяем canonical by_usage по кварталам пропорционально потреблению
                                if (
                                    resource_name == "electricity"
                                    and resources_data.get("electricity")
                                ):
                                    annual_by_usage = byu
                                    annual_total = float(annual_value)

                                    # Вычисляем общее потребление по всем кварталам
                                    total_quarterly_consumption = 0.0
                                    for quarter_key, quarter_data in resources_data[
                                        "electricity"
                                    ].items():
                                        if quarter_key == "ANNUAL":
                                            continue
                                        if isinstance(quarter_data, dict):
                                            quarter_total = quarter_data.get(
                                                "quarter_totals", {}
                                            ).get("active_kwh", 0)
                                            if quarter_total:
                                                total_quarterly_consumption += float(
                                                    quarter_total
                                                )

                                    # Распределяем by_usage по кварталам пропорционально потреблению
                                    if (
                                        total_quarterly_consumption > 0
                                        and annual_total > 0
                                    ):
                                        for quarter_key, quarter_data in resources_data[
                                            "electricity"
                                        ].items():
                                            if quarter_key == "ANNUAL":
                                                continue
                                            if isinstance(quarter_data, dict):
                                                quarter_total = quarter_data.get(
                                                    "quarter_totals", {}
                                                ).get("active_kwh", 0)
                                                if quarter_total and quarter_total > 0:
                                                    # Пропорциональное распределение
                                                    quarter_ratio = (
                                                        float(quarter_total)
                                                        / total_quarterly_consumption
                                                    )
                                                    quarter_by_usage = {
                                                        category: float(value)
                                                        * quarter_ratio
                                                        for category, value in annual_by_usage.items()
                                                    }
                                                    quarter_data["by_usage"] = (
                                                        quarter_by_usage
                                                    )
                                                    logger.debug(
                                                        "Распределен canonical by_usage для %s %s: %s (ratio=%.3f)",
                                                        resource_name,
                                                        quarter_key,
                                                        list(quarter_by_usage.keys()),
                                                        quarter_ratio,
                                                    )

                                        logger.info(
                                            "✅ Canonical by_usage распределен по %d кварталам electricity (annual_total=%.2f, quarterly_total=%.2f)",
                                            sum(
                                                1
                                                for k, v in resources_data[
                                                    "electricity"
                                                ].items()
                                                if k != "ANNUAL"
                                                and isinstance(v, dict)
                                                and v.get("by_usage")
                                            ),
                                            annual_total,
                                            total_quarterly_consumption,
                                        )
                                    else:
                                        logger.warning(
                                            "⚠️ Не удалось распределить canonical by_usage: annual_total=%.2f, quarterly_total=%.2f",
                                            annual_total,
                                            total_quarterly_consumption,
                                        )
                        except Exception as e:
                            logger.warning(
                                "Ошибка при добавлении canonical by_usage: %s", e
                            )
                        logger.info(
                            "Добавлен годовой итог Canonical для ресурса %s: %.2f (%s)",
                            resource_name,
                            float(annual_value),
                            key,
                        )
            except Exception as e:
                logger.warning(
                    "Не удалось применить годовые итоги Canonical для баланса: %s", e
                )

            # Заполняем лист "Расчет теплопотерь по зданиям"
            envelope_json_path = AGGREGATED_DIR / f"{batch_id}_envelope.json"
            if not envelope_json_path.exists() and enterprise_id:
                uploads = database.list_uploads_for_enterprise(enterprise_id)
                for upload_item in uploads:
                    upload_batch_id = upload_item.get("batch_id")
                    if upload_batch_id:
                        candidate_path = (
                            AGGREGATED_DIR / f"{upload_batch_id}_envelope.json"
                        )
                        if candidate_path.exists():
                            envelope_json_path = candidate_path
                            break

            if envelope_json_path.exists():
                logger.info(
                    "Используется JSON расчета теплопотерь по зданиям: %s",
                    envelope_json_path,
                )
                import json as json_module

                envelope_data = json_module.loads(
                    envelope_json_path.read_text(encoding="utf-8")
                )
                envelope_sheet_names = ["02_Исходные данные", "Ограждающие", "Envelope"]
                for sheet_name in envelope_sheet_names:
                    if sheet_name in workbook.sheetnames:
                        fill_building_envelope_sheet(
                            workbook, envelope_data, sheet_name=sheet_name
                        )
                        logger.info(
                            f"Лист '{sheet_name}' заполнен данными расчета теплопотерь по зданиям"
                        )
                        break

            # Заполняем лист "Баланс"
            balans_sheet_names = ["Баланс", "Баланс ", "04_Баланс", "Balance", "Balans"]
            balans_sheet = None
            for sheet_name in balans_sheet_names:
                if sheet_name in workbook.sheetnames:
                    balans_sheet = workbook[sheet_name]
                    break
            # Если не нашли точное совпадение, ищем по частичному
            if not balans_sheet:
                for ws_name in workbook.sheetnames:
                    if "баланс" in ws_name.lower() or "balance" in ws_name.lower():
                        balans_sheet = workbook[ws_name]
                        break

            # Создаем лист, если его нет
            if not balans_sheet:
                logger.info("Лист 'Баланс' не найден, создаем новый")
                balans_sheet = workbook.create_sheet(title="04_Баланс")
                # КРИТИЧЕСКОЕ ИСПРАВЛЕНИЕ: вызываем fill_balans_sheet для нового листа
                logger.info(f"Заполнение нового листа '{balans_sheet.title}'")
                fill_balans_sheet(balans_sheet, resources_data)
                logger.info(f"✅ Новый лист '{balans_sheet.title}' заполнен")

            if balans_sheet:
                # ========== ДЕТАЛЬНАЯ ДИАГНОСТИКА ПЕРЕД fill_balans_sheet ==========
                logger.info("=" * 80)
                logger.info(
                    "🔍 ДИАГНОСТИКА: Структура resources_data перед fill_balans_sheet"
                )
                logger.info("=" * 80)

                # 1. Проверяем тип и структуру resources_data
                logger.info(f"📋 Тип resources_data: {type(resources_data)}")
                logger.info(
                    f"📋 Ключи resources_data: {list(resources_data.keys())[:20]}"
                )

                # 2. Проверяем оба формата доступа к electricity
                electricity_direct = resources_data.get("electricity", {})
                electricity_via_resources = (
                    resources_data.get("resources", {}).get("electricity", {})
                    if isinstance(resources_data.get("resources"), dict)
                    else {}
                )

                logger.info(
                    f"📊 electricity (прямой доступ): type={type(electricity_direct)}, len={len(electricity_direct) if isinstance(electricity_direct, dict) else 'N/A'}"
                )
                logger.info(
                    f"📊 electricity (через resources): type={type(electricity_via_resources)}, len={len(electricity_via_resources) if isinstance(electricity_via_resources, dict) else 'N/A'}"
                )

                # Используем тот, который не пустой
                electricity = (
                    electricity_direct
                    if electricity_direct
                    else electricity_via_resources
                )

                if not electricity:
                    logger.error("❌ КРИТИЧНО: electricity пуст в обоих форматах!")
                    logger.error(
                        f"   resources_data.get('electricity'): {electricity_direct}"
                    )
                    logger.error(
                        f"   resources_data.get('resources', {{}}).get('electricity'): {electricity_via_resources}"
                    )
                    logger.error(
                        "   Полная структура resources_data (первые 3 уровня):"
                    )
                    import json

                    try:
                        # Безопасный вывод структуры (без рекурсии)
                        struct_preview = {}
                        for key, value in list(resources_data.items())[:10]:
                            if isinstance(value, dict):
                                struct_preview[key] = {
                                    "type": "dict",
                                    "keys": list(value.keys())[:10],
                                    "sample": {
                                        k: type(v).__name__
                                        for k, v in list(value.items())[:5]
                                    },
                                }
                            else:
                                struct_preview[key] = type(value).__name__
                        logger.error(
                            f"   {json.dumps(struct_preview, indent=2, ensure_ascii=False)}"
                        )
                    except Exception as e:
                        logger.error(f"   Ошибка при выводе структуры: {e}")
                else:
                    logger.info(f"✅ electricity найден: {len(electricity)} элементов")
                    logger.info(
                        f"   Ключи electricity: {list(electricity.keys())[:10]}"
                    )

                # 3. Детальный анализ кварталов
                quarters_with_by_usage = []
                quarters_without_by_usage = []
                annual_has_by_usage = False

                if electricity:
                    for quarter, quarter_data in electricity.items():
                        if quarter == "ANNUAL":
                            # ANNUAL обрабатываем отдельно
                            continue
                        if not isinstance(quarter_data, dict):
                            logger.warning(
                                f"   ⚠️ Квартал {quarter}: не является dict, type={type(quarter_data)}"
                            )
                            continue
                        by_usage = quarter_data.get("by_usage")
                        quarter_totals = quarter_data.get("quarter_totals", {})
                        active_kwh = quarter_totals.get("active_kwh", 0)

                        if (
                            by_usage
                            and isinstance(by_usage, dict)
                            and len(by_usage) > 0
                        ):
                            quarters_with_by_usage.append(quarter)
                            logger.debug(
                                f"   ✅ Квартал {quarter}: by_usage найден - {list(by_usage.keys())}, active_kwh={active_kwh}"
                            )
                        else:
                            quarters_without_by_usage.append(quarter)
                            logger.debug(
                                f"   ⚠️ Квартал {quarter}: by_usage отсутствует, active_kwh={active_kwh}, quarter_totals keys={list(quarter_totals.keys())}"
                            )

                    # Проверяем ANNUAL (может содержать by_usage из canonical)
                    annual_data = electricity.get("ANNUAL")
                    if annual_data and isinstance(annual_data, dict):
                        annual_by_usage = annual_data.get("by_usage")
                        annual_totals = annual_data.get("quarter_totals", {})
                        if (
                            annual_by_usage
                            and isinstance(annual_by_usage, dict)
                            and len(annual_by_usage) > 0
                        ):
                            annual_has_by_usage = True
                            logger.info(
                                f"   ✅ ANNUAL: by_usage найден - {list(annual_by_usage.keys())}, totals={list(annual_totals.keys())}"
                            )
                        else:
                            logger.info(
                                f"   ⚠️ ANNUAL: by_usage отсутствует, totals={list(annual_totals.keys())}"
                            )

                # 4. Итоговая статистика
                logger.info("=" * 80)
                logger.info(
                    f"📊 СТАТИСТИКА: кварталов с by_usage: {len(quarters_with_by_usage)}, без: {len(quarters_without_by_usage)}"
                )
                if annual_has_by_usage:
                    logger.info(
                        "   ✅ ANNUAL содержит by_usage из canonical (для годового итога)"
                    )

                if quarters_without_by_usage:
                    logger.warning(
                        f"   ⚠️ Кварталы без by_usage (будут пропущены в листе баланса): {quarters_without_by_usage}"
                    )
                    logger.warning(
                        "   Это может привести к ошибке валидации. Проверьте, что:"
                    )
                    logger.warning("     1. Файл pererashod.xlsx загружен и обработан")
                    logger.warning(
                        "     2. Функция aggregate_usage_categories() вернула данные"
                    )
                    logger.warning(
                        "     3. Функция distribute_categories_by_quarter() успешно выполнилась"
                    )
                    logger.warning(
                        "     4. Если используется canonical mode, проверьте наличие оборудования в CanonicalSourceData"
                    )

                if quarters_with_by_usage:
                    logger.info(
                        f"   ✅ Кварталы с by_usage (будут заполнены): {quarters_with_by_usage}"
                    )

                # 5. Дополнительная диагностика: проверяем источник данных
                if canonical_payload and canonical_payload.get("balance", {}).get(
                    "by_usage", {}
                ).get("electricity"):
                    canonical_by_usage = canonical_payload["balance"]["by_usage"][
                        "electricity"
                    ]
                    logger.info(
                        f"   📊 Canonical by_usage для electricity найден: {list(canonical_by_usage.keys())}"
                    )
                    if not annual_has_by_usage:
                        logger.warning(
                            "   ⚠️ Canonical by_usage существует, но не попал в ANNUAL. Проверьте интеграцию."
                        )
                elif excel_ai_mode_runtime in ("assist", "strict"):
                    logger.warning(
                        "   ⚠️ Canonical mode активен, но by_usage для electricity не найден в canonical_payload"
                    )
                    logger.warning("      Возможные причины:")
                    logger.warning(
                        "      - Оборудование отсутствует в CanonicalSourceData"
                    )
                    logger.warning("      - annual_totals.electricity отсутствует")
                    logger.warning(
                        "      - compute_electricity_by_usage вернул пустой словарь"
                    )

                logger.info("=" * 80)

                # 6. Нормализуем структуру resources_data для fill_balans_sheet
                # Убеждаемся, что electricity доступен в ожидаемом формате
                if not electricity_direct and electricity_via_resources:
                    # Если данные есть только через resources, нормализуем структуру
                    logger.info(
                        "🔄 Нормализация: переносим electricity из resources в корень resources_data"
                    )
                    resources_data["electricity"] = electricity_via_resources
                    electricity = electricity_via_resources

                # Также убеждаемся, что есть структура resources
                if "resources" not in resources_data:
                    resources_data["resources"] = {}
                if "electricity" not in resources_data["resources"]:
                    resources_data["resources"]["electricity"] = electricity

                logger.info("✅ Вызов fill_balans_sheet с нормализованными данными")
                logger.info(
                    f"   electricity доступен: {len(electricity) if electricity else 0} элементов"
                )
                logger.info("=" * 80)

                fill_balans_sheet(balans_sheet, resources_data)
                logger.info(f"✅ Лист '{balans_sheet.title}' заполнен")

            # Заполняем лист "Динамика ср"
            dinamika_sheet_names = [
                "Динамика ср",
                "Динамика ср ",
                "Динамика",
                "05_Динамика",
                "Dynamics",
            ]
            dinamika_sheet = None
            for sheet_name in dinamika_sheet_names:
                if sheet_name in workbook.sheetnames:
                    dinamika_sheet = workbook[sheet_name]
                    break
            # Если не нашли точное совпадение, ищем по частичному
            if not dinamika_sheet:
                for ws_name in workbook.sheetnames:
                    if "динамика" in ws_name.lower() or "dynamics" in ws_name.lower():
                        dinamika_sheet = workbook[ws_name]
                        break

            if dinamika_sheet:
                logger.info(f"Заполнение листа '{dinamika_sheet.title}'")
                fill_dinamika_sheet(dinamika_sheet, resources_data)
                logger.info(f"Лист '{dinamika_sheet.title}' заполнен")

            # Заполняем лист "мазут,уголь 5"
            fuel_dynamics_sheet_names = [
                "мазут,уголь 5",
                "мазут,уголь 5 ",
                "мазут,уголь",
                "06_Мазут_Уголь",
                "Fuel Dynamics",
            ]
            fuel_dynamics_sheet = None
            for sheet_name in fuel_dynamics_sheet_names:
                if sheet_name in workbook.sheetnames:
                    fuel_dynamics_sheet = workbook[sheet_name]
                    break
            # Если не нашли точное совпадение, ищем по частичному
            if not fuel_dynamics_sheet:
                for ws_name in workbook.sheetnames:
                    if (
                        "мазут" in ws_name.lower()
                        or "уголь" in ws_name.lower()
                        or "fuel" in ws_name.lower()
                    ):
                        fuel_dynamics_sheet = workbook[ws_name]
                        break

            if fuel_dynamics_sheet:
                logger.info(f"Заполнение листа '{fuel_dynamics_sheet.title}'")
                fill_fuel_dynamics_sheet(fuel_dynamics_sheet, resources_data)
                logger.info(f"Лист '{fuel_dynamics_sheet.title}' заполнен")

            # Заполняем лист "Расход на ед.п"
            specific_consumption_sheet_names = [
                "Расход на ед.п",
                "Расход  на ед.п",
                "Расход  на ед.п ",
                "07_Расход_на_ед",
                "Specific Consumption",
            ]
            specific_consumption_sheet = None
            for sheet_name in specific_consumption_sheet_names:
                if sheet_name in workbook.sheetnames:
                    specific_consumption_sheet = workbook[sheet_name]
                    break
            # Если не нашли точное совпадение, ищем по частичному
            if not specific_consumption_sheet:
                for ws_name in workbook.sheetnames:
                    if "расход" in ws_name.lower() and "ед" in ws_name.lower():
                        specific_consumption_sheet = workbook[ws_name]
                        break

            if specific_consumption_sheet:
                logger.info(f"Заполнение листа '{specific_consumption_sheet.title}'")
                fill_specific_consumption_sheet(
                    specific_consumption_sheet, resources_data
                )
                logger.info(f"Лист '{specific_consumption_sheet.title}' заполнен")

            # Заполняем лист "Мериаприятия 1"
            meropriyatiya_sheet_names = [
                "Мериаприятия 1",
                "Мериаприятия 1 ",
                "Мероприятия",
                "08_Мероприятия",
                "Measures",
            ]
            meropriyatiya_sheet = None
            # Сначала ищем точное совпадение (с учетом пробелов)
            for sheet_name in meropriyatiya_sheet_names:
                if sheet_name in workbook.sheetnames:
                    meropriyatiya_sheet = workbook[sheet_name]
                    break
            # Если не нашли, ищем с учетом пробелов в конце
            if not meropriyatiya_sheet:
                for ws_name in workbook.sheetnames:
                    ws_name_stripped = ws_name.strip()
                    for target_name in meropriyatiya_sheet_names:
                        if ws_name_stripped == target_name.strip():
                            meropriyatiya_sheet = workbook[ws_name]
                            break
                    if meropriyatiya_sheet:
                        break
            # Если не нашли точное совпадение, ищем по частичному
            if not meropriyatiya_sheet:
                for ws_name in workbook.sheetnames:
                    if (
                        "мероприятия" in ws_name.lower()
                        or "мериаприятия" in ws_name.lower()
                        or "measures" in ws_name.lower()
                    ):
                        meropriyatiya_sheet = workbook[ws_name]
                        break

            if meropriyatiya_sheet:
                logger.info(f"Заполнение листа '{meropriyatiya_sheet.title}'")
                fill_meropriyatiya_sheet(meropriyatiya_sheet)
                logger.info(f"Лист '{meropriyatiya_sheet.title}' заполнен")

            # Заполняем лист "Monthly" (месячные данные)
            monthly_sheet_names = [
                "Monthly",
                "MONTHLY",
                "Месячные данные",
                "Месячный",
                "месячные",
            ]
            monthly_sheet = None
            for sheet_name in monthly_sheet_names:
                if sheet_name in workbook.sheetnames:
                    monthly_sheet = workbook[sheet_name]
                    break

            # Если не нашли точное совпадение, ищем частичное
            if not monthly_sheet:
                for sheet_name in workbook.sheetnames:
                    sheet_lower = sheet_name.lower()
                    if (
                        "monthly" in sheet_lower
                        or "месяц" in sheet_lower
                        or "месячн" in sheet_lower
                    ):
                        monthly_sheet = workbook[sheet_name]
                        logger.info(
                            f"Найден лист месячных данных по частичному совпадению: '{sheet_name}'"
                        )
                        break

            if monthly_sheet:
                logger.info(f"Заполнение листа '{monthly_sheet.title}'")
                try:
                    fill_monthly_sheet(monthly_sheet, resources_data)
                    logger.info(f"Лист '{monthly_sheet.title}' заполнен успешно")
                except Exception as monthly_exc:
                    logger.error(
                        f"Ошибка при заполнении листа '{monthly_sheet.title}': {monthly_exc}",
                        exc_info=True,
                    )
            else:
                logger.warning(
                    "Лист 'Monthly' не найден. Доступные листы: "
                    + ", ".join(workbook.sheetnames[:10])
                )

            # Проверяем, все ли листы заполнены
            filled_sheets = set()
            equipment_sheet_name = (
                None  # Инициализируем переменную для имени листа оборудования
            )

            # Собираем список заполненных листов
            if struktura_sheet:
                filled_sheets.add(struktura_sheet.title)
            if nodes_sheet:
                filled_sheets.add(nodes_sheet.title)

            # Проверяем лист оборудования (перепроверяем, был ли он заполнен)
            if equipment_json_path.exists():
                for sheet_name in workbook.sheetnames:
                    sheet_lower = sheet_name.lower()
                    if any(
                        keyword in sheet_lower
                        for keyword in ["equipment", "оборудование", "анализ"]
                    ):
                        filled_sheets.add(sheet_name)
                        equipment_sheet_name = sheet_name
                        break

            # Проверяем лист расчета теплопотерь по зданиям
            envelope_sheet_name = None
            if envelope_json_path.exists():
                for sheet_name in workbook.sheetnames:
                    if any(
                        name in sheet_name
                        for name in ["02_Исходные данные", "Ограждающие", "Envelope"]
                    ):
                        filled_sheets.add(sheet_name)
                        envelope_sheet_name = sheet_name
                        break

            if balans_sheet:
                filled_sheets.add(balans_sheet.title)
            if dinamika_sheet:
                filled_sheets.add(dinamika_sheet.title)
            if fuel_dynamics_sheet:
                filled_sheets.add(fuel_dynamics_sheet.title)
            if specific_consumption_sheet:
                filled_sheets.add(specific_consumption_sheet.title)
            if meropriyatiya_sheet:
                filled_sheets.add(meropriyatiya_sheet.title)
            if monthly_sheet:
                filled_sheets.add(monthly_sheet.title)

            # Находим незаполненные листы
            all_sheets = set(workbook.sheetnames)
            unfilled_sheets = all_sheets - filled_sheets

            if unfilled_sheets:
                logger.warning(
                    f"⚠️  Найдены незаполненные листы ({len(unfilled_sheets)}): "
                    f"{', '.join(sorted(unfilled_sheets))}"
                )
                logger.info(
                    f"📊 Статистика заполнения: "
                    f"заполнено {len(filled_sheets)}/{len(all_sheets)} листов "
                    f"({round(len(filled_sheets) / len(all_sheets) * 100, 1)}%)"
                )
            else:
                logger.info(
                    f"✅ Все листы заполнены: {len(filled_sheets)}/{len(all_sheets)} листов "
                    f"({round(len(filled_sheets) / len(all_sheets) * 100, 1)}%)"
                )

            # Восстановление формул с помощью AI (если доступно)
            if HAS_FORMULA_RESTORER:
                try:
                    logger.info("Начинаю восстановление формул в паспорте...")
                    restorer = AIFormulaRestorer()

                    # Ищем все ячейки с ошибками #REF! и восстанавливаем их
                    restored_count = 0
                    total_ref_errors = 0

                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        for row in sheet.iter_rows():
                            for cell in row:
                                if cell.data_type == "f" and cell.value:
                                    formula_str = str(cell.value)
                                    if "#REF!" in formula_str:
                                        total_ref_errors += 1
                                        cell_coord = cell.coordinate

                                        restored_formula = restorer.restore_ref_error(
                                            workbook,
                                            sheet_name,
                                            cell_coord,
                                            formula_str,
                                        )

                                        if restored_formula:
                                            cell.value = restored_formula
                                            restored_count += 1
                                            logger.debug(
                                                f"Восстановлена формула в {sheet_name}!{cell_coord}: "
                                                f"{formula_str} -> {restored_formula}"
                                            )

                    if total_ref_errors > 0:
                        logger.info(
                            f"Восстановление формул завершено: "
                            f"восстановлено {restored_count}/{total_ref_errors} формул"
                        )
                    else:
                        logger.info(
                            "Ошибок #REF! не найдено, восстановление не требуется"
                        )
                except Exception as restore_exc:
                    logger.warning(
                        f"Ошибка при восстановлении формул: {restore_exc}. Продолжаю без восстановления."
                    )

            debug_output = AGGREGATED_DIR / f"{batch_id}_filler_passport.xlsx"
            workbook.save(debug_output)
            logger.info("Паспорт сохранён для отладки: %s", debug_output)
            workbook.save(output_file)
            logger.info(f"Паспорт сохранён: {output_file}")

            # Проверяем, что файл существует и не пустой
            if not output_file.exists():
                raise HTTPException(
                    status_code=500,
                    detail=f"Файл паспорта не был создан: {output_file}",
                )
            file_size = output_file.stat().st_size
            if file_size == 0:
                raise HTTPException(
                    status_code=500, detail=f"Файл паспорта пустой: {output_file}"
                )
            logger.info(
                f"✅ Паспорт сгенерирован (filler): {output_file} (размер: {file_size} байт)"
            )

            return FileResponse(
                path=str(output_file),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                filename=f"Энергопаспорт_{enterprise_data['name']}_{batch_id[:8]}.xlsx",
            )
        except Exception as filler_exc:
            logger.exception(
                "Ошибка генерации через fill_energy_passport: %s", filler_exc
            )
            # Если был указан template_name, не переходим к PKM690ExcelGenerator
            # (он не использует шаблоны)
            if template_name_final:
                raise HTTPException(
                    status_code=500,
                    detail=(
                        f"Ошибка генерации паспорта с шаблоном '{template_name_final}': {filler_exc}"
                    ),
                ) from filler_exc
            if not HAS_GENERATOR:
                raise HTTPException(
                    status_code=500, detail=f"Ошибка генерации: {filler_exc}"
                ) from filler_exc

    # Если был указан template_name, не используем PKM690ExcelGenerator
    # (он создает файл с нуля и не использует шаблоны)
    if template_name_final:
        raise HTTPException(
            status_code=503,
            detail=(
                f"Для использования шаблона '{template_name_final}' требуется fill_energy_passport. "
                "Модуль недоступен или произошла ошибка. Проверьте установку зависимостей и логи."
            ),
        )

    if not HAS_GENERATOR:
        raise HTTPException(status_code=503, detail="Генератор паспортов недоступен")

    try:
        logger.info(
            f"Генерация паспорта через PKM690ExcelGenerator для {enterprise_data['name']}"
        )
        generator = PKM690ExcelGenerator(
            enterprise_data=enterprise_data, energy_data=aggregated
        )

        success = generator.create_energy_passport(str(output_file))

        if not success or not output_file.exists():
            raise HTTPException(status_code=500, detail="Ошибка генерации паспорта")

        logger.info(f"✅ Паспорт сгенерирован: {output_file}")

        return FileResponse(
            path=str(output_file),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=f"Энергопаспорт_{enterprise_data['name']}_{batch_id[:8]}.xlsx",
        )

    except Exception as e:
        logger.exception(f"Ошибка генерации паспорта: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка генерации: {str(e)}")


@app.post("/api/generate-word-report/{batch_id}")
async def generate_word_report(
    batch_id: str,
    skip_readiness_check: bool = Query(
        False, description="Пропустить проверку готовности данных"
    ),
):
    """
    Генерация Word отчета энергоаудита для batch_id
    """
    if not HAS_WORD_GENERATOR:
        raise HTTPException(
            status_code=503,
            detail="Генератор Word отчетов недоступен (python-docx не установлен)",
        )

    # 1. Получаем загрузку
    upload = database.get_upload_by_batch(batch_id)
    if not upload:
        raise HTTPException(status_code=404, detail=f"Загрузка {batch_id} не найдена")

    # 2. Проверка готовности данных
    if not skip_readiness_check:
        enterprise_id = upload.get("enterprise_id")
        if enterprise_id:
            readiness = validate_generation_readiness(enterprise_id)
            if not readiness["ready"]:
                logger.warning(
                    f"Попытка генерации Word отчета для предприятия {enterprise_id} при неготовности данных"
                )
                detail = {
                    "message": "Данные не готовы для генерации Word отчета",
                    "missing_resources": readiness.get("missing_resources", []),
                    "warnings": readiness.get("warnings", []),
                }
                raise HTTPException(status_code=400, detail=detail)

    # 3. Агрегируем данные из исходных файлов
    enterprise_id = upload.get("enterprise_id")
    logger.info(
        f"Генерация Word отчета для предприятия {enterprise_id} (batch_id: {batch_id})"
    )
    logger.info("📊 Генерация Word отчета из исходных данных с применением AI-анализа")

    from utils.readiness_validator import _get_aggregated_data_for_enterprise

    aggregated = _get_aggregated_data_for_enterprise(enterprise_id)

    if not aggregated:
        raw_json = upload.get("raw_json")
        if not raw_json:
            raise HTTPException(status_code=400, detail="Данные не распарсены")
        from utils.energy_aggregator import aggregate_from_db_json

        aggregated = aggregate_from_db_json(raw_json)

    if not aggregated:
        raise HTTPException(status_code=400, detail="Не удалось агрегировать данные")

    # 3.1. Применяем AI-анализ к агрегированным данным (если доступен)
    try:
        from ai_energy_analysis import enhanced_energy_analysis

        HAS_AI_ENERGY_ANALYSIS = True
    except ImportError:
        HAS_AI_ENERGY_ANALYSIS = False
        logger.debug(
            "ai_energy_analysis модуль не найден. AI-анализ данных недоступен."
        )

    ai_processed_data = None
    if HAS_AI_ENERGY_ANALYSIS:
        try:
            logger.info(
                "🤖 Применение AI-анализа к агрегированным данным для улучшения качества отчета..."
            )
            # Создаем структуру данных для AI-анализа
            analysis_input = {
                "data": aggregated,
                "metadata": {
                    "enterprise_id": enterprise_id,
                    "batch_id": batch_id,
                    "source": "aggregated_from_files",
                },
            }
            ai_processed_data = enhanced_energy_analysis(analysis_input)
            if ai_processed_data and not ai_processed_data.get("error"):
                logger.info("✅ AI-анализ успешно обработал данные")
                # Обогащаем агрегированные данные результатами AI-анализа
                # Извлекаем инсайты из результатов AI-анализа
                ai_insights = {}
                ai_recommendations = []

                # Формируем инсайты из результатов верификации, аномалий, эффективности и соответствия
                summary = ai_processed_data.get("summary", {})
                if summary:
                    ai_insights["Статус верификации"] = (
                        "Данные валидны"
                        if summary.get("is_valid")
                        else "Обнаружены проблемы с данными"
                    )
                    ai_insights["Наличие аномалий"] = (
                        f"Обнаружено аномалий: {ai_processed_data.get('anomalies', {}).get('anomaly_count', 0)}"
                        if summary.get("has_anomalies")
                        else "Аномалий не обнаружено"
                    )
                    ai_insights["Соответствие нормам"] = (
                        "Соответствует"
                        if summary.get("is_compliant")
                        else "Требуется улучшение"
                    )
                    efficiency_class = summary.get("efficiency_class", "N/A")
                    if efficiency_class != "N/A":
                        ai_insights["Класс энергоэффективности"] = efficiency_class

                # Добавляем рекомендации из анализа эффективности
                efficiency = ai_processed_data.get("efficiency", {})
                if efficiency and "recommendations" in efficiency:
                    ai_recommendations.extend(efficiency.get("recommendations", []))

                # Добавляем рекомендации из проверки соответствия
                compliance = ai_processed_data.get("compliance", {})
                if compliance and "recommendations" in compliance:
                    ai_recommendations.extend(compliance.get("recommendations", []))

                # Сохраняем AI-инсайты и рекомендации в агрегированных данных
                if ai_insights:
                    aggregated["ai_insights"] = ai_insights
                if ai_recommendations:
                    aggregated["ai_recommendations"] = ai_recommendations[
                        :20
                    ]  # Первые 20 рекомендаций

                logger.info(
                    f"AI-анализ добавил {len(ai_insights)} инсайтов и {len(ai_recommendations)} рекомендаций"
                )
            else:
                error_msg = (
                    ai_processed_data.get("error", "Неизвестная ошибка")
                    if ai_processed_data
                    else "Нет данных"
                )
                logger.info(
                    f"AI-анализ не вернул обработанные данные ({error_msg}), используем исходные агрегированные данные"
                )
        except Exception as ai_exc:
            logger.warning(
                f"Ошибка при применении AI-анализа (продолжаем без AI): {ai_exc}"
            )
            # Продолжаем генерацию без AI-обогащения

    # 4. Подготовка данных предприятия
    enterprise = database.get_enterprise_by_id(enterprise_id) if enterprise_id else None
    enterprise_data = {
        "id": enterprise_id,
        "name": enterprise.get("name", "Неизвестное предприятие")
        if enterprise
        else upload.get("enterprise_name", "Неизвестное предприятие"),
        "address": enterprise.get("address", "не указан")
        if enterprise
        else "не указан",
    }

    # 5. Загружаем дополнительные данные (если есть)
    equipment_data = None
    nodes_data = None
    envelope_data = None

    # Оборудование
    equipment_json_path = AGGREGATED_DIR / f"{batch_id}_equipment.json"
    if not equipment_json_path.exists() and enterprise_id:
        uploads = database.list_uploads_for_enterprise(enterprise_id)
        for upload_item in uploads:
            upload_batch_id = upload_item.get("batch_id")
            if upload_batch_id:
                candidate_path = AGGREGATED_DIR / f"{upload_batch_id}_equipment.json"
                if candidate_path.exists():
                    equipment_json_path = candidate_path
                    break

    if equipment_json_path.exists():
        try:
            equipment_data = json.loads(equipment_json_path.read_text(encoding="utf-8"))
        except Exception as e:
            logger.warning(f"Не удалось загрузить данные оборудования: {e}")

    # Узлы учета
    nodes_json_path = AGGREGATED_DIR / f"{batch_id}_nodes.json"
    if not nodes_json_path.exists() and enterprise_id:
        uploads = database.list_uploads_for_enterprise(enterprise_id)
        for upload_item in uploads:
            upload_batch_id = upload_item.get("batch_id")
            if upload_batch_id:
                candidate_path = AGGREGATED_DIR / f"{upload_batch_id}_nodes.json"
                if candidate_path.exists():
                    nodes_json_path = candidate_path
                    break

    if nodes_json_path.exists():
        try:
            from tools.fill_energy_passport import load_nodes_from_json

            nodes_data = load_nodes_from_json(nodes_json_path)
        except Exception as e:
            logger.warning(f"Не удалось загрузить данные узлов учета: {e}")

    # Расчет теплопотерь по зданиям
    envelope_json_path = AGGREGATED_DIR / f"{batch_id}_envelope.json"
    if not envelope_json_path.exists() and enterprise_id:
        uploads = database.list_uploads_for_enterprise(enterprise_id)
        for upload_item in uploads:
            upload_batch_id = upload_item.get("batch_id")
            if upload_batch_id:
                candidate_path = AGGREGATED_DIR / f"{upload_batch_id}_envelope.json"
                if candidate_path.exists():
                    envelope_json_path = candidate_path
                    break

    if envelope_json_path.exists():
        try:
            envelope_data = json.loads(envelope_json_path.read_text(encoding="utf-8"))
        except Exception as e:
            logger.warning(
                f"Не удалось загрузить данные расчета теплопотерь по зданиям: {e}"
            )

    # 6. Генерируем Word отчет из исходных данных с применением AI-анализа
    try:
        # Генерация Word отчета выполняется на основе:
        # - Агрегированных данных из исходных файлов (Excel, PDF и т.д.)
        # - Результатов AI-анализа данных (если доступен)
        # - Дополнительных данных (оборудование, узлы учета, ограждающие конструкции)
        logger.info("📄 Генерация Word отчета из исходных данных с AI-обогащением...")

        generator = WordReportGenerator()

        # Создаем временный файл
        output_dir = AGGREGATED_DIR
        output_file = output_dir / f"{batch_id}_report.docx"

        # Генерируем документ из исходных агрегированных данных
        # (AI-обработанные данные уже обогатили aggregated, если AI был доступен)
        generator.generate_report(
            enterprise_data=enterprise_data,
            aggregated_data=aggregated,  # Включает AI-инсайты, если AI применялся
            equipment_data=equipment_data,
            nodes_data=nodes_data,
            envelope_data=envelope_data,
            output_path=output_file,
        )

        logger.info(f"✅ Word отчет сгенерирован из исходных данных: {output_file}")
        if ai_processed_data:
            logger.info("📊 Отчет обогащен результатами AI-анализа")

        return FileResponse(
            path=str(output_file),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"Отчет_энергоаудита_{enterprise_data['name']}_{batch_id[:8]}.docx",
        )
    except Exception as exc:
        logger.exception(f"Ошибка генерации Word отчета: {exc}")
        raise HTTPException(
            status_code=500, detail=f"Ошибка генерации Word отчета: {exc}"
        ) from exc


@app.post("/ingest/files")
async def ingest_files(file: UploadFile = File(...)):
    """API endpoint для загрузки файлов с валидацией согласно ТЗ"""
    is_valid, error_msg = validate_file(file)
    if not is_valid:
        raise HTTPException(status_code=400, detail=error_msg)

    size_valid, size_error = await validate_file_size(file)
    if not size_valid:
        raise HTTPException(status_code=400, detail=size_error)

    batch_id = str(uuid4())
    safe_filename = os.path.basename(file.filename)
    save_path = f"/tmp/ingest_{batch_id}_{safe_filename}"

    try:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Empty file not allowed")
        with open(save_path, "wb") as output_file:
            output_file.write(content)
    except HTTPException:
        raise
    except PermissionError as exc:
        raise HTTPException(
            status_code=500, detail=f"Permission denied when saving file: {exc}"
        )
    except OSError as exc:  # pragma: no cover - IO errors
        raise HTTPException(status_code=500, detail=f"Failed to save file: {exc}")
    except Exception as exc:  # pragma: no cover - unexpected errors
        raise HTTPException(
            status_code=500, detail=f"Unexpected error saving file: {exc}"
        )

    validate_resp: Dict[str, Any] = {}
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.post(
                "http://validate:8002/validate/run", json={"batchId": batch_id}
            )
            resp.raise_for_status()
            validate_resp = resp.json()
    except httpx.TimeoutException:
        validate_resp = {"error": "validate service timeout", "batchId": batch_id}
    except httpx.HTTPStatusError as exc:
        validate_resp = {
            "error": f"validate service returned {exc.response.status_code}",
            "batchId": batch_id,
        }
    except httpx.RequestError as exc:
        validate_resp = {
            "error": f"validate service connection failed: {exc}",
            "batchId": batch_id,
        }
    except Exception as exc:  # pragma: no cover - unexpected
        validate_resp = {"error": f"validate call failed: {exc}", "batchId": batch_id}

    parsing_result = None
    try:
        logger.info("Начинаю парсинг файла через API: %s", save_path)
        parsing_result = parse_file(save_path)
        status_value = "success" if parsing_result.get("parsed") else "partial"
        parsing_results_cache[batch_id] = {
            "batch_id": batch_id,
            "filename": safe_filename,
            "file_path": save_path,
            "parsing": parsing_result,
            "status": status_value,
        }
    except Exception as exc:  # pragma: no cover - parse failure path
        logger.exception("Ошибка при парсинге файла %s", save_path)
        parsing_results_cache[batch_id] = {
            "batch_id": batch_id,
            "filename": safe_filename,
            "file_path": save_path,
            "parsing": None,
            "error": str(exc),
            "status": "error",
        }

    parsing_status = (
        "success"
        if parsing_result and parsing_result.get("parsed")
        else ("error" if parsing_result is None else "pending")
    )

    return {
        "batchId": batch_id,
        "filename": safe_filename,
        "validate": validate_resp,
        "parsing_status": parsing_status,
    }


@app.post("/api/normative/upload")
async def upload_normative_document(
    file: UploadFile = File(...),
    title: Optional[str] = Form(None),
    document_type: Optional[str] = Form(None),
):
    """
    Загрузка нормативного документа с AI-анализом и сохранением в БД.

    Поддерживает форматы: PDF, Word, Excel.
    Документ парсится, анализируется AI для извлечения формул и нормативов,
    результат сохраняется в БД для использования в расчетах энергопаспорта.

    Args:
        file: Файл нормативного документа
        title: Название документа (опционально, если не указано - берется из имени файла)
        document_type: Тип документа (PKM690, GOST, SNiP и т.д., опционально)

    Returns:
        Результат импорта с количеством извлеченных правил
    """
    try:
        logger.info(
            f"Получен запрос на загрузку нормативного документа: {file.filename}"
        )

        # Валидация файла
        try:
            is_valid, error_msg = validate_file(file)
            if not is_valid:
                logger.warning(f"Валидация файла не пройдена: {error_msg}")
                raise HTTPException(status_code=400, detail=error_msg)
        except Exception as e:
            logger.error(f"Ошибка валидации файла: {e}")
            raise HTTPException(status_code=400, detail=f"Ошибка валидации файла: {e}")

        # Сохраняем размер файла для проверки дубликата
        file_size = None
        try:
            size_valid, size_error = await validate_file_size(file)
            if not size_valid:
                logger.warning(f"Проверка размера файла не пройдена: {size_error}")
                raise HTTPException(status_code=400, detail=size_error)
            # После проверки размера получаем размер файла
            await file.seek(0)
            content_for_size = await file.read()
            file_size = len(content_for_size)
            await file.seek(0)  # Сбрасываем позицию для дальнейшего чтения
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Ошибка проверки размера файла: {e}")
            raise HTTPException(
                status_code=400, detail=f"Ошибка проверки размера файла: {e}"
            )

        # Импорт модуля
        try:
            from domain.normative_importer import get_normative_importer

            importer = get_normative_importer()
            if not importer:
                logger.warning("Импортер нормативных документов недоступен")
                raise HTTPException(
                    status_code=503,
                    detail="Импортер нормативных документов недоступен. Проверьте настройки AI.",
                )
        except ImportError as e:
            logger.exception("Не удалось импортировать normative_importer")
            raise HTTPException(
                status_code=500,
                detail=f"Модуль импорта нормативных документов недоступен: {e}",
            )
        except HTTPException:
            raise
        except Exception as e:
            logger.exception(f"Неожиданная ошибка при импорте модуля: {e}")
            raise HTTPException(
                status_code=500, detail=f"Ошибка инициализации импортера: {e}"
            )

        # Сохраняем файл
        safe_filename = os.path.basename(file.filename)
        normative_dir = Path(INBOX_DIR) / "normative"
        normative_dir.mkdir(parents=True, exist_ok=True)

        file_path = normative_dir / safe_filename

        # Проверка на дубликат (по имени и размеру)
        if file_path.exists() and file_size is not None:
            existing_size = file_path.stat().st_size

            if existing_size == file_size:
                logger.warning(
                    f"Файл {safe_filename} уже существует, пропускаем загрузку"
                )
                # Но всё равно делаем импорт из существующего файла
                try:
                    result = importer.import_normative_document(
                        file_path=str(file_path),
                        title=title,
                        document_type=document_type,
                    )
                    return {
                        **result,
                        "message": "Файл уже существует, выполнен импорт из существующего",
                        "file_path": str(file_path),
                    }
                except Exception as e:
                    logger.error(f"Ошибка импорта существующего файла: {e}")
                    raise HTTPException(status_code=500, detail=f"Ошибка импорта: {e}")

        try:
            with open(file_path, "wb") as output_file:
                content = await file.read()
                output_file.write(content)
        except Exception as exc:
            logger.exception("Ошибка сохранения нормативного документа")
            raise HTTPException(
                status_code=500, detail=f"Не удалось сохранить файл: {exc}"
            )

        # Импортируем с AI-анализом
        try:
            logger.info(f"Начинаю импорт нормативного документа: {safe_filename}")
            result = importer.import_normative_document(
                file_path=str(file_path),
                title=title or safe_filename,
                document_type=document_type,
            )
            logger.info(
                f"✅ Импорт завершен: {result.get('rules_extracted', 0)} правил извлечено"
            )
            return {
                **result,
                "file_path": str(file_path),
                "filename": safe_filename,
            }
        except Exception as exc:
            logger.exception(f"Ошибка импорта нормативного документа: {exc}")
            raise HTTPException(
                status_code=500, detail=f"Ошибка импорта нормативного документа: {exc}"
            )

    except HTTPException:
        # Перебрасываем HTTPException как есть
        raise
    except Exception as exc:
        # Ловим любые другие неожиданные ошибки
        logger.exception(
            f"Неожиданная ошибка при загрузке нормативного документа: {exc}"
        )
        raise HTTPException(
            status_code=500, detail=f"Неожиданная ошибка при загрузке документа: {exc}"
        )


@app.get("/api/normative/documents")
def list_normative_documents():
    """Получить список всех загруженных нормативных документов"""
    try:
        documents = database.list_normative_documents()
        return {"documents": documents, "total": len(documents)}
    except Exception as exc:
        logger.exception("Ошибка получения списка нормативных документов")
        raise HTTPException(status_code=500, detail=f"Ошибка получения списка: {exc}")


@app.get("/api/normative/rules/{rule_type}")
def get_normative_rules_by_type(rule_type: str):
    """Получить правила определенного типа"""
    try:
        rules = database.get_normative_rules_by_type(rule_type)
        return {"rule_type": rule_type, "rules": rules, "total": len(rules)}
    except Exception as exc:
        logger.exception(f"Ошибка получения правил типа {rule_type}")
        raise HTTPException(status_code=500, detail=f"Ошибка получения правил: {exc}")


@app.get("/api/normative/rules/for-field/{field_name}")
def get_normative_rules_for_field(
    field_name: str,
    sheet_name: Optional[str] = Query(None),
):
    """Получить правила, связанные с конкретным полем энергопаспорта"""
    try:
        rules = database.get_normative_rules_for_field(field_name, sheet_name)
        return {
            "field_name": field_name,
            "sheet_name": sheet_name,
            "rules": rules,
            "total": len(rules),
        }
    except Exception as exc:
        logger.exception(f"Ошибка получения правил для поля {field_name}")
        raise HTTPException(status_code=500, detail=f"Ошибка получения правил: {exc}")


@app.get("/api/normative/documents/{document_id}")
def get_normative_document_by_id(document_id: int):
    """Получить нормативный документ по ID (включая метаданные)"""
    try:
        doc = database.get_normative_document(document_id)
        if not doc:
            raise HTTPException(status_code=404, detail=f"Документ с ID={document_id} не найден")
        
        # Возвращаем документ (без полного текста для экономии трафика)
        return {
            "document": {
                "id": doc.get("id"),
                "title": doc.get("title"),
                "document_type": doc.get("document_type"),
                "file_path": doc.get("file_path"),
                "file_size": doc.get("file_size"),
                "uploaded_at": doc.get("uploaded_at"),
                "ai_processed": doc.get("ai_processed"),
                "processing_status": doc.get("processing_status"),
            },
            "has_full_text": bool(doc.get("full_text")),
            "has_parsed_data": bool(doc.get("parsed_data_json")),
        }
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception(f"Ошибка получения документа {document_id}")
        raise HTTPException(status_code=500, detail=f"Ошибка получения документа: {exc}")


@app.get("/api/normative/documents/{document_id}/text")
def get_normative_document_text(document_id: int):
    """Получить полный текст нормативного документа"""
    try:
        doc = database.get_normative_document(document_id)
        if not doc:
            raise HTTPException(status_code=404, detail=f"Документ с ID={document_id} не найден")
        
        full_text = doc.get("full_text")
        if not full_text:
            return {
                "document_id": document_id,
                "title": doc.get("title"),
                "text": None,
                "message": "Полный текст не сохранен для этого документа",
            }
        
        return {
            "document_id": document_id,
            "title": doc.get("title"),
            "text": full_text,
            "text_length": len(full_text),
        }
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception(f"Ошибка получения текста документа {document_id}")
        raise HTTPException(status_code=500, detail=f"Ошибка получения текста: {exc}")


@app.post("/api/normative/validate-field")
def validate_field_against_normative(
    field_name: str,
    actual_value: float,
    sheet_name: Optional[str] = Query(None),
    tolerance_percent: float = Query(10.0),
):
    """Проверить соответствие фактического значения нормативу"""
    try:
        from domain.normative_validator import validate_against_normative
        
        result = validate_against_normative(
            actual_value=actual_value,
            field_name=field_name,
            sheet_name=sheet_name,
            tolerance_percent=tolerance_percent,
        )
        
        return {
            "field_name": field_name,
            "sheet_name": sheet_name,
            "validation": result,
        }
    except Exception as exc:
        logger.exception(f"Ошибка проверки поля {field_name}")
        raise HTTPException(status_code=500, detail=f"Ошибка проверки: {exc}")


@app.get("/api/normative/statistics")
def get_normative_statistics():
    """Получить статистику по нормативным документам"""
    try:
        from domain.normative_validator import get_normative_statistics
        
        stats = get_normative_statistics()
        return stats
    except Exception as exc:
        logger.exception("Ошибка получения статистики")
        raise HTTPException(status_code=500, detail=f"Ошибка получения статистики: {exc}")


@app.get("/api/normative/critical-fields/{enterprise_id}")
def check_critical_fields_for_enterprise(enterprise_id: int):
    """Проверить критические поля для предприятия"""
    try:
        from domain.normative_validator import check_critical_fields
        
        result = check_critical_fields(enterprise_id=enterprise_id)
        return result
    except Exception as exc:
        logger.exception(f"Ошибка проверки критических полей для предприятия {enterprise_id}")
        raise HTTPException(status_code=500, detail=f"Ошибка проверки: {exc}")


@app.get("/api/normative/violations")
def get_normative_violations(
    enterprise_id: Optional[int] = Query(None),
    batch_id: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    limit: int = Query(100, ge=1, le=1000),
):
    """Получить список нарушений нормативов"""
    try:
        violations = database.get_normative_violations(
            enterprise_id=enterprise_id,
            batch_id=batch_id,
            status=status,
            limit=limit,
        )
        return {
            "violations": violations,
            "total": len(violations),
            "filters": {
                "enterprise_id": enterprise_id,
                "batch_id": batch_id,
                "status": status,
            },
        }
    except Exception as exc:
        logger.exception("Ошибка получения нарушений нормативов")
        raise HTTPException(status_code=500, detail=f"Ошибка получения нарушений: {exc}")


@app.post("/api/normative/monitor-passport")
def monitor_passport_normatives(
    passport_path: str = Form(...),
    enterprise_id: Optional[int] = Form(None),
    batch_id: Optional[str] = Form(None),
):
    """Мониторинг критических полей из энергопаспорта"""
    try:
        from domain.normative_monitor import monitor_critical_fields_from_passport
        
        result = monitor_critical_fields_from_passport(
            passport_path=passport_path,
            enterprise_id=enterprise_id,
            batch_id=batch_id,
        )
        return result
    except Exception as exc:
        logger.exception(f"Ошибка мониторинга паспорта {passport_path}")
        raise HTTPException(status_code=500, detail=f"Ошибка мониторинга: {exc}")


@app.get("/api/normative/monitoring-summary")
def get_monitoring_summary(
    enterprise_id: Optional[int] = Query(None),
):
    """Получить сводку мониторинга критических полей"""
    try:
        from domain.normative_monitor import get_monitoring_summary
        
        summary = get_monitoring_summary(enterprise_id=enterprise_id)
        return summary
    except Exception as exc:
        logger.exception("Ошибка получения сводки мониторинга")
        raise HTTPException(status_code=500, detail=f"Ошибка получения сводки: {exc}")


@app.post("/api/enterprises/{enterprise_id}/reaggregate")
def api_reaggregate_enterprise(enterprise_id: int):
    """
    Переагрегирует все данные для предприятия.
    Используется для применения исправлений к уже загруженным файлам.
    """
    enterprise = database.get_enterprise_by_id(enterprise_id)
    if not enterprise:
        raise HTTPException(status_code=404, detail="Предприятие не найдено")

    try:
        from utils.reaggregate_all import reaggregate_enterprise

        logger.info(
            f"🔄 Запрос на переагрегацию для предприятия {enterprise_id} ({enterprise['name']})"
        )
        stats = reaggregate_enterprise(enterprise_id)
        logger.info(f"✅ Переагрегация завершена: {stats}")
        return {
            "enterprise_id": enterprise_id,
            "enterprise_name": enterprise["name"],
            "status": "success",
            "stats": stats,
        }
    except Exception as e:
        logger.exception(
            f"Ошибка при переагрегации для предприятия {enterprise_id}: {e}"
        )
        raise HTTPException(status_code=500, detail=f"Ошибка переагрегации: {e}")


@app.post("/api/validate-word-document")
async def validate_word_document(
    file: UploadFile = File(...),
    check_structure: bool = True,
    check_calculations: bool = True,
    check_compliance: bool = True
):
    """Проверить Word документ энергоаудита через AI"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Только .docx файлы")
    
    try:
        temp_path = DATA_DIR / f"temp_{uuid4().hex}.docx"
        with open(temp_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        from docx import Document
        doc = Document(temp_path)
        text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        tables_data = []
        for table in doc.tables:
            table_text = [" | ".join([cell.text.strip() for cell in row.cells]) for row in table.rows]
            tables_data.append("\n".join(table_text))
        
        prompt = f"""Проверь документ энергоаудита на ошибки.

ТЕКСТ ({len(text_content)} символов):
{text_content[:15000]}

ТАБЛИЦЫ ({len(tables_data)} шт):
{chr(10).join(tables_data[:5])}

ЗАДАЧИ:
{'✓ Структура ПКМ-690' if check_structure else ''}
{'✓ Расчеты' if check_calculations else ''}
{'✓ Нормативы' if check_compliance else ''}
✓ Орфография

JSON:
{{"overall_status": "OK|WARNINGS|ERRORS", "summary": "резюме", 
  "errors": [{{"type": "...", "severity": "...", "location": "...", "description": "...", "suggestion": "..."}}],
  "statistics": {{"critical_errors": 0, "warnings": 0}},
  "structure_check": {{"missing_sections": []}},
  "calculations_check": {{"inconsistencies": []}},
  "compliance_check": {{"pkm690_compliant": true}}
}}"""

        from ai_parser import AIParser
        ai_parser = AIParser()
        
        if not ai_parser.enabled:
            raise HTTPException(status_code=503, detail="AI не настроен")
        
        logger.info(f"🤖 Проверка через {ai_parser.provider}")
        ai_response = await ai_parser.parse_text(prompt)
        
        try:
            result = json.loads(ai_response)
        except:
            result = {"overall_status": "UNKNOWN", "summary": ai_response[:500]}
        
        temp_path.unlink()
        
        return {
            "status": "checked",
            "filename": file.filename,
            "validation_result": result,
            "ai_provider": ai_parser.provider
        }
        
    except Exception as e:
        logger.exception(f"Ошибка проверки: {e}")
        if 'temp_path' in locals() and temp_path.exists():
            temp_path.unlink()
        raise HTTPException(status_code=500, detail=str(e))

