"""
Pytest configuration and fixtures for EAIP ingest service tests.

This file provides reusable fixtures for:
- Test database (isolated SQLite)
- FastAPI test client
- Test files (Excel, PDF, images)
- Test enterprises
"""

import os
import sys
import tempfile
import shutil
from pathlib import Path
from typing import Generator, Dict, Any
import io

import pytest
from fastapi.testclient import TestClient
from openpyxl import Workbook


# ============================================================================
# Path Configuration
# ============================================================================

def pytest_configure(config):
    """
    Configure test environment before running tests.
    
    Sets up:
    - Test database path (isolated from production)
    - Module paths for imports
    """
    # Locate the ingest service root (parent of tests directory)
    test_dir = Path(__file__).resolve().parent
    ingest_root = test_dir.parent
    
    # Create temporary test database
    test_db_dir = test_dir / ".test_db"
    test_db_dir.mkdir(exist_ok=True)
    test_db_path = test_db_dir / "test_ingest.db"
    
    # Set environment variable for test database
    os.environ["INGEST_DB_PATH"] = str(test_db_path)
    os.environ["SYSTEM_MODE"] = "production"  # Use production mode to reuse existing uploads
    
    # Override database module if already imported
    try:
        import database as dbmod
        if hasattr(dbmod, "DB_PATH"):
            dbmod.DB_PATH = str(test_db_path)
    except ImportError:
        pass


# ============================================================================
# Database Fixtures
# ============================================================================

@pytest.fixture(scope="function")
def test_db() -> Generator[str, None, None]:
    """
    Create isolated test database for each test.
    
    Yields:
        str: Path to test database
        
    Cleanup:
        Database is cleared after each test
    """
    # Import database module (after pytest_configure sets path)
    import database
    
    # Initialize fresh database schema
    database.init_db()
    
    db_path = os.environ.get("INGEST_DB_PATH")
    
    yield db_path
    
    # Cleanup: Clear all tables after test
    with database.get_connection() as conn:
        cursor = conn.cursor()
        # Get all table names
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = cursor.fetchall()
        
        # Delete all data from tables
        for table in tables:
            table_name = table[0]
            if table_name != "sqlite_sequence":
                cursor.execute(f"DELETE FROM {table_name}")
        
        conn.commit()


# ============================================================================
# FastAPI Client Fixture
# ============================================================================

@pytest.fixture(scope="function")
def client(test_db) -> TestClient:
    """
    Create FastAPI test client with isolated database.
    
    Args:
        test_db: Test database fixture (ensures DB is ready)
        
    Returns:
        TestClient: FastAPI test client for making HTTP requests
        
    Example:
        def test_health_endpoint(client):
            response = client.get("/health")
            assert response.status_code == 200
    """
    # Import app after database is configured
    from main import app
    
    return TestClient(app)


# ============================================================================
# Enterprise Fixtures
# ============================================================================

@pytest.fixture(scope="function")
def test_enterprise(client) -> Dict[str, Any]:
    """
    Create test enterprise in database.
    
    Args:
        client: FastAPI test client
        
    Returns:
        dict: Enterprise data with id, name, created_at
        
    Example:
        def test_upload_file(client, test_enterprise):
            response = client.post(
                "/web/upload",
                data={"enterprise_id": test_enterprise["id"]},
                files={"file": ...}
            )
    """
    response = client.post(
        "/api/enterprises",
        json={"name": "Test Enterprise"}
    )
    
    assert response.status_code == 200, f"Failed to create test enterprise: {response.text}"
    
    return response.json()


# ============================================================================
# File Fixtures
# ============================================================================

@pytest.fixture(scope="function")
def test_excel_file() -> io.BytesIO:
    """
    Create minimal valid Excel file for testing.
    
    Returns:
        io.BytesIO: Excel file in memory
        
    Structure:
        Sheet1:
            A1: "Test"
            A2: "Data"
    """
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    
    # Add minimal data
    sheet["A1"] = "Test"
    sheet["A2"] = "Data"
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    workbook.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


@pytest.fixture(scope="function")
def test_excel_electricity_file() -> io.BytesIO:
    """
    Create Excel file with electricity consumption data.
    
    Returns:
        io.BytesIO: Excel file with electricity data
        
    Structure:
        Sheet1 (electricity consumption):
            Headers: Период, Активная (кВт*ч), Реактивная (кВАр*ч)
            Data: Q1 2024, 1000, 200
    """
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Электроэнергия"
    
    # Headers
    sheet["A1"] = "Период"
    sheet["B1"] = "Активная (кВт*ч)"
    sheet["C1"] = "Реактивная (кВАр*ч)"
    
    # Sample data
    sheet["A2"] = "2024-Q1"
    sheet["B2"] = 1000
    sheet["C2"] = 200
    
    # Save to BytesIO
    file_stream = io.BytesIO()
    workbook.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


@pytest.fixture(scope="function")
def test_large_file() -> io.BytesIO:
    """
    Create file larger than 50MB for size validation testing.
    
    Returns:
        io.BytesIO: File > 50MB
    """
    file_stream = io.BytesIO()
    
    # Write 51MB of data
    chunk_size = 1024 * 1024  # 1MB
    for _ in range(51):
        file_stream.write(b"X" * chunk_size)
    
    file_stream.seek(0)
    return file_stream


@pytest.fixture(scope="function")
def test_invalid_file() -> io.BytesIO:
    """
    Create invalid file (text file) for validation testing.
    
    Returns:
        io.BytesIO: Text file (not allowed)
    """
    file_stream = io.BytesIO()
    file_stream.write(b"This is a text file, not Excel")
    file_stream.seek(0)
    
    return file_stream


@pytest.fixture(scope="function")
def test_empty_file() -> io.BytesIO:
    """
    Create empty file for validation testing.
    
    Returns:
        io.BytesIO: Empty file (0 bytes)
    """
    return io.BytesIO()


@pytest.fixture(scope="function")
def test_pdf_file() -> io.BytesIO:
    """
    Create minimal valid PDF file for testing.
    
    Returns:
        io.BytesIO: PDF file in memory
        
    Note:
        This is a minimal valid PDF structure.
    """
    # Minimal valid PDF structure
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT
/F1 12 Tf
100 700 Td
(Test PDF) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000214 00000 n 
trailer
<< /Size 5 /Root 1 0 R >>
startxref
307
%%EOF
"""
    file_stream = io.BytesIO(pdf_content)
    file_stream.seek(0)
    return file_stream


@pytest.fixture(scope="function")
def test_docx_file() -> io.BytesIO:
    """
    Create minimal valid DOCX file for testing.
    
    Returns:
        io.BytesIO: DOCX file in memory
        
    Note:
        Uses python-docx if available, otherwise creates minimal ZIP structure.
    """
    try:
        from docx import Document
        
        doc = Document()
        doc.add_paragraph("Test Document")
        doc.add_paragraph("This is test content for EAIP testing.")
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except ImportError:
        # Fallback: create minimal DOCX (it's a ZIP with XML)
        import zipfile
        
        file_stream = io.BytesIO()
        with zipfile.ZipFile(file_stream, 'w', zipfile.ZIP_DEFLATED) as zf:
            # Minimal [Content_Types].xml
            content_types = '''<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
            zf.writestr('[Content_Types].xml', content_types)
            
            # Minimal _rels/.rels
            rels = '''<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>'''
            zf.writestr('_rels/.rels', rels)
            
            # Minimal word/document.xml
            document = '''<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>Test Document</w:t></w:r></w:p>
  </w:body>
</w:document>'''
            zf.writestr('word/document.xml', document)
        
        file_stream.seek(0)
        return file_stream


@pytest.fixture(scope="function")
def test_xlsm_file() -> io.BytesIO:
    """
    Create minimal valid XLSM file (Excel with macros) for testing.
    
    Returns:
        io.BytesIO: XLSM file in memory
        
    Note:
        XLSM is essentially XLSX with macro support flag.
    """
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sheet1"
    
    # Add data
    sheet["A1"] = "Test XLSM"
    sheet["A2"] = "Macro-enabled workbook"
    sheet["B1"] = 123
    sheet["B2"] = 456
    
    # Save as XLSM (openpyxl saves it correctly with .xlsm extension handling)
    file_stream = io.BytesIO()
    workbook.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


@pytest.fixture(scope="function")
def test_excel_with_multiple_sheets() -> io.BytesIO:
    """
    Create Excel file with multiple sheets for comprehensive testing.
    
    Returns:
        io.BytesIO: Excel file with electricity, gas, and equipment sheets
    """
    workbook = Workbook()
    
    # Sheet 1: Electricity
    ws1 = workbook.active
    ws1.title = "Электроэнергия"
    ws1["A1"] = "Период"
    ws1["B1"] = "Активная (кВт*ч)"
    ws1["C1"] = "Стоимость"
    ws1["A2"] = "2024-Q1"
    ws1["B2"] = 15000
    ws1["C2"] = 750000
    ws1["A3"] = "2024-Q2"
    ws1["B3"] = 18000
    ws1["C3"] = 900000
    
    # Sheet 2: Gas
    ws2 = workbook.create_sheet("Газ")
    ws2["A1"] = "Период"
    ws2["B1"] = "Объем (м³)"
    ws2["C1"] = "Стоимость"
    ws2["A2"] = "2024-Q1"
    ws2["B2"] = 5000
    ws2["C2"] = 250000
    
    # Sheet 3: Equipment
    ws3 = workbook.create_sheet("Оборудование")
    ws3["A1"] = "Наименование"
    ws3["B1"] = "Мощность (кВт)"
    ws3["C1"] = "Количество"
    ws3["A2"] = "Компрессор"
    ws3["B2"] = 75
    ws3["C2"] = 3
    
    file_stream = io.BytesIO()
    workbook.save(file_stream)
    file_stream.seek(0)
    
    return file_stream


# ============================================================================
# Temporary Directory Fixture
# ============================================================================

@pytest.fixture(scope="function")
def temp_dir() -> Generator[Path, None, None]:
    """
    Create temporary directory for test files.
    
    Yields:
        Path: Temporary directory path
        
    Cleanup:
        Directory is deleted after test
    """
    temp_path = Path(tempfile.mkdtemp(prefix="eaip_test_"))
    
    yield temp_path
    
    # Cleanup
    if temp_path.exists():
        shutil.rmtree(temp_path)


# ============================================================================
# Mock Data Fixtures
# ============================================================================

@pytest.fixture
def sample_aggregated_data() -> Dict[str, Any]:
    """
    Sample aggregated energy data for testing passport generation.
    
    Returns:
        dict: Aggregated data structure with electricity, gas, etc.
    """
    return {
        "resources": {
            "electricity": {
                "2024-Q1": {
                    "year": 2024,
                    "quarter": 1,
                    "quarter_totals": {
                        "active_kwh": 1000.0,
                        "reactive_kvarh": 200.0,
                        "cost_sum": 50000.0
                    },
                    "by_usage": {
                        "production": 800.0,
                        "lighting": 150.0,
                        "hvac": 50.0
                    }
                }
            }
        },
        "missing_sheets": []
    }


# ============================================================================
# Pytest Configuration Hooks
# ============================================================================

def pytest_collection_modifyitems(items):
    """
    Modify test collection to add markers automatically.
    
    Marks:
        - e2e: End-to-end integration tests
        - unit: Unit tests
        - slow: Tests that take > 1 second
    """
    for item in items:
        # Mark E2E tests
        if "e2e" in item.nodeid or "integration" in item.nodeid:
            item.add_marker(pytest.mark.e2e)
        
        # Mark unit tests
        if "unit" in item.nodeid:
            item.add_marker(pytest.mark.unit)
