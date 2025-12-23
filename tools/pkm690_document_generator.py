"""
Генератор стандартизированных документов энергоаудита по ПКМ 690 Узбекистан
Соблюдение всех требований форматирования и структуры
"""

import logging
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import sqlite3

# Настройка логирования
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)


class PKM690DocumentGenerator:
    """Генератор документов по стандарту ПКМ 690 Узбекистан"""

    def __init__(self, db_path: str = "energy_audit.db"):
        self.db_path = db_path
        self.doc = None

    def create_document(self, enterprise_id: int, output_path: str) -> bool:
        """Создание полного документа энергоаудита"""
        try:
            logger.info(f"📄 Создание документа для предприятия ID: {enterprise_id}")

            # Получаем данные предприятия
            enterprise_data = self.get_enterprise_data(enterprise_id)
            if not enterprise_data:
                logger.error(f"Предприятие с ID {enterprise_id} не найдено")
                return False

            # Создаем новый документ
            self.doc = Document()

            # Настраиваем стили документа
            self.setup_document_styles()

            # Создаем структуру документа по ПКМ 690
            self.create_title_page(enterprise_data)
            self.create_table_of_contents()
            self.create_introduction(enterprise_data)
            self.create_enterprise_info(enterprise_data)
            self.create_energy_consumption_analysis(enterprise_data)
            self.create_equipment_analysis(enterprise_data)
            self.create_energy_efficiency_measures(enterprise_data)
            self.create_economic_analysis(enterprise_data)
            self.create_conclusion(enterprise_data)
            self.create_appendix(enterprise_data)

            # Сохраняем документ
            self.doc.save(output_path)
            logger.info(f"✅ Документ сохранен: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Ошибка создания документа: {e}")
            return False

    def setup_document_styles(self):
        """Настройка стилей документа по стандарту ПКМ 690"""

        # Основной стиль текста
        normal_style = self.doc.styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "Times New Roman"
        normal_font.size = Pt(14)
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_after = Pt(6)

        # Стиль заголовков
        heading_styles = {
            "Heading 1": {
                "size": 16,
                "bold": True,
                "space_before": 12,
                "space_after": 6,
            },
            "Heading 2": {
                "size": 15,
                "bold": True,
                "space_before": 10,
                "space_after": 6,
            },
            "Heading 3": {
                "size": 14,
                "bold": True,
                "space_before": 8,
                "space_after": 6,
            },
        }

        for style_name, params in heading_styles.items():
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                font = style.font
                font.name = "Times New Roman"
                font.size = Pt(params["size"])
                font.bold = params["bold"]

                para_format = style.paragraph_format
                para_format.space_before = Pt(params["space_before"])
                para_format.space_after = Pt(params["space_after"])
                para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Создаем стиль для таблиц
        table_style = self.doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
        table_font = table_style.font
        table_font.name = "Times New Roman"
        table_font.size = Pt(12)
        table_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Стиль для заголовков таблиц
        table_header_style = self.doc.styles.add_style(
            "Table Header", WD_STYLE_TYPE.PARAGRAPH
        )
        table_header_font = table_header_style.font
        table_header_font.name = "Times New Roman"
        table_header_font.size = Pt(12)
        table_header_font.bold = True
        table_header_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def create_title_page(self, enterprise_data: dict):
        """Создание титульной страницы"""
        logger.info("📋 Создание титульной страницы...")

        # Заголовок документа
        title = self.doc.add_heading("ЭНЕРГЕТИЧЕСКИЙ АУДИТ", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Подзаголовок
        subtitle = self.doc.add_heading("Энергетический паспорт предприятия", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пустые строки
        for _ in range(3):
            self.doc.add_paragraph()

        # Информация о предприятии
        enterprise_info = f"""
НАЗВАНИЕ ПРЕДПРИЯТИЯ: {enterprise_data["name"]}
ИНН: {enterprise_data["inn"] or "не указан"}
АДРЕС: {enterprise_data["address"] or "не указан"}
ДИРЕКТОР: {enterprise_data["director_name"] or "не указан"}
ОТРАСЛЬ: {enterprise_data["industry"] or "не указана"}
ГОД ОТЧЕТА: {enterprise_data["reporting_year"] or datetime.now().year}
"""

        info_para = self.doc.add_paragraph(enterprise_info)
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пустые строки
        for _ in range(5):
            self.doc.add_paragraph()

        # Информация об аудиторе
        auditor_info = f"""
АУДИТОР: {enterprise_data.get("auditor", "Не указан")}
ДАТА ПРОВЕДЕНИЯ: {datetime.now().strftime("%d.%m.%Y")}
СТАНДАРТ: ПКМ 690 Узбекистан
"""

        auditor_para = self.doc.add_paragraph(auditor_info)
        auditor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Разрыв страницы
        self.doc.add_page_break()

    def create_table_of_contents(self):
        """Создание содержания"""
        logger.info("📑 Создание содержания...")

        title = self.doc.add_heading("СОДЕРЖАНИЕ", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Список разделов по ПКМ 690
        sections = [
            "1. ВВЕДЕНИЕ",
            "2. ОБЩИЕ СВЕДЕНИЯ О ПРЕДПРИЯТИИ",
            "3. АНАЛИЗ ЭНЕРГОПОТРЕБЛЕНИЯ",
            "4. АНАЛИЗ ОБОРУДОВАНИЯ",
            "5. МЕРОПРИЯТИЯ ПО ЭНЕРГОСБЕРЕЖЕНИЮ",
            "6. ЭКОНОМИЧЕСКИЙ АНАЛИЗ",
            "7. ЗАКЛЮЧЕНИЕ",
            "8. ПРИЛОЖЕНИЯ",
        ]

        for section in sections:
            para = self.doc.add_paragraph(section)
            para.style = "Normal"

        # Разрыв страницы
        self.doc.add_page_break()

    def create_introduction(self, enterprise_data: dict):
        """Создание введения"""
        logger.info("📝 Создание введения...")

        title = self.doc.add_heading("1. ВВЕДЕНИЕ", level=1)

        intro_text = f"""
Настоящий энергетический аудит проведен в соответствии с требованиями стандарта ПКМ 690 Узбекистан 
"Энергетические обследования. Общие требования" для предприятия "{enterprise_data["name"]}".

Целью энергетического аудита является:
- Определение фактического энергопотребления предприятия
- Выявление резервов энергосбережения
- Разработка мероприятий по повышению энергоэффективности
- Оценка экономической эффективности предлагаемых мероприятий

Объектом обследования является предприятие "{enterprise_data["name"]}", расположенное по адресу: 
{enterprise_data["address"] or "адрес не указан"}.

Период проведения аудита: {datetime.now().strftime("%d.%m.%Y")}
Аудитор: {enterprise_data.get("auditor", "Не указан")}

Методология проведения энергетического аудита основана на:
- Анализе энергопотребления за отчетный период
- Обследовании энергетического оборудования
- Расчете энергетических показателей
- Разработке рекомендаций по энергосбережению
"""

        self.doc.add_paragraph(intro_text)

    def create_enterprise_info(self, enterprise_data: dict):
        """Создание раздела общих сведений о предприятии"""
        logger.info("🏢 Создание раздела о предприятии...")

        title = self.doc.add_heading("2. ОБЩИЕ СВЕДЕНИЯ О ПРЕДПРИЯТИИ", level=1)

        # Основная информация
        basic_info = f"""
2.1 Основная информация

Полное наименование предприятия: {enterprise_data["name"]}
ИНН: {enterprise_data["inn"] or "не указан"}
Юридический адрес: {enterprise_data["address"] or "не указан"}
Руководитель: {enterprise_data["director_name"] or "не указан"}
Отрасль деятельности: {enterprise_data["industry"] or "не указана"}
Год основания: {enterprise_data.get("founding_year", "не указан")}
"""

        self.doc.add_paragraph(basic_info)

        # Создаем таблицу с характеристиками предприятия
        self.doc.add_heading("2.2 Характеристики предприятия", level=2)

        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Показатель"
        hdr_cells[1].text = "Значение"

        # Применяем стиль к заголовкам
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.style = "Table Header"

        # Данные таблицы
        characteristics = [
            ("Площадь территории", enterprise_data.get("territory_area", "не указана")),
            ("Площадь зданий", enterprise_data.get("building_area", "не указана")),
            (
                "Количество сотрудников",
                enterprise_data.get("employees_count", "не указано"),
            ),
            ("Режим работы", enterprise_data.get("work_schedule", "не указан")),
            ("Основная продукция", enterprise_data.get("main_products", "не указана")),
        ]

        for char_name, char_value in characteristics:
            row_cells = table.add_row().cells
            row_cells[0].text = char_name
            row_cells[1].text = char_value

            # Применяем стиль к ячейкам
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = "Table Text"

    def create_energy_consumption_analysis(self, enterprise_data: dict):
        """Создание анализа энергопотребления"""
        logger.info("⚡ Создание анализа энергопотребления...")

        title = self.doc.add_heading("3. АНАЛИЗ ЭНЕРГОПОТРЕБЛЕНИЯ", level=1)

        # Получаем данные по энергопотреблению
        energy_data = self.get_energy_consumption_data(enterprise_data["id"])

        analysis_text = f"""
3.1 Общая характеристика энергопотребления

Предприятие потребляет следующие виды энергетических ресурсов:
- Электрическая энергия
- Природный газ
- Вода (холодная и горячая)
- Топливно-энергетические ресурсы

Общее энергопотребление предприятия за отчетный период составляет:
- Электрическая энергия: {energy_data.get("electricity_total", "не указано")} кВт·ч/год
- Природный газ: {energy_data.get("gas_total", "не указано")} м³/год
- Вода: {energy_data.get("water_total", "не указано")} м³/год
"""

        self.doc.add_paragraph(analysis_text)

        # Таблица энергопотребления по месяцам
        self.doc.add_heading("3.2 Энергопотребление по месяцам", level=2)

        monthly_data = self.get_monthly_energy_data(enterprise_data["id"])
        if monthly_data:
            table = self.doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = [
                "Месяц",
                "Электричество (кВт·ч)",
                "Газ (м³)",
                "Вода (м³)",
                "Общая стоимость (руб)",
            ]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.style = "Table Header"

            # Данные
            for month_data in monthly_data:
                row_cells = table.add_row().cells
                row_cells[0].text = month_data["month"]
                row_cells[1].text = str(month_data.get("electricity", 0))
                row_cells[2].text = str(month_data.get("gas", 0))
                row_cells[3].text = str(month_data.get("water", 0))
                row_cells[4].text = str(month_data.get("total_cost", 0))

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = "Table Text"

    def create_equipment_analysis(self, enterprise_data: dict):
        """Создание анализа оборудования"""
        logger.info("🏭 Создание анализа оборудования...")

        title = self.doc.add_heading("4. АНАЛИЗ ОБОРУДОВАНИЯ", level=1)

        equipment_data = self.get_equipment_data(enterprise_data["id"])

        analysis_text = f"""
4.1 Общая характеристика оборудования

На предприятии установлено следующее энергопотребляющее оборудование:
- Производственное оборудование
- Электрооборудование
- Системы отопления и вентиляции
- Осветительные установки

Общая установленная мощность: {equipment_data.get("total_power", "не указана")} кВт
"""

        self.doc.add_paragraph(analysis_text)

        # Таблица оборудования
        if equipment_data.get("equipment_list"):
            self.doc.add_heading("4.2 Перечень основного оборудования", level=2)

            table = self.doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = ["Наименование", "Тип", "Мощность (кВт)", "Год установки"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.style = "Table Header"

            # Данные оборудования
            for equipment in equipment_data["equipment_list"][
                :10
            ]:  # Показываем первые 10
                row_cells = table.add_row().cells
                row_cells[0].text = equipment.get("name", "не указано")
                row_cells[1].text = equipment.get("type", "не указан")
                row_cells[2].text = str(equipment.get("power", 0))
                row_cells[3].text = str(equipment.get("year", "не указан"))

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = "Table Text"

    def create_energy_efficiency_measures(self, enterprise_data: dict):
        """Создание раздела мероприятий по энергосбережению"""
        logger.info("💡 Создание раздела мероприятий...")

        title = self.doc.add_heading("5. МЕРОПРИЯТИЯ ПО ЭНЕРГОСБЕРЕЖЕНИЮ", level=1)

        measures_data = self.get_energy_efficiency_measures(enterprise_data["id"])

        intro_text = """
5.1 Общие принципы энергосбережения

На основе проведенного анализа энергопотребления разработаны следующие мероприятия 
по повышению энергоэффективности предприятия:

1. Технические мероприятия
2. Организационные мероприятия
3. Информационные мероприятия
"""

        self.doc.add_paragraph(intro_text)

        # Таблица мероприятий
        if measures_data:
            self.doc.add_heading("5.2 Перечень мероприятий", level=2)

            table = self.doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"

            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = [
                "Мероприятие",
                "Тип",
                "Экономия (%)",
                "Стоимость (руб)",
                "Срок окупаемости",
            ]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    paragraph.style = "Table Header"

            # Данные мероприятий
            for measure in measures_data[:10]:  # Показываем первые 10
                row_cells = table.add_row().cells
                row_cells[0].text = measure.get("name", "не указано")
                row_cells[1].text = measure.get("type", "не указан")
                row_cells[2].text = str(measure.get("savings_percent", 0))
                row_cells[3].text = str(measure.get("cost", 0))
                row_cells[4].text = str(measure.get("payback_period", 0))

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.style = "Table Text"

    def create_economic_analysis(self, enterprise_data: dict):
        """Создание экономического анализа"""
        logger.info("💰 Создание экономического анализа...")

        title = self.doc.add_heading("6. ЭКОНОМИЧЕСКИЙ АНАЛИЗ", level=1)

        economic_text = f"""
6.1 Анализ затрат на энергоресурсы

Общие затраты предприятия на энергоресурсы за отчетный период:
- Электрическая энергия: {enterprise_data.get("electricity_cost", "не указано")} руб/год
- Природный газ: {enterprise_data.get("gas_cost", "не указано")} руб/год
- Вода: {enterprise_data.get("water_cost", "не указано")} руб/год
- Общие затраты: {enterprise_data.get("total_energy_cost", "не указано")} руб/год

6.2 Экономический эффект от мероприятий

Реализация предложенных мероприятий позволит:
- Снизить энергопотребление на 15-25%
- Экономить денежные средства на оплате энергоресурсов
- Повысить конкурентоспособность предприятия
- Улучшить экологическую ситуацию

6.3 Срок окупаемости мероприятий

Средний срок окупаемости предлагаемых мероприятий составляет 2-4 года.
"""

        self.doc.add_paragraph(economic_text)

    def create_conclusion(self, enterprise_data: dict):
        """Создание заключения"""
        logger.info("📝 Создание заключения...")

        title = self.doc.add_heading("7. ЗАКЛЮЧЕНИЕ", level=1)

        conclusion_text = f"""
На основании проведенного энергетического аудита предприятия "{enterprise_data["name"]}" 
можно сделать следующие выводы:

1. ОБЩАЯ ОЦЕНКА ЭНЕРГОПОТРЕБЛЕНИЯ
Предприятие потребляет значительное количество энергетических ресурсов. 
Анализ показал возможности для повышения энергоэффективности.

2. ОСНОВНЫЕ НАПРАВЛЕНИЯ ЭНЕРГОСБЕРЕЖЕНИЯ
- Оптимизация систем освещения
- Модернизация отопительных систем
- Внедрение энергосберегающего оборудования
- Улучшение теплоизоляции зданий

3. ЭКОНОМИЧЕСКИЙ ЭФФЕКТ
Реализация предложенных мероприятий позволит:
- Снизить энергопотребление на 15-25%
- Экономить денежные средства на оплате энергоресурсов
- Повысить конкурентоспособность предприятия

4. РЕКОМЕНДАЦИИ
Рекомендуется:
- Разработать программу энергосбережения
- Внедрить систему энергетического менеджмента
- Регулярно проводить энергетические аудиты
- Обучить персонал основам энергосбережения

Заключение подготовлено на основании данных энергетического аудита, 
проведенного в соответствии с требованиями стандарта ПКМ 690 Узбекистан.
"""

        self.doc.add_paragraph(conclusion_text)

    def create_appendix(self, enterprise_data: dict):
        """Создание приложений"""
        logger.info("📎 Создание приложений...")

        title = self.doc.add_heading("8. ПРИЛОЖЕНИЯ", level=1)

        # Приложение 1
        self.doc.add_heading("Приложение 1. Справочные данные", level=2)

        reference_text = f"""
Таблица 1.1 - Основные характеристики предприятия

| Показатель | Значение |
|------------|----------|
| Название предприятия | {enterprise_data["name"]} |
| ИНН | {enterprise_data["inn"] or "не указан"} |
| Адрес | {enterprise_data["address"] or "не указан"} |
| Директор | {enterprise_data["director_name"] or "не указан"} |
| Отрасль | {enterprise_data["industry"] or "не указана"} |
| Год отчета | {enterprise_data["reporting_year"] or "не указан"} |
| Аудитор | {enterprise_data.get("auditor", "не указан")} |
"""

        self.doc.add_paragraph(reference_text)

        # Приложение 2
        self.doc.add_heading("Приложение 2. Нормативные документы", level=2)

        normative_text = """
Список нормативных документов, использованных при проведении энергетического аудита:

1. ПКМ 690 Узбекистан "Энергетические обследования. Общие требования"
2. Закон Республики Узбекистан "Об энергосбережении"
3. ГОСТ Р 51387-99 "Энергосбережение. Нормативно-методическое обеспечение"
4. Методические рекомендации по проведению энергетических обследований
5. Санитарные правила и нормы (СанПиН)
6. Строительные нормы и правила (СНиП)
7. Правила устройства электроустановок (ПУЭ)
8. Правила технической эксплуатации электроустановок потребителей (ПТЭЭП)
"""

        self.doc.add_paragraph(normative_text)

    def get_enterprise_data(self, enterprise_id: int) -> dict:
        """Получение данных предприятия из базы данных"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT * FROM enterprises WHERE id = ?
            """,
                (enterprise_id,),
            )

            row = cursor.fetchone()
            if row:
                columns = [description[0] for description in cursor.description]
                enterprise_data = dict(zip(columns, row))
                enterprise_data["id"] = enterprise_id
                return enterprise_data

            conn.close()
            return None

        except Exception as e:
            logger.error(f"Ошибка получения данных предприятия: {e}")
            return None

    def get_energy_consumption_data(self, enterprise_id: int) -> dict:
        """Получение данных по энергопотреблению"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # Получаем данные по энергопотреблению
            cursor.execute(
                """
                SELECT 
                    SUM(electricity_consumption) as electricity_total,
                    SUM(gas_consumption) as gas_total,
                    SUM(water_consumption) as water_total
                FROM energy_consumption 
                WHERE enterprise_id = ?
            """,
                (enterprise_id,),
            )

            row = cursor.fetchone()
            if row:
                return {
                    "electricity_total": row[0] or 0,
                    "gas_total": row[1] or 0,
                    "water_total": row[2] or 0,
                }

            conn.close()
            return {}

        except Exception as e:
            logger.error(f"Ошибка получения данных энергопотребления: {e}")
            return {}

    def get_monthly_energy_data(self, enterprise_id: int) -> list:
        """Получение месячных данных по энергопотреблению"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT 
                    month,
                    electricity_consumption as electricity,
                    gas_consumption as gas,
                    water_consumption as water,
                    total_cost
                FROM energy_consumption 
                WHERE enterprise_id = ?
                ORDER BY month
            """,
                (enterprise_id,),
            )

            rows = cursor.fetchall()
            columns = [description[0] for description in cursor.description]

            monthly_data = []
            for row in rows:
                monthly_data.append(dict(zip(columns, row)))

            conn.close()
            return monthly_data

        except Exception as e:
            logger.error(f"Ошибка получения месячных данных: {e}")
            return []

    def get_equipment_data(self, enterprise_id: int) -> dict:
        """Получение данных по оборудованию"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # Получаем список оборудования
            cursor.execute(
                """
                SELECT name, type, power, year_installed as year
                FROM equipment 
                WHERE enterprise_id = ?
                ORDER BY power DESC
            """,
                (enterprise_id,),
            )

            equipment_list = []
            total_power = 0

            for row in cursor.fetchall():
                equipment = {
                    "name": row[0],
                    "type": row[1],
                    "power": row[2] or 0,
                    "year": row[3],
                }
                equipment_list.append(equipment)
                total_power += equipment["power"]

            conn.close()
            return {"equipment_list": equipment_list, "total_power": total_power}

        except Exception as e:
            logger.error(f"Ошибка получения данных оборудования: {e}")
            return {"equipment_list": [], "total_power": 0}

    def get_energy_efficiency_measures(self, enterprise_id: int) -> list:
        """Получение мероприятий по энергосбережению"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            cursor.execute(
                """
                SELECT 
                    name, 
                    type, 
                    savings_percent, 
                    cost, 
                    payback_period
                FROM energy_efficiency_measures 
                WHERE enterprise_id = ?
                ORDER BY savings_percent DESC
            """,
                (enterprise_id,),
            )

            measures = []
            for row in cursor.fetchall():
                measure = {
                    "name": row[0],
                    "type": row[1],
                    "savings_percent": row[2] or 0,
                    "cost": row[3] or 0,
                    "payback_period": row[4] or 0,
                }
                measures.append(measure)

            conn.close()
            return measures

        except Exception as e:
            logger.error(f"Ошибка получения мероприятий: {e}")
            return []


def main():
    """Основная функция"""
    print("\n" + "=" * 70)
    print("║     📄 ГЕНЕРАТОР ДОКУМЕНТОВ ПО ПКМ 690 УЗБЕКИСТАН        ║")
    print("=" * 70)

    # Создаем генератор
    generator = PKM690DocumentGenerator()

    # Список предприятий для генерации документов
    enterprises = [
        {"id": 1, "name": "Metin Iroda"},
        {"id": 2, "name": "Test Enterprise"},
    ]

    for enterprise in enterprises:
        output_path = f"energy_audit_report_{enterprise['id']}_{enterprise['name'].replace(' ', '_')}.docx"

        print(f"\n📄 Генерация документа для: {enterprise['name']}")

        success = generator.create_document(enterprise["id"], output_path)

        if success:
            print(f"✅ Документ создан: {output_path}")
        else:
            print(f"❌ Ошибка создания документа для {enterprise['name']}")

    print("\n" + "=" * 70)
    print("✅ ГЕНЕРАЦИЯ ДОКУМЕНТОВ ЗАВЕРШЕНА!")
    print("=" * 70)


if __name__ == "__main__":
    main()
