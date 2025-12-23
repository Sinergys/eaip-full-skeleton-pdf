"""
Конвертация результатов ручного распознавания из Excel/Word в JSON формат
"""
import sys
from pathlib import Path
import json

# Добавляем путь к проекту
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

def read_excel_file(file_path: str) -> dict:
    """Читает Excel файл и извлекает таблицы"""
    try:
        import openpyxl
    except ImportError:
        raise ImportError("Требуется openpyxl: pip install openpyxl")
    
    workbook = openpyxl.load_workbook(file_path, data_only=True)
    result = {
        "tables": [],
        "text": ""
    }
    
    # Обрабатываем все листы
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        
        # Находим таблицу (данные начинаются с первой строки)
        rows = []
        headers = []
        
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
            # Пропускаем пустые строки
            if not any(cell for cell in row if cell):
                continue
            
            # Конвертируем значения в строки
            row_data = [str(cell) if cell is not None else "" for cell in row]
            
            # Первая непустая строка - заголовки
            if not headers and row_data:
                headers = row_data
            else:
                rows.append(row_data)
        
        if headers or rows:
            result["tables"].append({
                "rows": [headers] + rows if headers else rows,
                "headers": headers if headers else []
            })
    
    return result

def read_word_file(file_path: str) -> dict:
    """Читает Word файл и извлекает таблицы и текст"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Требуется python-docx: pip install python-docx")
    
    doc = Document(file_path)
    result = {
        "tables": [],
        "text": ""
    }
    
    # Извлекаем текст
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    result["text"] = "\n".join(text_parts)
    
    # Извлекаем таблицы
    for table in doc.tables:
        rows = []
        headers = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            
            # Первая строка - заголовки
            if row_idx == 0:
                headers = row_data
            else:
                rows.append(row_data)
        
        if headers or rows:
            result["tables"].append({
                "rows": [headers] + rows if headers else rows,
                "headers": headers if headers else []
            })
    
    return result

def convert_file_to_json(file_path: str, output_path: str = None) -> str:
    """Конвертирует файл в JSON формат"""
    file_path_obj = Path(file_path)
    
    if not file_path_obj.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")
    
    # Определяем тип файла
    if file_path_obj.suffix.lower() in ['.xlsx', '.xls']:
        data = read_excel_file(str(file_path_obj))
    elif file_path_obj.suffix.lower() == '.docx':
        data = read_word_file(str(file_path_obj))
    else:
        raise ValueError(f"Неподдерживаемый формат: {file_path_obj.suffix}")
    
    # Сохраняем в JSON
    if output_path is None:
        output_path = file_path_obj.parent / f"{file_path_obj.stem}_converted.json"
    else:
        output_path = Path(output_path)
    
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    return str(output_path)

def main():
    """Основная функция"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Конвертация результатов ручного распознавания в JSON')
    parser.add_argument('files', nargs='+', help='Пути к файлам для конвертации')
    parser.add_argument('--output-dir', type=str, help='Директория для сохранения JSON файлов')
    
    args = parser.parse_args()
    
    print("=" * 80)
    print("КОНВЕРТАЦИЯ РЕЗУЛЬТАТОВ РУЧНОГО РАСПОЗНАВАНИЯ В JSON")
    print("=" * 80)
    print()
    
    output_dir = Path(args.output_dir) if args.output_dir else project_root / "reports" / "ocr"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    converted_files = []
    
    for file_path in args.files:
        print(f"📄 Обработка: {file_path}")
        try:
            output_file = output_dir / f"{Path(file_path).stem}_converted.json"
            converted_path = convert_file_to_json(file_path, str(output_file))
            converted_files.append(converted_path)
            print(f"✅ Конвертировано: {converted_path}")
            
            # Показываем статистику
            with open(converted_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                print(f"   - Таблиц: {len(data.get('tables', []))}")
                if data.get('tables'):
                    for i, table in enumerate(data['tables'], 1):
                        print(f"     Таблица {i}: {len(table.get('rows', []))} строк, {len(table.get('headers', []))} столбцов")
                print(f"   - Длина текста: {len(data.get('text', ''))} символов")
            print()
            
        except Exception as e:
            print(f"❌ Ошибка обработки {file_path}: {e}")
            print()
    
    if converted_files:
        print("=" * 80)
        print("✅ КОНВЕРТАЦИЯ ЗАВЕРШЕНА")
        print("=" * 80)
        print()
        print("📋 Конвертированные файлы:")
        for f in converted_files:
            print(f"  - {f}")
        print()
        print("📌 Теперь можно запустить сравнение:")
        if len(converted_files) == 1:
            print(f"  python tools/compare_recognition_results.py --manual \"{converted_files[0]}\"")
        else:
            print(f"  python tools/compare_recognition_results.py --manual1 \"{converted_files[0]}\" --manual2 \"{converted_files[1]}\"")

if __name__ == "__main__":
    main()

