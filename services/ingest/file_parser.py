"""
Модуль для парсинга загруженных файлов согласно ТЗ (раздел 4.1, Модуль 1)
Поддерживает распознавание и структурирование данных из различных форматов
"""

import os
import re
from pathlib import Path
from typing import Dict, Any, Optional
import logging

logger = logging.getLogger(__name__)

# Импорт для проверки отмены обработки
try:
    from utils.progress_tracker import get_progress_tracker

    HAS_PROGRESS_TRACKER = True
except ImportError:
    HAS_PROGRESS_TRACKER = False
    logger.warning(
        "progress_tracker недоступен. Проверка отмены обработки будет пропущена."
    )

# Импорт ИИ парсера (опционально)
try:
    from ai_parser import get_ai_parser

    HAS_AI_PARSER = True
except ImportError:
    HAS_AI_PARSER = False
    logger.warning("ai_parser модуль не найден. ИИ функции недоступны.")

# Импорт AI-модулей для усиления OCR (опционально)
try:
    from ai_ocr_enhancer import get_ocr_ai_enhancer

    HAS_OCR_AI_ENHANCER = True
except ImportError:
    HAS_OCR_AI_ENHANCER = False
    logger.debug("ai_ocr_enhancer модуль не найден. AI-усиление OCR недоступно.")

try:
    from ai_data_validator import get_ai_data_validator

    HAS_AI_DATA_VALIDATOR = True
except ImportError:
    HAS_AI_DATA_VALIDATOR = False
    logger.debug("ai_data_validator модуль не найден. AI-валидация данных недоступна.")

try:
    from ai_table_parser import get_ai_table_parser

    HAS_AI_TABLE_PARSER = True
except ImportError:
    HAS_AI_TABLE_PARSER = False
    logger.debug(
        "ai_table_parser модуль не найден. AI-структурирование таблиц недоступно."
    )

# Импорт AI-модулей для полного анализа энергетических данных (опционально)
try:
    from ai_energy_analysis import enhanced_energy_analysis

    HAS_AI_ENERGY_ANALYSIS = True
except ImportError:
    HAS_AI_ENERGY_ANALYSIS = False
    logger.debug(
        "ai_energy_analysis модуль не найден. AI-анализ энергетических данных недоступен."
    )

try:
    from ai_quality_reporter import get_quality_reporter

    HAS_AI_QUALITY_REPORTER = True
except ImportError:
    HAS_AI_QUALITY_REPORTER = False
    logger.debug(
        "ai_quality_reporter модуль не найден. Генерация отчетов о качестве недоступна."
    )

try:
    from openpyxl import load_workbook

    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    logger.warning("openpyxl не установлен. Парсинг Excel недоступен.")

try:
    import pandas  # noqa: F401

    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logger.warning("pandas не установлен. Расширенный парсинг Excel недоступен.")

try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx не установлен. Парсинг DOCX недоступен.")

try:
    import fitz  # PyMuPDF

    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF не установлен. Базовый парсинг PDF недоступен.")

try:
    import pdfplumber

    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    logger.warning("pdfplumber не установлен. Расширенный парсинг PDF недоступен.")

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image, ImageEnhance

    # Сохраняем Image для использования в функциях
    PIL_Image = Image
    PIL_ImageEnhance = ImageEnhance
    
    # Увеличиваем лимит PIL для больших изображений (для сканированных PDF с высоким DPI)
    PIL_Image.MAX_IMAGE_PIXELS = 200000000  # 200 миллионов пикселей (было 89 миллионов)
    
    # Автоматическое определение пути к Tesseract (если не в PATH)
    tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        r"C:\Tesseract-OCR\tesseract.exe",
    ]

    # Проверяем, установлен ли путь к tesseract и существует ли файл
    current_cmd = pytesseract.pytesseract.tesseract_cmd
    # Если путь не указан явно или файл не существует, ищем в стандартных местах
    if not current_cmd or current_cmd == "tesseract" or not os.path.exists(current_cmd):
        # Ищем Tesseract в стандартных путях
        for tesseract_path in tesseract_paths:
            if os.path.exists(tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = tesseract_path
                logger.info(f"Автоматически найден Tesseract: {tesseract_path}")
                break
        else:
            # Если не нашли, но tesseract может быть в PATH, проверяем через which/where
            logger.warning(
                "Tesseract не найден в стандартных путях. Убедитесь, что он установлен и добавлен в PATH."
            )

    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    PIL_Image = None
    PIL_ImageEnhance = None
    logger.warning("pytesseract или pdf2image не установлены. OCR недоступен.")


def parse_excel_file(
    file_path: str, template_name: Optional[str] = None
) -> Dict[str, Any]:
    """
    Парсинг Excel файла согласно ТЗ (раздел 4.1)

    Args:
        file_path: Путь к файлу
        template_name: Имя эталонного шаблона (например, 'EnergyPassport_v1.1.2.xlsx')

    Returns:
        Словарь с распознанными данными
    """
    if not HAS_OPENPYXL:
        raise ImportError("openpyxl не установлен. Установите: pip install openpyxl")

    result = {
        "file_path": file_path,
        "template_detected": False,
        "sheets": [],
        "data": {},
    }

    try:
        wb = load_workbook(file_path, data_only=True)

        # Проверка на эталонный шаблон EnergyPassport_v1.1.2.xlsx
        if template_name and template_name in file_path:
            result["template_detected"] = True

        # Обработка всех листов
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            sheet_data = {
                "name": sheet_name,
                "rows": [],
                "max_row": sheet.max_row,
                "max_column": sheet.max_column,
            }

            # Чтение данных из листа
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    sheet_data["rows"].append(list(row))

            result["sheets"].append(sheet_data)
            result["data"][sheet_name] = sheet_data

        logger.info(
            f"Успешно распарсен Excel файл: {file_path}, листов: {len(result['sheets'])}"
        )

        # Optional: collect CanonicalSourceData (deterministic-only for now)
        try:
            from utils.canonical_collector import collect_canonical_from_workbook

            canonical = collect_canonical_from_workbook(file_path)
            if canonical:
                # Store as lightweight dict for JSON-serializability
                result["canonical_source"] = canonical.dict()
                logger.debug("CanonicalSourceData collected for %s", file_path)
        except Exception as _exc:
            logger.debug(f"Canonical collection skipped due to error: {_exc}")

    except Exception as e:
        logger.error(f"Ошибка при парсинге Excel файла {file_path}: {e}")
        raise

    return result


def parse_docx_file(file_path: str) -> Dict[str, Any]:
    """
    Парсинг DOCX файла согласно ТЗ (раздел 4.1)

    Если в файле уже есть таблицы с данными (готовый отчет), они извлекаются
    и помечаются как готовые для использования без повторной обработки.

    Args:
        file_path: Путь к файлу

    Returns:
        Словарь с распознанными данными, включая метаданные о готовых таблицах
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx не установлен. Установите: pip install python-docx"
        )

    result = {
        "file_path": file_path,
        "paragraphs": [],
        "tables": [],
        "is_ready_report": False,  # Флаг готового отчета
        "table_count": 0,
    }

    try:
        doc = Document(file_path)
        file_name_lower = Path(file_path).name.lower()

        # Проверяем, является ли файл готовым отчетом по имени
        is_report_file = any(
            keyword in file_name_lower
            for keyword in ["отчет", "report", "финал", "final", "паспорт", "passport"]
        )

        # Извлечение параграфов
        for para in doc.paragraphs:
            if para.text.strip():
                result["paragraphs"].append(
                    {
                        "text": para.text,
                        "style": para.style.name if para.style else None,
                    }
                )

        # Извлечение таблиц
        table_index = 0
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                "index": table_idx,
                "rows": [],
                "is_data_table": False,  # Флаг таблицы с данными (не просто форматирование)
            }

            row_count = 0
            for row_idx, row in enumerate(table.rows):
                try:
                    row_data = []
                    # Безопасная обработка ячеек с проверкой на ошибки
                    for cell_idx, cell in enumerate(row.cells):
                        try:
                            text = cell.text.strip() if cell.text else ""
                            row_data.append(text)
                        except (AttributeError, IndexError, Exception) as cell_error:
                            logger.warning(
                                f"Ошибка обработки ячейки [{row_idx}, {cell_idx}] в таблице {table_idx}: {cell_error}. "
                                f"Пропускаю ячейку."
                            )
                            row_data.append("")
                    
                    # Пропускаем полностью пустые строки
                    if any(row_data):
                        table_data["rows"].append(row_data)
                        row_count += 1
                except (IndexError, AttributeError, Exception) as row_error:
                    logger.warning(
                        f"Ошибка обработки строки {row_idx} в таблице {table_idx}: {row_error}. "
                        f"Пропускаю строку и продолжаю обработку остальных."
                    )
                    continue

            # Если таблица содержит данные (более 1 строки, обычно минимум 2 - заголовок + данные)
            if row_count >= 2:
                table_data["is_data_table"] = True
                result["tables"].append(table_data)
                table_index += 1

        result["table_count"] = table_index

        # Если файл содержит таблицы с данными и похож на отчет, помечаем как готовый отчет
        if is_report_file and table_index > 0:
            result["is_ready_report"] = True
            logger.info(
                f"Обнаружен готовый Word отчет: {file_path} ({table_index} таблиц с данными)"
            )
        elif table_index > 0:
            logger.info(
                f"Word файл содержит {table_index} таблиц с данными: {file_path}"
            )

        logger.info(f"Успешно распарсен DOCX файл: {file_path}, таблиц: {table_index}")

    except Exception as e:
        logger.error(f"Ошибка при парсинге DOCX файла {file_path}: {e}")
        raise

    return result


def parse_pdf_file(file_path: str, batch_id: Optional[str] = None) -> Dict[str, Any]:
    """
    Улучшенный парсинг PDF файла с предварительной классификацией типа

    Новая архитектура:
    1. Предварительная классификация типа PDF (текст/изображение)
    2. Выбор стратегии обработки на основе классификации
    3. Оптимизированный пайплайн для каждого типа

    Args:
        file_path: Путь к файлу
        batch_id: Идентификатор загрузки для проверки отмены обработки

    Returns:
        Словарь с распознанными данными
    """
    result = {
        "file_path": file_path,
        "pages": [],
        "text": "",
        "tables": [],
        "metadata": {},
        "pdf_type": "unknown",
        "processing_strategy": "hybrid",
    }

    # ШАГ 1: Предварительная классификация типа PDF
    try:
        from utils.pdf_classifier import classify_pdf_type, get_pdf_processing_strategy

        pdf_classification = classify_pdf_type(file_path)
        strategy = get_pdf_processing_strategy(pdf_classification)

        result["pdf_type"] = pdf_classification.get("type", "unknown")
        result["pdf_type_confidence"] = pdf_classification.get("confidence", "low")
        result["processing_strategy"] = strategy
        result["text_ratio"] = pdf_classification.get("text_ratio", 0.0)

        logger.info(
            f"Классификация PDF: type={result['pdf_type']}, "
            f"strategy={strategy}, confidence={result['pdf_type_confidence']}"
        )

        # Если PDF явно на основе изображений, сразу переходим к OCR
        if strategy == "ocr_first" and result["pdf_type_confidence"] in [
            "high",
            "medium",
        ]:
            logger.info("PDF определен как image-based, применяю OCR-first стратегию")
            return _process_pdf_with_ocr_first(
                file_path, batch_id, result, pdf_classification
            )

    except ImportError:
        logger.warning(
            "Модуль pdf_classifier недоступен, используем стандартный пайплайн"
        )
    except Exception as e:
        logger.warning(
            f"Ошибка при классификации PDF: {e}, используем стандартный пайплайн"
        )

    # ШАГ 2: Стандартный пайплайн (text-first или hybrid)
    # Попытка использовать pdfplumber (лучше для таблиц)
    if HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(file_path) as pdf:
                result["metadata"] = {
                    "num_pages": len(pdf.pages),
                    "info": pdf.metadata or {},
                }

                full_text = []
                all_tables = []

                for page_num, page in enumerate(pdf.pages, 1):
                    # Проверяем, не была ли отменена обработка
                    if batch_id and HAS_PROGRESS_TRACKER:
                        tracker = get_progress_tracker(batch_id)
                        if tracker and tracker.is_cancelled():
                            raise InterruptedError(
                                "Обработка файла отменена пользователем"
                            )

                    # Извлечение текста
                    page_text = page.extract_text() or ""
                    full_text.append(page_text)

                    # Извлечение таблиц (базовый метод через pdfplumber)
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table:  # Проверяем, что таблица не пустая
                            all_tables.append(
                                {
                                    "page": page_num,
                                    "table_index": table_idx,
                                    "method": "pdfplumber",
                                    "rows": table,
                                    "row_count": len(table),
                                    "col_count": len(table[0]) if table else 0,
                                }
                            )

                    result["pages"].append(
                        {
                            "page_number": page_num,
                            "text": page_text,
                            "char_count": len(page_text),
                            "table_count": len(page_tables),
                        }
                    )

                # Дополнительное извлечение таблиц через улучшенный модуль
                try:
                    from utils.table_detector import extract_tables_from_pdf

                    enhanced_tables = extract_tables_from_pdf(
                        file_path, prefer_camelot=True
                    )
                    if enhanced_tables:
                        # Объединяем с уже найденными таблицами
                        # Используем улучшенные таблицы, если они найдены
                        if len(enhanced_tables) > len(all_tables):
                            logger.info(
                                f"Улучшенный модуль table_detector нашел {len(enhanced_tables)} таблиц (больше чем базовый метод)"
                            )
                            all_tables = enhanced_tables
                        else:
                            # Добавляем только те таблицы, которых нет в базовом результате
                            existing_pages = {
                                (t["page"], t["row_count"]) for t in all_tables
                            }
                            for table in enhanced_tables:
                                key = (table.get("page"), table.get("row_count", 0))
                                if key not in existing_pages:
                                    all_tables.append(table)
                except ImportError:
                    logger.debug(
                        "Модуль table_detector недоступен, используем только базовый метод"
                    )
                except Exception as e:
                    logger.warning(
                        f"Ошибка при использовании улучшенного извлечения таблиц: {e}"
                    )

                result["text"] = "\n\n".join(full_text)
                result["tables"] = all_tables
                result["total_characters"] = len(result["text"])
                result["total_tables"] = len(all_tables)

                # Улучшенное определение сканированного PDF
                # Используем несколько критериев:
                # 1. Среднее количество символов на страницу
                # 2. Количество страниц с текстом
                # 3. Соотношение страниц с текстом к общему количеству
                avg_chars_per_page = (
                    result["total_characters"] / len(pdf.pages) if pdf.pages else 0
                )
                pages_with_text = sum(
                    1 for p in result["pages"] if p.get("char_count", 0) > 0
                )
                text_coverage = (
                    (pages_with_text / len(pdf.pages) * 100) if pdf.pages else 0
                )

                # PDF считается сканированным если:
                # - Меньше 50 символов на страницу ИЛИ
                # - Меньше 30% страниц содержат текст
                is_likely_scanned = avg_chars_per_page < 50 or text_coverage < 30
                is_definitely_scanned = avg_chars_per_page < 10 or text_coverage < 10

                result["is_scanned"] = is_likely_scanned
                result["scanned_confidence"] = (
                    "high"
                    if is_definitely_scanned
                    else ("medium" if is_likely_scanned else "low")
                )
                result["text_coverage"] = text_coverage
                result["pages_with_text"] = pages_with_text

                if is_likely_scanned and HAS_OCR:
                    logger.info(
                        f"Обнаружено мало текста ({avg_chars_per_page:.0f} символов/страницу), применяю OCR..."
                    )
                    result["ocr_attempted"] = True  # Помечаем, что OCR был попытка
                    ocr_text = apply_ocr_to_pdf(file_path, batch_id=batch_id)
                    if ocr_text and len(ocr_text) > result["total_characters"]:
                        result["text"] = ocr_text
                        result["total_characters"] = len(ocr_text)
                        result["ocr_used"] = True
                        result["ocr_success"] = True
                        logger.info(f"OCR извлек {len(ocr_text)} символов текста")

                        # AI-ВАЛИДАЦИЯ: проверяем правдоподобность извлеченных данных
                        if HAS_AI_DATA_VALIDATOR:
                            try:
                                data_validator = get_ai_data_validator()
                                if data_validator:
                                    logger.info(
                                        "Применяю AI-валидацию извлеченных данных (pdfplumber)..."
                                    )
                                    validation_result = (
                                        data_validator.validate_extracted_data(
                                            extracted_data=ocr_text,
                                            document_type="energy_document",
                                        )
                                    )

                                    result["ai_validation"] = {
                                        "is_valid": validation_result.get(
                                            "is_valid", True
                                        ),
                                        "confidence": validation_result.get(
                                            "confidence", 0.5
                                        ),
                                        "issues": validation_result.get("issues", []),
                                        "warnings": validation_result.get(
                                            "warnings", []
                                        ),
                                        "suggestions": validation_result.get(
                                            "suggestions", []
                                        ),
                                        "ai_used": validation_result.get(
                                            "ai_used", False
                                        ),
                                    }

                                    logger.info(
                                        f"AI-валидация завершена: "
                                        f"валидно: {result['ai_validation']['is_valid']}, "
                                        f"проблем: {len(result['ai_validation']['issues'])}, "
                                        f"уверенность: {result['ai_validation']['confidence']:.2f}"
                                    )
                            except Exception as e:
                                logger.warning(f"Ошибка при AI-валидации данных: {e}")
                    else:
                        result["ocr_used"] = False
                        result["ocr_success"] = False
                        if avg_chars_per_page < 10:
                            result["ocr_error"] = "poppler_not_installed"
                            logger.warning(
                                "PDF похож на сканированный документ, но OCR недоступен или не извлек текст. "
                                "Установите poppler для обработки сканированных PDF."
                            )
                else:
                    result["ocr_used"] = False
                    result["ocr_attempted"] = False
                    result["is_scanned"] = False

                logger.info(
                    f"Успешно распарсен PDF файл через pdfplumber: {file_path}, страниц: {len(pdf.pages)}, таблиц: {len(all_tables)}, символов: {result['total_characters']}"
                )
                return result

        except Exception as e:
            logger.warning(
                f"Ошибка при парсинге PDF через pdfplumber: {e}, пробую PyPDF2"
            )

    # Fallback на PyPDF2
    if HAS_PYPDF2:
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                result["metadata"] = {
                    "num_pages": len(pdf_reader.pages),
                    "info": pdf_reader.metadata or {},
                }

                full_text = []
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    # Проверяем, не была ли отменена обработка
                    if batch_id and HAS_PROGRESS_TRACKER:
                        tracker = get_progress_tracker(batch_id)
                        if tracker and tracker.is_cancelled():
                            raise InterruptedError(
                                "Обработка файла отменена пользователем"
                            )

                    page_text = page.extract_text() or ""
                    full_text.append(page_text)
                    result["pages"].append(
                        {
                            "page_number": page_num,
                            "text": page_text,
                            "char_count": len(page_text),
                        }
                    )

                result["text"] = "\n\n".join(full_text)
                result["total_characters"] = len(result["text"])
                result["total_tables"] = 0  # PyPDF2 не извлекает таблицы

                # Улучшенное определение сканированного PDF (PyPDF2 fallback)
                avg_chars_per_page = (
                    result["total_characters"] / len(pdf_reader.pages)
                    if pdf_reader.pages
                    else 0
                )
                pages_with_text = sum(
                    1 for p in result["pages"] if p.get("char_count", 0) > 0
                )
                text_coverage = (
                    (pages_with_text / len(pdf_reader.pages) * 100)
                    if pdf_reader.pages
                    else 0
                )

                is_likely_scanned = avg_chars_per_page < 50 or text_coverage < 30
                is_definitely_scanned = avg_chars_per_page < 10 or text_coverage < 10

                result["is_scanned"] = is_likely_scanned
                result["scanned_confidence"] = (
                    "high"
                    if is_definitely_scanned
                    else ("medium" if is_likely_scanned else "low")
                )
                result["text_coverage"] = text_coverage
                result["pages_with_text"] = pages_with_text

                if is_likely_scanned and HAS_OCR:
                    logger.info(
                        f"Обнаружено мало текста ({avg_chars_per_page:.0f} символов/страницу), применяю OCR..."
                    )
                    result["ocr_attempted"] = True  # Помечаем, что OCR был попытка
                    ocr_text = apply_ocr_to_pdf(file_path, batch_id=batch_id)
                    if ocr_text and len(ocr_text) > result["total_characters"]:
                        result["text"] = ocr_text
                        result["total_characters"] = len(ocr_text)
                        result["ocr_used"] = True
                        result["ocr_success"] = True
                        logger.info(f"OCR извлек {len(ocr_text)} символов текста")

                        # AI-ВАЛИДАЦИЯ: проверяем правдоподобность извлеченных данных
                        if HAS_AI_DATA_VALIDATOR:
                            try:
                                data_validator = get_ai_data_validator()
                                if data_validator:
                                    logger.info(
                                        "Применяю AI-валидацию извлеченных данных..."
                                    )
                                    validation_result = (
                                        data_validator.validate_extracted_data(
                                            extracted_data=ocr_text,
                                            document_type="energy_document",
                                        )
                                    )

                                    result["ai_validation"] = {
                                        "is_valid": validation_result.get(
                                            "is_valid", True
                                        ),
                                        "confidence": validation_result.get(
                                            "confidence", 0.5
                                        ),
                                        "issues": validation_result.get("issues", []),
                                        "warnings": validation_result.get(
                                            "warnings", []
                                        ),
                                        "suggestions": validation_result.get(
                                            "suggestions", []
                                        ),
                                        "ai_used": validation_result.get(
                                            "ai_used", False
                                        ),
                                    }

                                    logger.info(
                                        f"AI-валидация завершена: "
                                        f"валидно: {result['ai_validation']['is_valid']}, "
                                        f"проблем: {len(result['ai_validation']['issues'])}, "
                                        f"уверенность: {result['ai_validation']['confidence']:.2f}"
                                    )
                            except Exception as e:
                                logger.warning(f"Ошибка при AI-валидации данных: {e}")

                        # Пытаемся извлечь таблицы из OCR текста
                        try:
                            from utils.ocr_table_extractor import (
                                extract_tables_from_ocr_text,
                            )

                            ocr_tables = extract_tables_from_ocr_text(
                                ocr_text, page_num=1
                            )
                            if ocr_tables:
                                result["tables"].extend(ocr_tables)
                                result["total_tables"] = len(result["tables"])
                                logger.info(
                                    f"Добавлено {len(ocr_tables)} таблиц из OCR текста"
                                )

                                # AI-СТРУКТУРИРОВАНИЕ ТАБЛИЦ: улучшаем структуру таблиц через AI
                                if HAS_AI_TABLE_PARSER and ocr_tables:
                                    try:
                                        table_parser = get_ai_table_parser()
                                        if table_parser:
                                            logger.info(
                                                "Применяю AI-структурирование таблиц..."
                                            )
                                            structured_tables = (
                                                table_parser.structure_multiple_tables(
                                                    ocr_tables
                                                )
                                            )
                                            if structured_tables:
                                                result["tables"] = structured_tables
                                                result["ai_tables_structured"] = True
                                                logger.info(
                                                    f"AI-структурировано {len(structured_tables)} таблиц"
                                                )
                                    except Exception as e:
                                        logger.warning(
                                            f"Ошибка при AI-структурировании таблиц: {e}"
                                        )
                        except ImportError:
                            logger.debug("Модуль ocr_table_extractor недоступен")
                        except Exception as e:
                            logger.warning(f"Ошибка при извлечении таблиц из OCR: {e}")
                    else:
                        result["ocr_used"] = False
                        result["ocr_success"] = False
                        if avg_chars_per_page < 10:
                            result["ocr_error"] = "poppler_not_installed"
                            logger.warning(
                                "PDF похож на сканированный документ, но OCR недоступен или не извлек текст. "
                                "Установите poppler для обработки сканированных PDF."
                            )
                else:
                    result["ocr_used"] = False
                    result["ocr_attempted"] = False
                    result["is_scanned"] = False

                logger.info(
                    f"Успешно распарсен PDF файл через PyPDF2: {file_path}, страниц: {len(pdf_reader.pages)}, символов: {result['total_characters']}"
                )
                return result

        except Exception as e:
            logger.error(f"Ошибка при парсинге PDF через PyPDF2: {e}")
            raise

    raise ImportError(
        "Не установлены библиотеки для парсинга PDF. Установите: pip install PyPDF2 pdfplumber"
    )


def _process_pdf_with_ocr_first(
    file_path: str,
    batch_id: Optional[str],
    result: Dict[str, Any],
    pdf_classification: Dict[str, Any],
) -> Dict[str, Any]:
    """
    Обработка PDF с OCR-first стратегией для сканированных документов

    Args:
        file_path: Путь к PDF файлу
        batch_id: Идентификатор загрузки
        result: Базовый результат
        pdf_classification: Результат классификации

    Returns:
        Результат обработки
    """
    logger.info("Применяю OCR-first стратегию для image-based PDF")

    # Сразу применяем OCR
    if HAS_OCR:
        ocr_text = apply_ocr_to_pdf(file_path, batch_id=batch_id)
        if ocr_text:
            result["text"] = ocr_text
            result["total_characters"] = len(ocr_text)
            result["ocr_used"] = True
            result["ocr_success"] = True
            result["ocr_attempted"] = True

            # AI-ВАЛИДАЦИЯ: проверяем правдоподобность извлеченных данных
            if HAS_AI_DATA_VALIDATOR:
                try:
                    data_validator = get_ai_data_validator()
                    if data_validator:
                        logger.info(
                            "Применяю AI-валидацию извлеченных данных (OCR-first)..."
                        )
                        validation_result = data_validator.validate_extracted_data(
                            extracted_data=ocr_text, document_type="energy_document"
                        )

                        result["ai_validation"] = {
                            "is_valid": validation_result.get("is_valid", True),
                            "confidence": validation_result.get("confidence", 0.5),
                            "issues": validation_result.get("issues", []),
                            "warnings": validation_result.get("warnings", []),
                            "suggestions": validation_result.get("suggestions", []),
                            "ai_used": validation_result.get("ai_used", False),
                        }

                        logger.info(
                            f"AI-валидация завершена: "
                            f"валидно: {result['ai_validation']['is_valid']}, "
                            f"проблем: {len(result['ai_validation']['issues'])}, "
                            f"уверенность: {result['ai_validation']['confidence']:.2f}"
                        )
                except Exception as e:
                    logger.warning(f"Ошибка при AI-валидации данных: {e}")

            # Извлекаем таблицы из OCR текста
            try:
                from utils.ocr_table_extractor import (
                    extract_tables_from_ocr_text,
                    structure_ocr_data,
                )

                # Пытаемся извлечь таблицы из каждой страницы
                # Разбиваем OCR текст по страницам (маркеры "--- Страница X ---")
                ocr_tables = []
                pages = re.split(r"---\s*Страница\s+(\d+)\s+---", ocr_text)
                
                # Если текст разбит по страницам, обрабатываем каждую отдельно
                if len(pages) > 1:
                    for i in range(1, len(pages), 2):  # Нечетные индексы - номера страниц
                        if i + 1 < len(pages):
                            page_num = int(pages[i])
                            page_text = pages[i + 1]
                            page_tables = extract_tables_from_ocr_text(page_text, page_num=page_num)
                            ocr_tables.extend(page_tables)
                else:
                    # Если нет маркеров страниц, обрабатываем весь текст
                    ocr_tables = extract_tables_from_ocr_text(ocr_text, page_num=1)
                if ocr_tables:
                    result["tables"] = ocr_tables
                    result["total_tables"] = len(ocr_tables)
                    logger.info(f"Извлечено {len(ocr_tables)} таблиц из OCR текста")

                    # AI-СТРУКТУРИРОВАНИЕ ТАБЛИЦ: улучшаем структуру таблиц через AI
                    if HAS_AI_TABLE_PARSER and ocr_tables:
                        try:
                            table_parser = get_ai_table_parser()
                            if table_parser:
                                logger.info(
                                    "Применяю AI-структурирование таблиц (OCR-first)..."
                                )
                                structured_tables = (
                                    table_parser.structure_multiple_tables(ocr_tables)
                                )
                                if structured_tables:
                                    result["tables"] = structured_tables
                                    result["ai_tables_structured"] = True
                                    logger.info(
                                        f"AI-структурировано {len(structured_tables)} таблиц"
                                    )
                        except Exception as e:
                            logger.warning(
                                f"Ошибка при AI-структурировании таблиц: {e}"
                            )

                # Структурируем данные
                structured = structure_ocr_data(ocr_text, ocr_tables)
                result["structured_data"] = structured.get("structured_sections", [])

            except ImportError:
                logger.debug(
                    "Модуль ocr_table_extractor недоступен, пропускаем извлечение таблиц"
                )
            except Exception as e:
                logger.warning(f"Ошибка при извлечении таблиц из OCR: {e}")

            # Получаем метаданные
            if HAS_PYPDF2:
                try:
                    with open(file_path, "rb") as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        result["metadata"] = {
                            "num_pages": len(pdf_reader.pages),
                            "info": pdf_reader.metadata or {},
                        }
                        result["pages"] = [
                            {
                                "page_number": i + 1,
                                "text": "",
                                "char_count": 0,
                                "ocr_used": True,
                            }
                            for i in range(len(pdf_reader.pages))
                        ]
                except Exception as e:
                    logger.warning(f"Не удалось получить метаданные PDF: {e}")

            logger.info(
                f"OCR-first обработка завершена: {len(ocr_text)} символов, {len(result.get('tables', []))} таблиц"
            )
            return result
        else:
            result["ocr_used"] = False
            result["ocr_success"] = False
            result["ocr_attempted"] = True
            result["ocr_error"] = "ocr_failed"
            logger.warning("OCR не извлек текст из image-based PDF")
    else:
        result["ocr_used"] = False
        result["ocr_attempted"] = False
        result["ocr_error"] = "ocr_not_available"
        logger.warning("OCR недоступен для image-based PDF")

    return result


def apply_ocr_to_pdf(
    file_path: str, languages: str = "rus+eng", batch_id: Optional[str] = None
) -> str:
    """
    Применяет OCR (оптическое распознавание символов) к PDF файлу
    Используется для сканированных документов

    Args:
        file_path: Путь к PDF файлу
        languages: Языки для распознавания (по умолчанию русский + английский)
        batch_id: Идентификатор загрузки для проверки отмены обработки

    Returns:
        Распознанный текст
    """
    if not HAS_OCR:
        logger.warning("OCR библиотеки не установлены")
        return ""

    try:
        logger.info(f"Начинаю OCR обработку PDF: {file_path}")

        # Конвертируем PDF в изображения
        try:
            # Проверяем наличие Poppler в стандартных путях
            poppler_paths = [
                r"C:\poppler\Library\bin",  # Conda установка
                r"C:\poppler\bin",  # Стандартная установка
            ]

            poppler_path = None
            for path in poppler_paths:
                if os.path.exists(path) and os.path.exists(
                    os.path.join(path, "pdftoppm.exe")
                ):
                    poppler_path = path
                    break

            if poppler_path:
                # Добавляем путь к Poppler в PATH для текущего процесса
                current_path = os.environ.get("PATH", "")
                if poppler_path not in current_path:
                    os.environ["PATH"] = poppler_path + os.pathsep + current_path
                    logger.debug(f"Добавлен Poppler в PATH: {poppler_path}")

            # Используем DPI=200 для ускорения (достаточно для OCR, быстрее обработка)
            images = convert_from_path(file_path, dpi=200, poppler_path=poppler_path)
            logger.info(f"PDF конвертирован в {len(images)} изображений")
        except Exception as e:
            error_msg = str(e).lower()
            if "poppler" in error_msg or "page count" in error_msg:
                logger.warning(
                    "Poppler не установлен или не найден в PATH. OCR для PDF недоступен.\n"
                    "Для установки на Windows:\n"
                    "  1. Скачайте poppler: https://github.com/oschwartz10612/poppler-windows/releases/\n"
                    "  2. Распакуйте и добавьте папку 'bin' в PATH\n"
                    "  3. Или установите через conda: conda install -c conda-forge poppler\n"
                    "  4. Или через chocolatey: choco install poppler\n"
                    "Файл будет обработан без OCR."
                )
            else:
                logger.error(f"Ошибка конвертации PDF в изображения: {e}")
            return ""

        # Применяем OCR к каждому изображению
        ocr_results = []
        tesseract_available = True
        max_pages_for_ocr = 10  # Ограничиваем количество страниц для OCR

        for i, image in enumerate(images[:max_pages_for_ocr], 1):
            # Проверяем, не была ли отменена обработка перед обработкой каждой страницы
            if batch_id and HAS_PROGRESS_TRACKER:
                tracker = get_progress_tracker(batch_id)
                if tracker and tracker.is_cancelled():
                    logger.info(f"OCR обработка отменена пользователем на странице {i}")
                    raise InterruptedError("Обработка файла отменена пользователем")

            try:
                # Улучшаем изображение перед OCR (использует улучшенный модуль если доступен)
                try:
                    enhanced_image = preprocess_image_for_ocr(image)
                except Exception as e:
                    logger.warning(
                        f"Ошибка улучшения изображения страницы {i}: {e}, использую оригинал"
                    )
                    enhanced_image = image

                # Автоопределение поворота на 90°/180°/270° через Tesseract OSD
                try:
                    osd = pytesseract.image_to_osd(enhanced_image)
                    angle_match = re.search(r'(?<=Rotate: )\d+', osd)
                    if angle_match:
                        angle = int(angle_match.group(0))
                        if angle != 0:
                            enhanced_image = enhanced_image.rotate(-angle, expand=True)
                            logger.debug(f"Страница {i} повернута на {angle}° (определено через OSD)")
                except Exception as e:
                    logger.debug(f"OSD не смог определить поворот страницы {i}: {e}")

                # Исправление небольшого наклона (deskew)
                try:
                    from utils.image_enhancement import deskew_image
                    enhanced_image = deskew_image(enhanced_image)
                    logger.debug(f"Применено автоисправление наклона для страницы {i}")
                except ImportError:
                    logger.debug("Модуль deskew_image недоступен, пропускаем исправление наклона")
                except Exception as e:
                    logger.warning(f"Ошибка исправления наклона страницы {i}: {e}")

                # Используем pytesseract для распознавания текста
                text = pytesseract.image_to_string(enhanced_image, lang=languages)
                if text.strip():
                    ocr_results.append(f"--- Страница {i} ---\n{text}")
                    logger.info(f"OCR страница {i}: извлечено {len(text)} символов")
            except pytesseract.TesseractNotFoundError:
                if tesseract_available:  # Выводим сообщение только один раз
                    logger.error(
                        "Tesseract OCR не установлен или не найден в PATH.\n"
                        "Установите Tesseract для работы OCR:\n"
                        "  Windows: https://github.com/UB-Mannheim/tesseract/wiki\n"
                        "  Или через chocolatey: choco install tesseract\n"
                        "  После установки перезапустите сервер/скрипт"
                    )
                    tesseract_available = False
                break  # Прерываем цикл, так как Tesseract недоступен
            except Exception as e:
                error_msg = str(e).lower()
                if "tesseract" in error_msg and (
                    "not installed" in error_msg or "not in your path" in error_msg
                ):
                    if tesseract_available:  # Выводим сообщение только один раз
                        logger.error(
                            f"Tesseract OCR не установлен или не найден в PATH на странице {i}.\n"
                            f"Установите Tesseract: https://github.com/UB-Mannheim/tesseract/wiki\n"
                            f"См. INSTALL_TESSERACT.md для подробных инструкций"
                        )
                        tesseract_available = False
                    break  # Прерываем цикл
                else:
                    logger.error(f"Ошибка OCR на странице {i}: {e}")
                    continue

        full_text = "\n\n".join(ocr_results)
        if full_text:
            logger.info(
                f"OCR завершен: всего извлечено {len(full_text)} символов из {len(images)} страниц"
            )

            # AI-УСИЛЕНИЕ OCR: улучшаем качество распознавания через AI
            if HAS_OCR_AI_ENHANCER:
                try:
                    ocr_enhancer = get_ocr_ai_enhancer()
                    if ocr_enhancer:
                        logger.info(
                            "Применяю AI-усиление OCR для улучшения качества распознавания..."
                        )
                        enhancement_result = ocr_enhancer.enhance_ocr_accuracy(
                            initial_ocr_text=full_text,
                            context="Энергетический документ с данными о потреблении энергии",
                        )

                        if enhancement_result.get("ai_used") and enhancement_result.get(
                            "enhanced_text"
                        ):
                            enhanced_text = enhancement_result.get(
                                "enhanced_text", full_text
                            )
                            improvements_count = len(
                                enhancement_result.get("improvements", [])
                            )
                            confidence = enhancement_result.get("confidence_score", 0.0)

                            logger.info(
                                f"AI-усиление OCR завершено: "
                                f"улучшений: {improvements_count}, "
                                f"уверенность: {confidence:.2f}, "
                                f"символов: {len(enhanced_text)}"
                            )

                            # Используем улучшенный текст
                            full_text = enhanced_text
                        else:
                            logger.debug(
                                "AI-усиление OCR недоступно или не применилось, используем оригинальный OCR текст"
                            )
                except Exception as e:
                    logger.warning(
                        f"Ошибка при AI-усилении OCR: {e}, используем оригинальный OCR текст"
                    )
        else:
            logger.warning("OCR не извлек текст из PDF")
        return full_text

    except Exception as e:
        error_msg = str(e).lower()
        if "poppler" in error_msg or "page count" in error_msg:
            logger.warning(
                "Poppler не установлен. OCR для PDF недоступен.\n"
                "Инструкции по установке см. в логах выше."
            )
        else:
            logger.error(f"Ошибка при применении OCR к PDF {file_path}: {e}")
        return ""


def preprocess_image_for_ocr(image, dpi: int = 200):
    """
    Комплексная предобработка изображения для улучшения качества OCR
    
    Порядок обработки:
    1. Нормализация освещения (если доступен OpenCV)
    2. Удаление шумов
    3. Улучшение контраста
    4. Адаптивная бинаризация (если доступен OpenCV)
    5. Улучшение резкости
    
    Использует улучшенный модуль image_enhancement если доступен

    Args:
        image: Объект PIL Image
        dpi: Разрешение изображения (для оптимизации, не используется напрямую)

    Returns:
        Обработанное изображение
    """
    try:
        # Пробуем использовать улучшенный модуль (комплексная обработка)
        from utils.image_enhancement import enhance_image_for_ocr

        logger.debug("Используется улучшенный модуль image_enhancement (комплексная обработка)")
        enhanced = enhance_image_for_ocr(image)
        return enhanced
    except ImportError:
        # Fallback на базовую предобработку через PIL
        logger.debug("Используется базовая предобработка изображения (PIL only)")
        
        # Конвертируем в RGB если нужно
        if image.mode != "RGB":
            image = image.convert("RGB")

        # Улучшение контраста (более агрессивное для плохих сканов)
        try:
            if PIL_ImageEnhance:
                enhancer = PIL_ImageEnhance.Contrast(image)
                image = enhancer.enhance(1.5)
        except Exception as e:
            logger.debug(f"Ошибка улучшения контраста: {e}")

        # Улучшение резкости
        try:
            if PIL_ImageEnhance:
                enhancer = PIL_ImageEnhance.Sharpness(image)
                image = enhancer.enhance(1.3)  # Увеличено с 1.2 до 1.3
        except Exception as e:
            logger.debug(f"Ошибка улучшения резкости: {e}")

        logger.debug("Изображение предобработано: контраст +50%, резкость +30%")

        return image
    except Exception as e:
        logger.warning(f"Ошибка предобработки изображения: {e}, возвращаю оригинал")
        return image


def apply_ocr_to_image(image_path: str, languages: str = "rus+eng") -> Dict[str, Any]:
    """
    Применяет OCR (оптическое распознавание символов) к изображению
    Используется для JPG, PNG и других форматов изображений

    Args:
        image_path: Путь к файлу изображения
        languages: Языки для распознавания (по умолчанию русский + английский)

    Returns:
        Словарь с распознанным текстом и метаданными
    """
    if not HAS_OCR:
        logger.warning("OCR библиотеки не установлены")
        return {
            "text": "",
            "char_count": 0,
            "ocr_used": False,
            "error": "OCR библиотеки не установлены",
        }

    try:
        logger.info(f"Начинаю OCR обработку изображения: {image_path}")

        # Открываем изображение
        if not PIL_Image:
            raise ImportError("PIL Image не доступен")
        image = PIL_Image.open(image_path)
        logger.info(
            f"Изображение загружено: {image.size[0]}x{image.size[1]} пикселей, режим: {image.mode}"
        )

        # Предобработка изображения для лучшего распознавания
        image = preprocess_image_for_ocr(image)

        # Применяем OCR
        text = pytesseract.image_to_string(image, lang=languages)

        char_count = len(text)
        logger.info(f"OCR завершен: извлечено {char_count} символов")

        return {
            "text": text,
            "char_count": char_count,
            "ocr_used": True,
            "image_size": image.size,
            "image_mode": image.mode,
        }

    except Exception as e:
        logger.error(f"Ошибка при применении OCR к изображению {image_path}: {e}")
        return {"text": "", "char_count": 0, "ocr_used": False, "error": str(e)}


def detect_file_type(file_path: str) -> str:
    """
    Определение типа файла по расширению

    Returns:
        Тип файла: 'excel', 'docx', 'pdf', 'image', 'unknown'
    """
    ext = Path(file_path).suffix.lower()

    if ext in [".xlsx", ".xlsm", ".xls"]:
        return "excel"
    elif ext == ".docx":
        return "docx"
    elif ext == ".pdf":
        return "pdf"
    elif ext in [".jpg", ".jpeg", ".png"]:
        return "image"
    else:
        return "unknown"


def parse_file(
    file_path: str, template_name: Optional[str] = None, batch_id: Optional[str] = None
) -> Dict[str, Any]:
    """
    Универсальная функция парсинга файла согласно ТЗ (раздел 4.1)

    Args:
        file_path: Путь к файлу
        template_name: Имя эталонного шаблона (для Excel)
        batch_id: Идентификатор загрузки для проверки отмены обработки

    Returns:
        Словарь с распознанными данными
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    file_type = detect_file_type(file_path)

    result = {
        "file_path": file_path,
        "file_type": file_type,
        "parsed": False,
        "data": {},
    }

    try:
        if file_type == "excel":
            parsed_data = parse_excel_file(file_path, template_name)
            result["data"] = parsed_data
            result["parsed"] = True

        elif file_type == "docx":
            parsed_data = parse_docx_file(file_path)
            result["data"] = parsed_data
            result["parsed"] = True

        elif file_type == "pdf":
            # Проверяем, нужно ли использовать ИИ вместо традиционного парсинга
            use_ai = os.getenv("AI_ENABLED", "false").lower() == "true"
            prefer_ai = os.getenv("AI_PREFER_FOR_PDF", "false").lower() == "true"

            if use_ai and HAS_AI_PARSER and prefer_ai:
                # Используем ИИ для распознавания
                try:
                    ai_parser = get_ai_parser()
                    if ai_parser:
                        logger.info("Используется ИИ для распознавания PDF")
                        parsed_data = ai_parser.recognize_document(file_path, "pdf")
                        result["data"] = parsed_data
                        result["parsed"] = True
                        result["ai_used"] = True
                    else:
                        # Fallback на традиционный парсинг
                        parsed_data = parse_pdf_file(file_path, batch_id=batch_id)
                        result["data"] = parsed_data
                        result["parsed"] = True
                        result["ai_used"] = False
                except Exception as e:
                    logger.warning(
                        f"Ошибка ИИ распознавания, используем традиционный парсинг: {e}"
                    )
                    parsed_data = parse_pdf_file(file_path, batch_id=batch_id)
                    result["data"] = parsed_data
                    result["parsed"] = True
                    result["ai_used"] = False
            else:
                # Традиционный парсинг
                parsed_data = parse_pdf_file(file_path, batch_id=batch_id)
                result["data"] = parsed_data
                result["parsed"] = True
                result["ai_used"] = False

        elif file_type == "image":
            # Проверяем, нужно ли использовать ИИ для изображений
            use_ai = os.getenv("AI_ENABLED", "false").lower() == "true"

            if use_ai and HAS_AI_PARSER:
                try:
                    ai_parser = get_ai_parser()
                    if ai_parser:
                        logger.info("Используется ИИ для распознавания изображения")
                        parsed_data = ai_parser.recognize_document(file_path, "image")
                        result["data"] = parsed_data
                        result["parsed"] = True
                        result["ai_used"] = True
                    else:
                        # Fallback на OCR
                        logger.info("ИИ не настроен, используем OCR для изображения")
                        ocr_data = apply_ocr_to_image(file_path)
                        result["data"] = {
                            "text": ocr_data.get("text", ""),
                            "char_count": ocr_data.get("char_count", 0),
                            "ocr_used": ocr_data.get("ocr_used", False),
                            "image_size": ocr_data.get("image_size"),
                            "image_mode": ocr_data.get("image_mode"),
                        }
                        result["parsed"] = ocr_data.get("ocr_used", False)
                        result["ai_used"] = False
                except Exception as e:
                    logger.warning(
                        f"Ошибка ИИ распознавания изображения: {e}, используем OCR"
                    )
                    # Fallback на OCR
                    ocr_data = apply_ocr_to_image(file_path)
                    result["data"] = {
                        "text": ocr_data.get("text", ""),
                        "char_count": ocr_data.get("char_count", 0),
                        "ocr_used": ocr_data.get("ocr_used", False),
                        "error": ocr_data.get("error"),
                    }
                    result["parsed"] = ocr_data.get("ocr_used", False)
                    result["ai_used"] = False
            else:
                # Используем OCR напрямую (с улучшенной предобработкой)
                logger.info("Используется OCR для распознавания изображения")
                ocr_data = apply_ocr_to_image(file_path)
                result["data"] = {
                    "text": ocr_data.get("text", ""),
                    "char_count": ocr_data.get("char_count", 0),
                    "ocr_used": ocr_data.get("ocr_used", False),
                    "image_size": ocr_data.get("image_size"),
                    "image_mode": ocr_data.get("image_mode"),
                }
                result["parsed"] = ocr_data.get("ocr_used", False)
                result["ai_used"] = False

        else:
            result["data"] = {"error": f"Неподдерживаемый тип файла: {file_type}"}
            result["parsed"] = False

        # ПОЛНЫЙ AI-АНАЛИЗ ЭНЕРГЕТИЧЕСКИХ ДАННЫХ (после успешного парсинга)
        if result.get("parsed") and HAS_AI_ENERGY_ANALYSIS:
            try:
                # Проверяем, является ли это энергетическим документом
                data = result.get("data", {})
                is_energy_document = (
                    "electricity" in data
                    or "gas" in data
                    or "water" in data
                    or "fuel" in data
                    or "energy_resources" in data
                    or "equipment" in data
                )

                if is_energy_document:
                    logger.info(
                        "Обнаружен энергетический документ, запускаю полный AI-анализ..."
                    )

                    # Запускаем полный AI-анализ
                    ai_analysis = enhanced_energy_analysis(data)

                    # Генерируем отчет о качестве данных
                    if HAS_AI_QUALITY_REPORTER:
                        try:
                            quality_reporter = get_quality_reporter()
                            quality_report = quality_reporter.generate_quality_report(
                                ai_analysis, result
                            )
                            result["quality_report"] = quality_report
                        except Exception as e:
                            logger.warning(
                                f"Ошибка при генерации отчета о качестве: {e}"
                            )

                    # Добавляем результаты AI-анализа в результат парсинга
                    result["ai_analysis"] = ai_analysis
                    result["data_quality_score"] = ai_analysis.get(
                        "confidence_score", 0.0
                    )
                    result["is_compliant"] = ai_analysis.get("summary", {}).get(
                        "is_compliant", True
                    )

                    # Логируем результаты
                    logger.info(
                        f"Полный AI-анализ завершен: "
                        f"валидно: {ai_analysis.get('summary', {}).get('is_valid', True)}, "
                        f"аномалий: {ai_analysis.get('anomalies', {}).get('anomaly_count', 0)}, "
                        f"соответствует ПКМ №690: {ai_analysis.get('summary', {}).get('is_compliant', True)}, "
                        f"уверенность: {ai_analysis.get('confidence_score', 0.0):.2f}"
                    )

                    # Предупреждение при низком качестве данных
                    if ai_analysis.get("confidence_score", 1.0) < 0.7:
                        logger.warning(
                            f"Низкое качество данных (confidence_score: {ai_analysis.get('confidence_score', 0.0):.2f}). "
                            f"Рекомендуется ручная проверка."
                        )

            except Exception as e:
                logger.warning(
                    f"Ошибка при полном AI-анализе энергетических данных: {e}"
                )
                # Не прерываем обработку, просто логируем ошибку

    except Exception as e:
        logger.error(f"Ошибка при парсинге файла {file_path}: {e}")
        result["data"] = {"error": str(e)}
        result["parsed"] = False

    return result
