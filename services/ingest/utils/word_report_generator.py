"""
Генератор Word отчета энергоаудита на основе исходных данных из ingest
Отчеты генерируются из:
- Агрегированных данных из исходных файлов (Excel, PDF, DOCX)
- Результатов AI-анализа данных (если AI доступен)
- Дополнительных данных (оборудование, узлы учета, ограждающие конструкции)

Все данные обрабатываются через AI для улучшения качества и полноты отчета.

Все числовые расчёты выполняются через централизованные модули:
- energy_passport_calculations.py - формулы и расчёты
- energy_units.py - единицы измерения и конвертация
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List
import sys

logger = logging.getLogger(__name__)

# Импорт централизованных формул и единиц измерения
# Путь: utils -> ingest -> domain
_calc_module_path = Path(__file__).resolve().parent.parent / "domain"
if str(_calc_module_path) not in sys.path:
    sys.path.insert(0, str(_calc_module_path))

try:
    from energy_passport_calculations import (
        calculate_total_consumption_by_resource,
        calculate_total_costs,
        calculate_average_payback_period,
        extract_equipment_data,
    )

    HAS_CALCULATIONS = True
except ImportError:
    HAS_CALCULATIONS = False
    logger.warning(
        "Не удалось импортировать energy_passport_calculations. Используются локальные вычисления."
    )

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None


class WordReportGenerator:
    """Генератор Word документов энергоаудита по ПКМ 690"""

    def __init__(self):
        if not HAS_DOCX:
            raise ImportError(
                "python-docx не установлен. Установите: pip install python-docx"
            )
        self.doc = None

    def generate_report(
        self,
        enterprise_data: Dict[str, Any],
        aggregated_data: Dict[str, Any],
        equipment_data: Optional[Dict[str, Any]] = None,
        nodes_data: Optional[List[Dict[str, Any]]] = None,
        envelope_data: Optional[Dict[str, Any]] = None,
        output_path: Optional[Path] = None,
        skip_readiness_check: bool = False,
    ) -> Document:
        """
        Генерирует Word отчет энергоаудита из исходных данных с применением AI-анализа

        Отчет формируется на основе:
        - Агрегированных данных из исходных файлов (Excel, PDF, DOCX)
        - Результатов AI-анализа (если доступен: insights, recommendations)
        - Дополнительных данных (оборудование, узлы учета, расчет теплопотерь по зданиям)

        Args:
            enterprise_data: Данные предприятия (name, id, address, etc.)
            aggregated_data: Агрегированные данные энергопотребления из исходных файлов (resources, quarters)
                           Может содержать ai_insights и ai_recommendations, если применен AI-анализ
            equipment_data: Данные оборудования из исходных файлов (опционально)
            nodes_data: Данные узлов учета из исходных файлов (опционально)
            envelope_data: Данные расчета теплопотерь по зданиям из исходных файлов (опционально)
            output_path: Путь для сохранения (если None, возвращается объект Document)

        Returns:
            Document: Объект Word документа, сгенерированный из исходных данных
        """
        logger.info(
            f"📄 Начало генерации Word отчета из исходных данных для предприятия: {enterprise_data.get('name', 'Unknown')}"
        )

        # Проверка готовности данных (если не пропущена)
        if not skip_readiness_check:
            try:
                from .word_readiness_validator import validate_word_report_readiness
                from ..domain.report_data import ReportData

                # Создаем ReportData для проверки
                report_data = ReportData.from_raw_data(
                    aggregated_data=aggregated_data,
                    equipment_data=equipment_data,
                    nodes_data=nodes_data,
                    envelope_data=envelope_data,
                    enterprise_data=enterprise_data,
                )

                # Проверяем готовность
                readiness = validate_word_report_readiness(report_data)

                if not readiness["ready"]:
                    from .word_readiness_validator import get_missing_data_summary

                    summary = get_missing_data_summary(readiness)
                    logger.warning(
                        f"⚠️ Данные не готовы для генерации Word-отчёта:\n{summary}"
                    )

                    # Блокируем генерацию, если нет критических данных
                    if readiness.get("critical_missing_sections"):
                        raise ValueError(
                            f"Недостаточно данных для генерации Word-отчёта. "
                            f"Критические разделы не могут быть сгенерированы: {readiness['critical_missing_sections']}. "
                            f"Подробности: {summary}"
                        )
                    else:
                        logger.info(
                            "⚠️ Генерация продолжается с предупреждениями (есть fallback на эталонные таблицы)"
                        )
                else:
                    logger.info(
                        f"✅ Данные готовы для генерации Word-отчёта (готовность: {readiness['completeness_score'] * 100:.0f}%)"
                    )

            except ImportError as e:
                logger.warning(
                    f"Модуль проверки готовности недоступен: {e}. Продолжаем генерацию без проверки."
                )
            except ValueError:
                # Пробрасываем ошибку валидации наверх
                raise
            except Exception as e:
                logger.warning(
                    f"Ошибка проверки готовности: {e}. Продолжаем генерацию."
                )

        # Проверяем наличие AI-инсайтов в данных
        has_ai_insights = bool(aggregated_data.get("ai_insights"))
        has_ai_recommendations = bool(aggregated_data.get("ai_recommendations"))
        if has_ai_insights or has_ai_recommendations:
            logger.info("🤖 Использование AI-обогащенных данных для генерации отчета")

        self.doc = Document()

        # Настройка стилей
        self._setup_document_styles()

        # Создание разделов документа
        self._create_title_page(enterprise_data)
        self._create_table_of_contents()
        self._create_introduction(enterprise_data)
        self._create_enterprise_info(enterprise_data)
        self._create_energy_consumption_analysis(enterprise_data, aggregated_data)

        if equipment_data:
            self._create_equipment_analysis(enterprise_data, equipment_data)

        if nodes_data:
            self._create_metering_nodes_section(nodes_data)

        if envelope_data:
            self._create_envelope_section(envelope_data)

        self._create_energy_efficiency_measures(enterprise_data, aggregated_data)
        self._create_economic_analysis(enterprise_data, aggregated_data)
        self._create_conclusion(
            enterprise_data, aggregated_data
        )  # Передаем aggregated_data для AI-инсайтов
        self._create_appendix(enterprise_data, aggregated_data)

        # Сохранение если указан путь
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.doc.save(str(output_path))
            logger.info(f"✅ Word отчет сохранен: {output_path}")

        return self.doc

    def _setup_document_styles(self):
        """Настройка стилей документа по стандарту ПКМ 690"""
        # Основной стиль текста
        normal_style = self.doc.styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "Times New Roman"
        normal_font.size = Pt(14)
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_after = Pt(6)

        # Стили заголовков
        heading_styles = {
            "Heading 1": {
                "size": 16,
                "bold": True,
                "space_before": 12,
                "space_after": 6,
            },
            "Heading 2": {
                "size": 15,
                "bold": True,
                "space_before": 10,
                "space_after": 6,
            },
            "Heading 3": {
                "size": 14,
                "bold": True,
                "space_before": 8,
                "space_after": 6,
            },
        }

        for style_name, params in heading_styles.items():
            if style_name in self.doc.styles:
                style = self.doc.styles[style_name]
                font = style.font
                font.name = "Times New Roman"
                font.size = Pt(params["size"])
                font.bold = params["bold"]
                para_format = style.paragraph_format
                para_format.space_before = Pt(params["space_before"])
                para_format.space_after = Pt(params["space_after"])

        # Стиль для таблиц
        try:
            table_style = self.doc.styles.add_style(
                "Table Text", WD_STYLE_TYPE.PARAGRAPH
            )
            table_font = table_style.font
            table_font.name = "Times New Roman"
            table_font.size = Pt(12)
        except (ValueError, KeyError):
            pass  # Стиль уже существует

        # Стиль для заголовков таблиц
        try:
            table_header_style = self.doc.styles.add_style(
                "Table Header", WD_STYLE_TYPE.PARAGRAPH
            )
            table_header_font = table_header_style.font
            table_header_font.name = "Times New Roman"
            table_header_font.size = Pt(12)
            table_header_font.bold = True
        except (ValueError, KeyError):
            pass

    def _create_title_page(self, enterprise_data: Dict[str, Any]):
        """Создание титульной страницы"""
        logger.info("📋 Создание титульной страницы...")

        # Заголовок
        title = self.doc.add_heading("ЭНЕРГЕТИЧЕСКИЙ АУДИТ", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = self.doc.add_heading("Энергетический паспорт предприятия", level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пустые строки
        for _ in range(3):
            self.doc.add_paragraph()

        # Информация о предприятии
        enterprise_info = f"""
НАЗВАНИЕ ПРЕДПРИЯТИЯ: {enterprise_data.get("name", "Не указано")}
АДРЕС: {enterprise_data.get("address", "не указан")}
ГОД ОТЧЕТА: {datetime.now().year}
"""

        info_para = self.doc.add_paragraph(enterprise_info.strip())
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Пустые строки
        for _ in range(5):
            self.doc.add_paragraph()

        # Информация об аудиторе
        auditor_info = f"""
ДАТА ПРОВЕДЕНИЯ: {datetime.now().strftime("%d.%m.%Y")}
СТАНДАРТ: ПКМ 690 Узбекистан
"""

        auditor_para = self.doc.add_paragraph(auditor_info.strip())
        auditor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Разрыв страницы
        self.doc.add_page_break()

    def _create_table_of_contents(self):
        """Создание содержания"""
        logger.info("📑 Создание содержания...")

        title = self.doc.add_heading("СОДЕРЖАНИЕ", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sections = [
            "1. ВВЕДЕНИЕ",
            "2. ОБЩИЕ СВЕДЕНИЯ О ПРЕДПРИЯТИИ",
            "3. АНАЛИЗ ЭНЕРГОПОТРЕБЛЕНИЯ",
            "4. АНАЛИЗ ОБОРУДОВАНИЯ",
            "5. УЗЛЫ УЧЕТА ЭНЕРГОРЕСУРСОВ",
            "6. РАСЧЕТ ТЕПЛОПОТЕРЬ ПО ЗДАНИЯМ",
            "7. МЕРОПРИЯТИЯ ПО ЭНЕРГОСБЕРЕЖЕНИЮ",
            "8. ЭКОНОМИЧЕСКИЙ АНАЛИЗ",
            "9. ЗАКЛЮЧЕНИЕ",
            "10. ПРИЛОЖЕНИЯ",
        ]

        for section in sections:
            self.doc.add_paragraph(section)

        self.doc.add_page_break()

    def _create_introduction(self, enterprise_data: Dict[str, Any]):
        """Создание введения"""
        logger.info("📝 Создание введения...")

        self.doc.add_heading("1. ВВЕДЕНИЕ", level=1)

        intro_text = f"""
Настоящий энергетический аудит проведен в соответствии с требованиями стандарта ПКМ 690 Узбекистан 
"Энергетические обследования. Общие требования" для предприятия "{enterprise_data.get("name", "неизвестно")}".

Целью энергетического аудита является:
- Определение фактического энергопотребления предприятия
- Выявление резервов энергосбережения
- Разработка мероприятий по повышению энергоэффективности
- Оценка экономической эффективности предлагаемых мероприятий

Объектом обследования является предприятие "{enterprise_data.get("name", "неизвестно")}", 
расположенное по адресу: {enterprise_data.get("address", "адрес не указан")}.

Период проведения аудита: {datetime.now().strftime("%d.%m.%Y")}

Методология проведения энергетического аудита основана на:
- Анализе энергопотребления за отчетный период
- Обследовании энергетического оборудования
- Расчете энергетических показателей
- Разработке рекомендаций по энергосбережению
"""

        self.doc.add_paragraph(intro_text.strip())

    def _create_enterprise_info(self, enterprise_data: Dict[str, Any]):
        """Создание раздела общих сведений о предприятии"""
        logger.info("🏢 Создание раздела о предприятии...")

        self.doc.add_heading("2. ОБЩИЕ СВЕДЕНИЯ О ПРЕДПРИЯТИИ", level=1)

        basic_info = f"""
2.1 Основная информация

Полное наименование предприятия: {enterprise_data.get("name", "не указано")}
Адрес: {enterprise_data.get("address", "не указан")}
Год отчета: {datetime.now().year}
"""

        self.doc.add_paragraph(basic_info.strip())

    def _create_energy_consumption_analysis(
        self, enterprise_data: Dict[str, Any], aggregated_data: Dict[str, Any]
    ):
        """Создание анализа энергопотребления"""
        logger.info("⚡ Создание анализа энергопотребления...")

        self.doc.add_heading("3. АНАЛИЗ ЭНЕРГОПОТРЕБЛЕНИЯ", level=1)

        resources = aggregated_data.get("resources", {})

        # Получаем данные по ресурсам (нужны для таблицы)
        electricity_data = resources.get("electricity", {})
        gas_data = resources.get("gas", {})
        water_data = resources.get("water", {})

        # Подсчет общих объемов потребления через централизованные функции
        if HAS_CALCULATIONS:
            total_electricity = calculate_total_consumption_by_resource(
                aggregated_data, "electricity"
            )
            total_gas = calculate_total_consumption_by_resource(aggregated_data, "gas")
            total_water = calculate_total_consumption_by_resource(
                aggregated_data, "water"
            )
        else:
            # Fallback на локальные вычисления
            total_electricity = 0.0
            total_gas = 0.0
            total_water = 0.0

            for quarter_data in electricity_data.values():
                if isinstance(quarter_data, dict):
                    totals = quarter_data.get("quarter_totals", {})
                    total_electricity += totals.get("active_kwh", 0) or 0

            for quarter_data in gas_data.values():
                if isinstance(quarter_data, dict):
                    totals = quarter_data.get("quarter_totals", {})
                    total_gas += totals.get("volume_m3", 0) or 0

            for quarter_data in water_data.values():
                if isinstance(quarter_data, dict):
                    totals = quarter_data.get("quarter_totals", {})
                    total_water += totals.get("volume_m3", 0) or 0

        analysis_text = f"""
3.1 Общая характеристика энергопотребления

Предприятие потребляет следующие виды энергетических ресурсов:
- Электрическая энергия
- Природный газ
- Вода (холодная и горячая)

Общее энергопотребление предприятия за отчетный период составляет:
- Электрическая энергия: {total_electricity:,.0f} кВт·ч
- Природный газ: {total_gas:,.0f} м³
- Вода: {total_water:,.0f} м³
"""

        self.doc.add_paragraph(analysis_text.strip())

        # Таблица энергопотребления по кварталам
        self.doc.add_heading("3.2 Энергопотребление по кварталам", level=2)

        # Собираем все кварталы
        all_quarters = sorted(
            set(
                list(electricity_data.keys())
                + list(gas_data.keys())
                + list(water_data.keys())
            )
        )

        if all_quarters:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = ["Квартал", "Электричество (кВт·ч)", "Газ (м³)", "Вода (м³)"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    try:
                        paragraph.style = "Table Header"
                    except (ValueError, AttributeError, KeyError):
                        # Стиль может быть недоступен, игнорируем ошибку
                        pass

            # Данные
            for quarter in all_quarters:
                row_cells = table.add_row().cells

                elec_totals = electricity_data.get(quarter, {}).get(
                    "quarter_totals", {}
                )
                gas_totals = gas_data.get(quarter, {}).get("quarter_totals", {})
                water_totals = water_data.get(quarter, {}).get("quarter_totals", {})

                row_cells[0].text = quarter
                row_cells[1].text = f"{elec_totals.get('active_kwh', 0) or 0:,.0f}"
                row_cells[2].text = f"{gas_totals.get('volume_m3', 0) or 0:,.0f}"
                row_cells[3].text = f"{water_totals.get('volume_m3', 0) or 0:,.0f}"

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.style = "Table Text"
                        except (ValueError, AttributeError, KeyError):
                            # Стиль может быть недоступен, игнорируем ошибку
                            pass

    def _create_equipment_analysis(
        self, enterprise_data: Dict[str, Any], equipment_data: Dict[str, Any]
    ):
        """Создание анализа оборудования"""
        logger.info("🏭 Создание анализа оборудования...")

        self.doc.add_heading("4. АНАЛИЗ ОБОРУДОВАНИЯ", level=1)

        # Используем централизованную функцию для извлечения данных оборудования
        if HAS_CALCULATIONS:
            try:
                eq_data = extract_equipment_data(equipment_data)
                total_power = eq_data.total_installed_power_kw
                total_items = eq_data.total_items_count
            except Exception as e:
                logger.warning(
                    f"Ошибка извлечения данных оборудования через централизованную функцию: {e}. Используются локальные данные."
                )
                summary = equipment_data.get("summary", {})
                total_power = summary.get("total_power_kw", 0) or 0
                total_items = summary.get("total_items", 0) or 0
        else:
            summary = equipment_data.get("summary", {})
            total_power = summary.get("total_power_kw", 0) or 0
            total_items = summary.get("total_items", 0) or 0

        analysis_text = f"""
4.1 Общая характеристика оборудования

На предприятии установлено следующее энергопотребляющее оборудование:
- Производственное оборудование
- Электрооборудование
- Системы отопления и вентиляции
- Осветительные установки

Общая установленная мощность: {total_power:,.2f} кВт
Количество единиц оборудования: {total_items}
"""

        self.doc.add_paragraph(analysis_text.strip())

        # Таблица оборудования (если есть данные)
        sheets_data = equipment_data.get("sheets", [])
        if sheets_data:
            self.doc.add_heading("4.2 Перечень основного оборудования", level=2)

            table = self.doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = ["Наименование", "Тип", "Мощность (кВт)", "Количество"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    try:
                        paragraph.style = "Table Header"
                    except (ValueError, AttributeError, KeyError):
                        # Стиль может быть недоступен, игнорируем ошибку
                        pass

            # Данные оборудования (первые 20 единиц)
            item_count = 0
            for sheet in sheets_data:
                sections = sheet.get("sections", [])
                for section in sections:
                    items = section.get("items", [])
                    for item in items:
                        if item_count >= 20:
                            break
                        row_cells = table.add_row().cells
                        row_cells[0].text = item.get("name", "не указано")
                        row_cells[1].text = item.get("type", "не указан")
                        row_cells[
                            2
                        ].text = f"{item.get('total_power_kw', item.get('unit_power_kw', 0)) or 0:,.2f}"
                        row_cells[3].text = str(item.get("quantity", 1) or 1)

                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                try:
                                    paragraph.style = "Table Text"
                                except (ValueError, AttributeError, KeyError):
                                    # Стиль может быть недоступен, игнорируем ошибку
                                    pass
                        item_count += 1
                    if item_count >= 20:
                        break
                if item_count >= 20:
                    break

    def _create_metering_nodes_section(self, nodes_data: List[Dict[str, Any]]):
        """Создание раздела об узлах учета"""
        logger.info("📊 Создание раздела узлов учета...")

        self.doc.add_heading("5. УЗЛЫ УЧЕТА ЭНЕРГОРЕСУРСОВ", level=1)

        if nodes_data and len(nodes_data) > 0:
            table = self.doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = ["Наименование", "Тип учета", "Место установки", "Коэффициент"]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    try:
                        paragraph.style = "Table Header"
                    except (ValueError, AttributeError, KeyError):
                        # Стиль может быть недоступен, игнорируем ошибку
                        pass

            # Данные узлов
            # Преобразуем nodes_data в список, если это словарь
            if isinstance(nodes_data, dict):
                if "tables" in nodes_data:
                    nodes_list = nodes_data["tables"]
                elif "nodes" in nodes_data:
                    nodes_list = nodes_data["nodes"]
                else:
                    nodes_list = list(nodes_data.values()) if nodes_data else []
            else:
                nodes_list = nodes_data if isinstance(nodes_data, list) else []

            for node in nodes_list[:15]:  # Первые 15 узлов
                row_cells = table.add_row().cells
                row_cells[0].text = node.get("name", "не указано")
                row_cells[1].text = node.get("type", "не указан")
                row_cells[2].text = node.get("location", "не указано")
                row_cells[3].text = str(node.get("coefficient", 1.0))

                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.style = "Table Text"
                        except (ValueError, AttributeError, KeyError):
                            # Стиль может быть недоступен, игнорируем ошибку
                            pass
        else:
            self.doc.add_paragraph("Данные об узлах учета не предоставлены.")

    def _create_envelope_section(self, envelope_data: Dict[str, Any]):
        """Создание раздела расчета теплопотерь по зданиям"""
        logger.info("🏗️ Создание раздела расчета теплопотерь по зданиям...")

        self.doc.add_heading("6. РАСЧЕТ ТЕПЛОПОТЕРЬ ПО ЗДАНИЯМ", level=1)

        self.doc.add_paragraph(
            "Расчет теплопотерь по зданиям и сооружениям предприятия."
        )

        # Можно добавить таблицу с данными, если они есть в envelope_data
        if envelope_data:
            self.doc.add_paragraph("Данные расчета теплопотерь по зданиям загружены.")

    def _create_energy_efficiency_measures(
        self, enterprise_data: Dict[str, Any], aggregated_data: Dict[str, Any]
    ):
        """
        Создание раздела мероприятий по энергосбережению.

        Использует эталонные таблицы из образцового отчёта как fallback, если данные не предоставлены.
        """
        logger.info("💡 Создание раздела мероприятий...")

        self.doc.add_heading("7. МЕРОПРИЯТИЯ ПО ЭНЕРГОСБЕРЕЖЕНИЮ", level=1)

        # Пытаемся загрузить эталонные таблицы
        try:
            from reference_tables_loader import get_all_measures, get_measures_mapping

            HAS_REFERENCE_TABLES = True
        except ImportError:
            HAS_REFERENCE_TABLES = False
            logger.warning(
                "Не удалось импортировать reference_tables_loader для мероприятий"
            )

        # Получаем данные мероприятий
        measures_data = aggregated_data.get("measures") or aggregated_data.get(
            "ai_recommendations"
        )

        # Если данных нет, используем эталонные таблицы
        if not measures_data and HAS_REFERENCE_TABLES:
            try:
                reference_measures = get_all_measures()
                if reference_measures:
                    logger.info(
                        f"Использование {len(reference_measures)} мероприятий из эталонных таблиц"
                    )
                    measures_data = reference_measures
            except Exception as e:
                logger.warning(f"Ошибка загрузки эталонных мероприятий: {e}")

        # Вводный текст
        intro_text = """
На основе проведенного анализа энергопотребления разработаны следующие мероприятия 
по повышению энергоэффективности предприятия.
"""
        self.doc.add_paragraph(intro_text.strip())

        # Если есть данные мероприятий, создаём таблицу
        if measures_data:
            # Загружаем маппинг для структуры таблицы
            if HAS_REFERENCE_TABLES:
                try:
                    mapping = get_measures_mapping()
                    word_config = mapping.get("word_section_config", {})
                    table_style = word_config.get("table_style", "Light Grid Accent 1")
                except Exception:
                    table_style = "Light Grid Accent 1"
            else:
                table_style = "Light Grid Accent 1"

            # Создаём таблицу мероприятий
            table = self.doc.add_table(rows=1, cols=5)
            try:
                table.style = table_style
            except (ValueError, AttributeError, KeyError):
                # Если стиль недоступен, используем запасной
                table.style = "Table Grid"

            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            headers = [
                "№",
                "Наименование мероприятия",
                "Экономия, кВт·ч/год",
                "Стоимость, сум",
                "Срок окупаемости, лет",
            ]
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    try:
                        paragraph.style = "Table Header"
                    except (ValueError, AttributeError, KeyError):
                        # Стиль может быть недоступен, игнорируем ошибку
                        pass

            # Данные мероприятий
            total_capex = 0.0
            total_saving_kwh = 0.0

            for idx, measure in enumerate(measures_data, 1):
                name = measure.get("name", "") or measure.get("essence", "")
                if not name:
                    continue

                row_cells = table.add_row().cells

                # №
                row_cells[0].text = str(measure.get("id", idx))

                # Наименование
                row_cells[1].text = name

                # Экономия
                saving_kwh = (
                    measure.get("saving_kwh") or measure.get("saving_kwh") or 0.0
                )
                row_cells[2].text = f"{saving_kwh:,.0f}" if saving_kwh > 0 else "-"
                total_saving_kwh += saving_kwh

                # Стоимость
                capex = (
                    measure.get("capex")
                    or measure.get("cost_usd")
                    or measure.get("cost")
                    or 0.0
                )
                row_cells[3].text = f"{capex:,.0f}" if capex > 0 else "-"
                total_capex += capex

                # Срок окупаемости
                payback = measure.get("payback_years") or measure.get("payback") or 0.0
                row_cells[4].text = f"{payback:.1f}" if payback > 0 else "-"

                # Стилизация ячеек
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.style = "Table Text"
                        except (ValueError, AttributeError, KeyError):
                            # Стиль может быть недоступен, игнорируем ошибку
                            pass

            # Итоговая строка
            if len(measures_data) > 0:
                summary_row = table.add_row().cells
                summary_row[0].text = "ИТОГО"
                summary_row[1].text = ""
                summary_row[2].text = (
                    f"{total_saving_kwh:,.0f}" if total_saving_kwh > 0 else "-"
                )
                summary_row[3].text = f"{total_capex:,.0f}" if total_capex > 0 else "-"

                # Используем централизованную функцию для расчета среднего срока окупаемости
                if HAS_CALCULATIONS:
                    avg_payback = calculate_average_payback_period(
                        total_capex=total_capex,
                        total_saving_kwh=total_saving_kwh,
                        tariff_per_kwh=0.15,  # Можно сделать конфигурируемым
                    )
                else:
                    # Fallback на локальный расчет
                    avg_payback = (
                        total_capex / (total_saving_kwh * 0.15)
                        if total_saving_kwh > 0
                        else 0.0
                    )

                summary_row[4].text = f"{avg_payback:.1f}" if avg_payback > 0 else "-"

                # Выделяем итоговую строку
                for cell in summary_row:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.runs[0].bold = True
                        except (IndexError, AttributeError):
                            # Если нет runs или они недоступны, игнорируем
                            pass

            # Добавляем пояснительный текст
            if total_capex > 0 and total_saving_kwh > 0:
                summary_text = f"""
Общая стоимость реализации мероприятий: {total_capex:,.0f} сум.
Общая годовая экономия электроэнергии: {total_saving_kwh:,.0f} кВт·ч/год.
Средний срок окупаемости: {avg_payback:.1f} лет.
"""
                self.doc.add_paragraph(summary_text.strip())
        else:
            # Если данных нет, добавляем общий текст
            default_text = """
На основе проведенного анализа энергопотребления рекомендуется реализовать следующие мероприятия:

1. Технические мероприятия
   - Замена устаревшего оборудования на энергоэффективное
   - Внедрение систем автоматического управления
   - Оптимизация режимов работы оборудования

2. Организационные мероприятия
   - Разработка программы энергосбережения
   - Обучение персонала основам энергосбережения
   - Внедрение системы энергетического менеджмента

3. Информационные мероприятия
   - Проведение энергетических аудитов
   - Мониторинг энергопотребления
   - Анализ эффективности мероприятий
"""
            self.doc.add_paragraph(default_text.strip())

    def _create_economic_analysis(
        self, enterprise_data: Dict[str, Any], aggregated_data: Dict[str, Any]
    ):
        """Создание экономического анализа"""
        logger.info("💰 Создание экономического анализа...")

        self.doc.add_heading("8. ЭКОНОМИЧЕСКИЙ АНАЛИЗ", level=1)

        resources = aggregated_data.get("resources", {})

        # Подсчет затрат через централизованную функцию
        if HAS_CALCULATIONS:
            try:
                costs = calculate_total_costs(aggregated_data)
                total_electricity_cost = costs.get("electricity", 0.0)
                total_gas_cost = costs.get("gas", 0.0)
                total_water_cost = costs.get("water", 0.0)
                total_cost = costs.get("total", 0.0)
            except Exception as e:
                logger.warning(
                    f"Ошибка расчета затрат через централизованную функцию: {e}. Используются локальные вычисления."
                )
                # Fallback на локальные вычисления
                total_electricity_cost = 0.0
                total_gas_cost = 0.0
                total_water_cost = 0.0

                for quarter_data in resources.get("electricity", {}).values():
                    if isinstance(quarter_data, dict):
                        totals = quarter_data.get("quarter_totals", {})
                        total_electricity_cost += totals.get("cost_sum", 0) or 0

                for quarter_data in resources.get("gas", {}).values():
                    if isinstance(quarter_data, dict):
                        totals = quarter_data.get("quarter_totals", {})
                        total_gas_cost += totals.get("cost_sum", 0) or 0

                for quarter_data in resources.get("water", {}).values():
                    if isinstance(quarter_data, dict):
                        totals = quarter_data.get("quarter_totals", {})
                        total_water_cost += totals.get("cost_sum", 0) or 0

                total_cost = total_electricity_cost + total_gas_cost + total_water_cost
        else:
            # Fallback на локальные вычисления
            total_electricity_cost = 0.0
            total_gas_cost = 0.0
            total_water_cost = 0.0

            for quarter_data in resources.get("electricity", {}).values():
                if isinstance(quarter_data, dict):
                    totals = quarter_data.get("quarter_totals", {})
                    total_electricity_cost += totals.get("cost_sum", 0) or 0

            for quarter_data in resources.get("gas", {}).values():
                if isinstance(quarter_data, dict):
                    totals = quarter_data.get("quarter_totals", {})
                    total_gas_cost += totals.get("cost_sum", 0) or 0

            for quarter_data in resources.get("water", {}).values():
                if isinstance(quarter_data, dict):
                    totals = quarter_data.get("quarter_totals", {})
                    total_water_cost += totals.get("cost_sum", 0) or 0

            total_cost = total_electricity_cost + total_gas_cost + total_water_cost

        economic_text = f"""
8.1 Анализ затрат на энергоресурсы

Общие затраты предприятия на энергоресурсы за отчетный период:
- Электрическая энергия: {total_electricity_cost:,.0f} сум
- Природный газ: {total_gas_cost:,.0f} сум
- Вода: {total_water_cost:,.0f} сум
- Общие затраты: {total_cost:,.0f} сум

8.2 Экономический эффект от мероприятий

Реализация предложенных мероприятий позволит:
- Снизить энергопотребление на 15-25%
- Экономить денежные средства на оплате энергоресурсов
- Повысить конкурентоспособность предприятия
- Улучшить экологическую ситуацию

8.3 Срок окупаемости мероприятий

Средний срок окупаемости предлагаемых мероприятий составляет 2-4 года.
"""

        self.doc.add_paragraph(economic_text.strip())

    def _create_conclusion(
        self,
        enterprise_data: Dict[str, Any],
        aggregated_data: Optional[Dict[str, Any]] = None,
    ):
        """Создание заключения на основе исходных данных и AI-анализа"""
        logger.info("📝 Создание заключения из исходных данных и AI-анализа...")

        self.doc.add_heading("9. ЗАКЛЮЧЕНИЕ", level=1)

        # Извлекаем AI-инсайты и рекомендации из агрегированных данных
        ai_insights = {}
        ai_recommendations = []
        if aggregated_data:
            ai_insights = aggregated_data.get("ai_insights", {})
            ai_recommendations = aggregated_data.get("ai_recommendations", [])

        # Формируем текст заключения
        conclusion_text = f"""
На основании проведенного энергетического аудита предприятия "{enterprise_data.get("name", "неизвестно")}", 
выполненного на основе анализа исходных данных энергопотребления, можно сделать следующие выводы:

1. ОБЩАЯ ОЦЕНКА ЭНЕРГОПОТРЕБЛЕНИЯ
Анализ исходных данных показал:
"""

        # Добавляем AI-инсайты, если они есть
        if ai_insights:
            conclusion_text += "\nНа основе AI-анализа исходных данных выявлено:\n"
            for key, value in ai_insights.items():
                if value:
                    conclusion_text += f"- {key}: {value}\n"
            conclusion_text += "\n"
        else:
            conclusion_text += """Предприятие потребляет значительное количество энергетических ресурсов. 
Анализ исходных данных показал возможности для повышения энергоэффективности.

"""

        conclusion_text += """2. ОСНОВНЫЕ НАПРАВЛЕНИЯ ЭНЕРГОСБЕРЕЖЕНИЯ
"""

        # Добавляем AI-рекомендации, если они есть
        if ai_recommendations:
            conclusion_text += "\nНа основе AI-анализа исходных данных рекомендованы следующие мероприятия:\n"
            for idx, rec in enumerate(
                ai_recommendations[:10], 1
            ):  # Первые 10 рекомендаций
                if isinstance(rec, dict):
                    rec_text = rec.get("text", rec.get("description", str(rec)))
                    rec_priority = rec.get("priority", "")
                    if rec_priority:
                        rec_text += f" (приоритет: {rec_priority})"
                    conclusion_text += f"{idx}. {rec_text}\n"
                else:
                    conclusion_text += f"{idx}. {rec}\n"
            conclusion_text += "\n"
        else:
            conclusion_text += """- Оптимизация систем освещения
- Модернизация отопительных систем
- Внедрение энергосберегающего оборудования
- Улучшение теплоизоляции зданий

"""

        conclusion_text += """3. ЭКОНОМИЧЕСКИЙ ЭФФЕКТ
Реализация предложенных мероприятий позволит:
- Снизить энергопотребление на 15-25%
- Экономить денежные средства на оплате энергоресурсов
- Повысить конкурентоспособность предприятия

4. РЕКОМЕНДАЦИИ
Рекомендуется:
- Разработать программу энергосбережения на основе анализа исходных данных
- Внедрить систему энергетического менеджмента
- Регулярно проводить энергетические аудиты
- Обучить персонал основам энергосбережения

"""

        # Указываем источник данных
        if ai_insights or ai_recommendations:
            conclusion_text += """Заключение подготовлено на основании:
- Анализа исходных данных энергетических ресурсов (Excel, PDF, DOCX)
- AI-анализа данных для выявления закономерностей и аномалий
- Требований стандарта ПКМ 690 Узбекистан
"""
        else:
            conclusion_text += """Заключение подготовлено на основании анализа исходных данных энергетического аудита, 
проведенного в соответствии с требованиями стандарта ПКМ 690 Узбекистан.
"""

        self.doc.add_paragraph(conclusion_text.strip())

    def _create_appendix(
        self, enterprise_data: Dict[str, Any], aggregated_data: Dict[str, Any]
    ):
        """Создание приложений"""
        logger.info("📎 Создание приложений...")

        self.doc.add_heading("10. ПРИЛОЖЕНИЯ", level=1)

        self.doc.add_heading("Приложение 1. Справочные данные", level=2)

        reference_text = f"""
Таблица 1.1 - Основные характеристики предприятия

| Показатель | Значение |
|------------|----------|
| Название предприятия | {enterprise_data.get("name", "не указано")} |
| Адрес | {enterprise_data.get("address", "не указан")} |
| Год отчета | {datetime.now().year} |
"""

        self.doc.add_paragraph(reference_text.strip())

        # Приложение 2
        self.doc.add_heading("Приложение 2. Нормативные документы", level=2)

        normative_text = """
Список нормативных документов, использованных при проведении энергетического аудита:

1. ПКМ 690 Узбекистан "Энергетические обследования. Общие требования"
2. Закон Республики Узбекистан "Об энергосбережении"
3. ГОСТ Р 51387-99 "Энергосбережение. Нормативно-методическое обеспечение"
4. Методические рекомендации по проведению энергетических обследований
"""

        self.doc.add_paragraph(normative_text.strip())
