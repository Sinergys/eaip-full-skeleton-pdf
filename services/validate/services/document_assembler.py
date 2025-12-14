"""
DocumentAssembler - Сборка финального DOCX документа на основе GOST шаблона.
Соответствует разделу 3.1.7 ТЗ.

Основные задачи:
- Загрузка GOST шаблона
- Вставка исправленного текста с применением стилей
- Восстановление объектов (замена [[OBJ_XXX]] на оригинальные)
- Добавление секции "AI Summary and Recommendations"
- Сохранение финального документа
"""
import io
import re
import logging
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from core.models import ExtractedObject, ProcessingSummary
from core.constants import OBJECT_MARKER_PATTERN
from utils.exceptions import TemplateError, DocumentAssemblyError

logger = logging.getLogger(__name__)


class DocumentAssembler:
    """
    Сборка финального DOCX документа на основе шаблона.
    Соответствует разделу 3.1.7 ТЗ.
    """

    def __init__(self, template_path: Path):
        """
        Инициализация assembler.

        Args:
            template_path: Путь к GOST шаблону

        Raises:
            TemplateError: Если шаблон не найден
        """
        self.template_path = template_path

        # Валидация существования шаблона
        if not template_path.exists():
            raise TemplateError(f"Template not found: {template_path}")

        logger.info(f"DocumentAssembler initialized with template: {template_path}")

    async def assemble_document(
        self,
        corrected_text: str,
        objects: Dict[str, ExtractedObject],
        recommendations: List[str],
        summary: ProcessingSummary,
        original_filename: str,
        output_dir: Optional[Path] = None
    ) -> str:
        """
        Создать финальный документ на основе шаблона.

        Args:
            corrected_text: Исправленный текст с маркерами [[OBJ_XXX]]
            objects: Словарь извлечённых объектов
            recommendations: Список рекомендаций
            summary: Итоговая сводка
            original_filename: Имя оригинального файла
            output_dir: Директория для сохранения (по умолчанию рядом с original)

        Returns:
            Путь к созданному файлу [Original]_Проверенный.docx

        Raises:
            DocumentAssemblyError: При ошибках сборки
        """
        try:
            logger.info(f"Starting document assembly for: {original_filename}")

            # 1. Загрузить шаблон
            document = self._load_template()

            # 2. Применить стили ГОСТ
            self._apply_gost_formatting(document)

            # 3. Вставить исправленный текст
            self._insert_corrected_text(document, corrected_text)

            # 4. Заменить маркеры на объекты
            self._restore_objects(document, objects)

            # 5. Добавить секцию рекомендаций
            self._add_recommendations_section(document, recommendations, summary)

            # 6. Сохранить документ
            output_path = self._save_document(document, original_filename, output_dir)

            logger.info(f"Document assembly complete: {output_path}")
            return str(output_path)

        except (TemplateError, DocumentAssemblyError):
            raise
        except Exception as e:
            logger.error(f"Unexpected error during document assembly: {str(e)}", exc_info=True)
            raise DocumentAssemblyError(f"Document assembly failed: {str(e)}")

    def _load_template(self) -> Document:
        """
        Загрузить GOST шаблон.

        Returns:
            Document: python-docx Document объект

        Raises:
            TemplateError: При ошибке загрузки
        """
        try:
            document = Document(str(self.template_path))
            logger.info("Template loaded successfully")
            return document

        except Exception as e:
            raise TemplateError(f"Failed to load template: {str(e)}")

    def _insert_corrected_text(
        self,
        document: Document,
        text: str
    ) -> None:
        """
        Вставить исправленный текст в документ.

        Применить стили ГОСТ:
        - Шрифт: Times New Roman
        - Размер: 14pt для основного текста
        - Межстрочный интервал: 1.5
        - Отступы по ГОСТ

        Args:
            document: python-docx Document объект
            text: Исправленный текст с маркерами
        """
        try:
            # Разбить текст на параграфы
            paragraphs = text.split('\n\n')

            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                # Добавить параграф
                paragraph = document.add_paragraph()

                # Применить стили ГОСТ
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.line_spacing = 1.5
                paragraph.paragraph_format.first_line_indent = Cm(1.25)  # Красная строка

                # Добавить текст
                run = paragraph.add_run(para_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)

            logger.info(f"Inserted {len(paragraphs)} paragraphs")

        except Exception as e:
            raise DocumentAssemblyError(f"Failed to insert corrected text: {str(e)}")

    def _restore_objects(
        self,
        document: Document,
        objects: Dict[str, ExtractedObject]
    ) -> None:
        """
        Найти все маркеры [[OBJ_XXX]] и заменить на оригинальные объекты.

        ВАЖНО:
        - Сохранить подписи к объектам
        - Применить форматирование по ГОСТ

        Args:
            document: python-docx Document объект
            objects: Словарь извлечённых объектов
        """
        try:
            objects_restored = 0

            # Iterate через все параграфы
            for paragraph in document.paragraphs:
                para_text = paragraph.text

                # Найти все маркеры в параграфе
                markers = re.findall(OBJECT_MARKER_PATTERN, para_text)

                for marker in markers:
                    # Извлечь obj_id из маркера [[OBJ_XXX]]
                    obj_id = marker.strip('[]')

                    if obj_id in objects:
                        obj = objects[obj_id]

                        # Заменить маркер на объект
                        if obj.object_type == "image":
                            self._insert_image(paragraph, obj, marker)
                            objects_restored += 1

                        elif obj.object_type == "table":
                            self._insert_table(document, paragraph, obj, marker)
                            objects_restored += 1

                        elif obj.object_type == "chart":
                            # Charts пока пропускаем (сложная вставка)
                            logger.warning(f"Chart restoration not implemented: {obj_id}")

            logger.info(f"Restored {objects_restored} objects")

        except Exception as e:
            logger.warning(f"Error restoring objects: {str(e)}")
            # Не поднимаем exception, просто логируем warning

    def _insert_image(
        self,
        paragraph,
        obj: ExtractedObject,
        marker: str
    ) -> None:
        """
        Вставить изображение в параграф.

        Args:
            paragraph: python-docx Paragraph объект
            obj: ExtractedObject с изображением
            marker: Маркер для замены
        """
        try:
            # Очистить параграф от маркера
            paragraph.clear()

            # Вставить изображение
            if obj.binary_data:
                # Создать BytesIO объект
                image_stream = io.BytesIO(obj.binary_data)

                # Добавить изображение с шириной по умолчанию
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Cm(14))  # 14cm width для A4

                # Выравнивание по центру
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Если есть caption, добавить его
                if obj.caption:
                    caption_para = paragraph._element.getparent().add_p()
                    caption_run = caption_para.add_r()
                    caption_run.text = obj.caption
                    # Стили для caption
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run.font.size = Pt(12)
                    caption_run.font.italic = True

                logger.debug(f"Inserted image: {obj.id}")

        except Exception as e:
            logger.warning(f"Failed to insert image {obj.id}: {str(e)}")
            # Восстановить маркер если не удалось вставить
            paragraph.add_run(marker)

    def _insert_table(
        self,
        document: Document,
        paragraph,
        obj: ExtractedObject,
        marker: str
    ) -> None:
        """
        Вставить таблицу в документ.

        Args:
            document: python-docx Document объект
            paragraph: python-docx Paragraph объект
            obj: ExtractedObject с таблицей
            marker: Маркер для замены
        """
        try:
            # Получить данные таблицы
            table_data = obj.metadata.get('data', [])

            if not table_data:
                logger.warning(f"No table data for {obj.id}")
                return

            # Очистить параграф от маркера
            paragraph.clear()

            # Создать таблицу
            rows_count = len(table_data)
            cols_count = len(table_data[0]) if table_data else 0

            if rows_count == 0 or cols_count == 0:
                logger.warning(f"Empty table data for {obj.id}")
                return

            table = document.add_table(rows=rows_count, cols=cols_count)
            table.style = 'Table Grid'  # Стиль с границами

            # Заполнить таблицу
            for row_idx, row_data in enumerate(table_data):
                for col_idx, cell_text in enumerate(row_data):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text

                    # Применить форматирование
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

            # Если есть caption, добавить его
            if obj.caption:
                caption_para = document.add_paragraph()
                caption_para.add_run(obj.caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.runs[0].font.size = Pt(12)
                caption_para.runs[0].font.italic = True

            logger.debug(f"Inserted table: {obj.id} ({rows_count}x{cols_count})")

        except Exception as e:
            logger.warning(f"Failed to insert table {obj.id}: {str(e)}")

    def _add_recommendations_section(
        self,
        document: Document,
        recommendations: List[str],
        summary: ProcessingSummary
    ) -> None:
        """
        Добавить секцию "AI Summary and Recommendations" в конец документа.

        Формат:
        - Заголовок секции
        - Краткая сводка (summary)
        - Список рекомендаций

        Args:
            document: python-docx Document объект
            recommendations: Список рекомендаций
            summary: Итоговая сводка
        """
        try:
            # Добавить разделитель
            document.add_page_break()

            # Заголовок секции
            heading = document.add_heading('AI Summary and Recommendations', level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Краткая сводка
            summary_para = document.add_paragraph()
            summary_para.add_run('Сводка обработки:').bold = True
            summary_para.paragraph_format.space_after = Pt(6)

            # Детали сводки
            details = [
                f"Обработано чанков: {summary.total_chunks}",
                f"Обнаружено проблем: {summary.total_issues_found}",
                f"Рекомендаций: {summary.total_recommendations}",
                f"Время обработки: {summary.processing_time_seconds:.1f} сек"
            ]

            for detail in details:
                para = document.add_paragraph(detail, style='List Bullet')
                para.runs[0].font.name = 'Times New Roman'
                para.runs[0].font.size = Pt(12)

            # Список рекомендаций
            if recommendations:
                document.add_paragraph()
                rec_heading = document.add_paragraph()
                rec_heading.add_run('Рекомендации по доработке:').bold = True
                rec_heading.paragraph_format.space_after = Pt(6)

                for idx, rec in enumerate(recommendations, 1):
                    rec_para = document.add_paragraph(f"{idx}. {rec}")
                    rec_para.runs[0].font.name = 'Times New Roman'
                    rec_para.runs[0].font.size = Pt(12)
                    rec_para.paragraph_format.left_indent = Cm(1)

            logger.info(f"Added recommendations section: {len(recommendations)} recommendations")

        except Exception as e:
            logger.warning(f"Failed to add recommendations section: {str(e)}")

    def _apply_gost_formatting(self, document: Document) -> None:
        """
        Применить форматирование по ГОСТ:
        - Поля: 2см сверху/снизу, 3см слева, 1см справа
        - Шрифт: Times New Roman 14pt
        - Интервал: 1.5
        - Выравнивание: по ширине

        Args:
            document: python-docx Document объект
        """
        try:
            # Установка полей для всех секций
            for section in document.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(3)
                section.right_margin = Cm(1)

            logger.info("Applied GOST formatting")

        except Exception as e:
            logger.warning(f"Failed to apply GOST formatting: {str(e)}")

    def _save_document(
        self,
        document: Document,
        original_filename: str,
        output_dir: Optional[Path] = None
    ) -> Path:
        """
        Сохранить финальный документ.

        Имя файла: [Original]_Проверенный.docx

        Args:
            document: python-docx Document объект
            original_filename: Имя оригинального файла
            output_dir: Директория для сохранения

        Returns:
            Path: Путь к сохранённому файлу

        Raises:
            DocumentAssemblyError: При ошибке сохранения
        """
        try:
            # Определить output директорию
            if output_dir is None:
                output_dir = Path.cwd() / "output"

            output_dir.mkdir(parents=True, exist_ok=True)

            # Сформировать имя файла
            original_stem = Path(original_filename).stem
            output_filename = f"{original_stem}_Проверенный.docx"
            output_path = output_dir / output_filename

            # Сохранить документ
            document.save(str(output_path))

            logger.info(f"Document saved: {output_path}")
            return output_path

        except Exception as e:
            raise DocumentAssemblyError(f"Failed to save document: {str(e)}")
