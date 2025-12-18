"""
Validation Module - Проверка целостности данных Word Document Validator.
Соответствует требованиям безопасности и целостности данных.

Основные задачи:
- Проверка количества объектов (input = output)
- Валидация безопасности (размер, ZIP, память)
- Проверка целостности данных
- Интеграция с pipeline
"""
import logging
import os
import zipfile
import io
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from docx import Document

from core.models import ExtractedObject
from utils.exceptions import ValidationError, SecurityError

logger = logging.getLogger(__name__)


class DocumentValidator:
    """
    Валидатор целостности данных для Word Document Validator.
    
    Обеспечивает:
    - Проверку соответствия количества объектов
    - Валидацию безопасности
    - Проверку целостности данных
    """

    def __init__(self, max_file_size_mb: int = 100):
        """
        Инициализация валидатора.

        Args:
            max_file_size_mb: Максимальный размер файла в МБ
        """
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.validation_stats = {
            'files_checked': 0,
            'objects_input': 0,
            'objects_output': 0,
            'validation_passed': 0,
            'validation_failed': 0,
            'security_violations': 0
        }

    async def validate_extraction_integrity(
        self,
        input_objects: Dict[str, ExtractedObject],
        output_objects: Dict[str, ExtractedObject],
        file_path: str,
        extraction_stats: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Проверить целостность извлечения объектов.

        КРИТИЧЕСКАЯ ПРОВЕРКА: 47 объектов input → 47 объектов output

        Args:
            input_objects: Объекты, извлеченные из оригинального документа
            output_objects: Объекты для восстановления в финальном документе
            file_path: Путь к файлу для логирования
            extraction_stats: Статистика извлечения (опционально)

        Returns:
            Dict с результатами валидации

        Raises:
            ValidationError: При критических ошибках целостности
        """
        try:
            logger.info(f"Starting integrity validation for: {file_path}")
            
            # Базовые проверки
            input_count = len(input_objects)
            output_count = len(output_objects)
            
            self.validation_stats['files_checked'] += 1
            self.validation_stats['objects_input'] += input_count
            self.validation_stats['objects_output'] += output_count

            validation_result = {
                'file_path': file_path,
                'input_objects_count': input_count,
                'output_objects_count': output_count,
                'integrity_status': 'PASS',
                'validation_details': [],
                'security_checks': {}
            }

            # КРИТИЧЕСКАЯ ПРОВЕРКА: Соответствие количества объектов
            if input_count != output_count:
                error_msg = f"Object count mismatch: {input_count} input vs {output_count} output"
                validation_result['integrity_status'] = 'FAIL'
                validation_result['validation_details'].append(f"CRITICAL: {error_msg}")
                self.validation_stats['validation_failed'] += 1
                logger.error(error_msg)
                raise ValidationError(error_msg)

            # Проверка соответствия ID объектов
            input_ids = set(input_objects.keys())
            output_ids = set(output_objects.keys())
            
            if input_ids != output_ids:
                missing_ids = input_ids - output_ids
                extra_ids = output_ids - input_ids
                
                error_msg = f"Object ID mismatch - Missing: {missing_ids}, Extra: {extra_ids}"
                validation_result['integrity_status'] = 'FAIL'
                validation_result['validation_details'].append(f"CRITICAL: {error_msg}")
                self.validation_stats['validation_failed'] += 1
                logger.error(error_msg)
                raise ValidationError(error_msg)

            # Проверка типа объектов
            type_mismatches = []
            for obj_id in input_ids:
                input_obj = input_objects[obj_id]
                output_obj = output_objects[obj_id]
                
                if input_obj.object_type != output_obj.object_type:
                    type_mismatches.append(f"{obj_id}: {input_obj.object_type} vs {output_obj.object_type}")

            if type_mismatches:
                error_msg = f"Object type mismatches: {type_mismatches}"
                validation_result['integrity_status'] = 'FAIL'
                validation_result['validation_details'].append(f"CRITICAL: {error_msg}")
                self.validation_stats['validation_failed'] += 1
                logger.error(error_msg)
                raise ValidationError(error_msg)

            # Проверка размеров binary_data для изображений
            size_issues = []
            for obj_id in input_ids:
                input_obj = input_objects[obj_id]
                output_obj = output_objects[obj_id]
                
                if input_obj.object_type == "image" and input_obj.binary_data:
                    input_size = len(input_obj.binary_data)
                    output_size = len(output_obj.binary_data) if output_obj.binary_data else 0
                    
                    if input_size != output_size:
                        size_issues.append(f"{obj_id}: {input_size} vs {output_size} bytes")

            if size_issues:
                error_msg = f"Binary data size mismatches: {size_issues}"
                validation_result['integrity_status'] = 'FAIL'
                validation_result['validation_details'].append(f"CRITICAL: {error_msg}")
                self.validation_stats['validation_failed'] += 1
                logger.error(error_msg)
                raise ValidationError(error_msg)

            # Если все проверки прошли успешно
            validation_result['validation_details'].append(f"SUCCESS: All {input_count} objects match perfectly")
            self.validation_stats['validation_passed'] += 1
            
            logger.info(f"Integrity validation PASSED: {input_count} objects verified")
            return validation_result

        except ValidationError:
            raise
        except Exception as e:
            error_msg = f"Unexpected error during integrity validation: {str(e)}"
            validation_result['integrity_status'] = 'ERROR'
            validation_result['validation_details'].append(f"ERROR: {error_msg}")
            self.validation_stats['validation_failed'] += 1
            logger.error(error_msg, exc_info=True)
            raise ValidationError(error_msg)

    async def validate_security(
        self,
        file_path: str,
        max_file_size_mb: Optional[int] = None
    ) -> Dict[str, Any]:
        """
        Проверить безопасность файла.

        Проверяет:
        - Размер файла
        - ZIP структуру
        - Потенциальные угрозы

        Args:
            file_path: Путь к файлу для проверки
            max_file_size_mb: Максимальный размер файла (по умолчанию из __init__)

        Returns:
            Dict с результатами проверки безопасности

        Raises:
            SecurityError: При нарушениях безопасности
        """
        try:
            logger.info(f"Starting security validation for: {file_path}")
            
            path = Path(file_path)
            security_result = {
                'file_path': file_path,
                'security_status': 'PASS',
                'security_details': [],
                'file_size_mb': 0,
                'is_valid_zip': False
            }

            # Проверка существования файла
            if not path.exists():
                error_msg = f"File not found: {file_path}"
                security_result['security_status'] = 'FAIL'
                security_result['security_details'].append(error_msg)
                self.validation_stats['security_violations'] += 1
                logger.error(error_msg)
                raise SecurityError(error_msg)

            # Проверка размера файла
            file_size_bytes = path.stat().st_size
            file_size_mb = file_size_bytes / (1024 * 1024)
            security_result['file_size_mb'] = round(file_size_mb, 2)

            max_size_mb = max_file_size_mb or (self.max_file_size_bytes / (1024 * 1024))
            
            if file_size_bytes > self.max_file_size_bytes:
                error_msg = f"File too large: {file_size_mb:.2f} MB > {max_size_mb} MB limit"
                security_result['security_status'] = 'FAIL'
                security_result['security_details'].append(error_msg)
                self.validation_stats['security_violations'] += 1
                logger.error(error_msg)
                raise SecurityError(error_msg)

            # Проверка ZIP структуры (DOCX - это ZIP архив)
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # Проверить основные файлы DOCX
                    required_files = ['[Content_Types].xml', '_rels/.rels']
                    zip_files = zip_file.namelist()
                    
                    missing_files = [f for f in required_files if f not in zip_files]
                    if missing_files:
                        error_msg = f"Invalid DOCX structure - missing files: {missing_files}"
                        security_result['security_status'] = 'FAIL'
                        security_result['security_details'].append(error_msg)
                        self.validation_stats['security_violations'] += 1
                        logger.error(error_msg)
                        raise SecurityError(error_msg)

                    # Проверить на подозрительные файлы
                    suspicious_patterns = ['.exe', '.bat', '.cmd', '.ps1', '.vbs']
                    suspicious_files = [f for f in zip_files 
                                      if any(pattern in f.lower() for pattern in suspicious_patterns)]
                    
                    if suspicious_files:
                        error_msg = f"Potentially dangerous files found: {suspicious_files}"
                        security_result['security_status'] = 'FAIL'
                        security_result['security_details'].append(error_msg)
                        self.validation_stats['security_violations'] += 1
                        logger.error(error_msg)
                        raise SecurityError(error_msg)

                    security_result['is_valid_zip'] = True
                    security_result['security_details'].append(f"ZIP structure valid, {len(zip_files)} files")

            except zipfile.BadZipFile:
                error_msg = "Invalid ZIP file format - not a valid DOCX"
                security_result['security_status'] = 'FAIL'
                security_result['security_details'].append(error_msg)
                self.validation_stats['security_violations'] += 1
                logger.error(error_msg)
                raise SecurityError(error_msg)

            # Проверка памяти (оценочная)
            estimated_memory_mb = file_size_mb * 3  # Оценка: в 3 раза больше для обработки
            if estimated_memory_mb > 500:  # Лимит 500MB памяти
                warning_msg = f"High memory usage expected: ~{estimated_memory_mb:.1f} MB"
                security_result['security_details'].append(f"WARNING: {warning_msg}")
                logger.warning(warning_msg)

            security_result['security_details'].append(f"Security validation PASSED: {file_size_mb:.2f} MB")
            logger.info(f"Security validation PASSED: {file_path}")
            return security_result

        except SecurityError:
            raise
        except Exception as e:
            error_msg = f"Unexpected error during security validation: {str(e)}"
            security_result['security_status'] = 'ERROR'
            security_result['security_details'].append(error_msg)
            self.validation_stats['security_violations'] += 1
            logger.error(error_msg, exc_info=True)
            raise SecurityError(error_msg)

    async def validate_document_structure(
        self,
        file_path: str,
        expected_objects: Optional[Dict[str, ExtractedObject]] = None
    ) -> Dict[str, Any]:
        """
        Проверить структуру документа и содержимое.

        Args:
            file_path: Путь к документу
            expected_objects: Ожидаемые объекты для проверки

        Returns:
            Dict с результатами проверки структуры
        """
        try:
            logger.info(f"Starting document structure validation for: {file_path}")
            
            structure_result = {
                'file_path': file_path,
                'structure_status': 'PASS',
                'structure_details': [],
                'document_stats': {}
            }

            # Попытка открыть документ
            try:
                document = Document(file_path)
            except Exception as e:
                error_msg = f"Cannot open document: {str(e)}"
                structure_result['structure_status'] = 'FAIL'
                structure_result['structure_details'].append(error_msg)
                logger.error(error_msg)
                raise ValidationError(error_msg)

            # Подсчет элементов документа
            paragraphs_count = len(document.paragraphs)
            tables_count = len(document.tables)
            
            structure_result['document_stats'] = {
                'paragraphs': paragraphs_count,
                'tables': tables_count,
                'sections': len(document.sections)
            }

            # Проверка на наличие контента
            if paragraphs_count == 0:
                warning_msg = "Document has no paragraphs"
                structure_result['structure_details'].append(f"WARNING: {warning_msg}")
                logger.warning(warning_msg)

            if tables_count == 0:
                structure_result['structure_details'].append("INFO: No tables found")
            
            # Проверка изображений в документе
            try:
                images_count = len(document.part.rels)
                structure_result['document_stats']['relationships'] = images_count
                structure_result['structure_details'].append(f"Found {images_count} relationships")
            except Exception:
                structure_result['structure_details'].append("WARNING: Could not count images")

            # Проверка соответствия ожидаемым объектам
            if expected_objects:
                object_types = {}
                for obj in expected_objects.values():
                    obj_type = obj.object_type
                    object_types[obj_type] = object_types.get(obj_type, 0) + 1
                
                structure_result['document_stats']['expected_objects'] = object_types
                structure_result['structure_details'].append(f"Expected objects: {object_types}")

            structure_result['structure_details'].append(f"Structure validation PASSED: {paragraphs_count} paragraphs, {tables_count} tables")
            logger.info(f"Document structure validation PASSED: {file_path}")
            return structure_result

        except Exception as e:
            error_msg = f"Unexpected error during structure validation: {str(e)}"
            structure_result['structure_status'] = 'ERROR'
            structure_result['structure_details'].append(error_msg)
            logger.error(error_msg, exc_info=True)
            raise ValidationError(error_msg)

    def get_validation_stats(self) -> Dict[str, Any]:
        """
        Получить статистику валидации.

        Returns:
            Dict со статистикой
        """
        return self.validation_stats.copy()

    def reset_stats(self) -> None:
        """Сбросить статистику валидации."""
        self.validation_stats = {
            'files_checked': 0,
            'objects_input': 0,
            'objects_output': 0,
            'validation_passed': 0,
            'validation_failed': 0,
            'security_violations': 0
        }
        logger.info("Validation statistics reset")


class ValidationResult:
    """
    Результат валидации для удобного доступа к данным.
    """
    
    def __init__(self, validation_data: Dict[str, Any]):
        self.data = validation_data
        
    @property
    def is_valid(self) -> bool:
        """Проверить, прошла ли валидация успешно."""
        return self.data.get('integrity_status', 'FAIL') == 'PASS'
        
    @property
    def is_secure(self) -> bool:
        """Проверить, прошла ли проверка безопасности."""
        return self.data.get('security_status', 'FAIL') == 'PASS'
        
    @property
    def objects_count(self) -> int:
        """Получить количество объектов."""
        return self.data.get('input_objects_count', 0)
        
    def get_details(self) -> List[str]:
        """Получить детали валидации."""
        return self.data.get('validation_details', [])