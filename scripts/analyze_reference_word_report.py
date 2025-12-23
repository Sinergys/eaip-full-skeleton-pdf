"""
Анализ образцового Word-отчёта "МЕТИН ИРОДА ОТЧЕТ 1107.docx".

Извлекает структуру, таблицы, графики и создаёт контент-гайд для генерации отчётов.
"""

import json
import sys
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.document import Document as DocumentType

PROJECT_ROOT = Path(__file__).resolve().parent.parent
REFERENCE_DOC_PATH = Path(r"C:\Users\DELL\Downloads\Telegram Desktop\МЕТИН ИРОДА ОТЧЕТ 1107.docx")


def analyze_document_structure(doc: DocumentType) -> Dict[str, Any]:
    """Анализирует структуру документа: разделы, заголовки, параграфы."""
    structure = {
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "total_sections": len(doc.sections),
        "headings": [],
        "sections": [],
    }
    
    # Извлекаем заголовки
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading'):
            heading_level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 0
            structure["headings"].append({
                "index": i,
                "level": heading_level,
                "text": para.text.strip(),
                "style": para.style.name
            })
    
    return structure


def extract_tables(doc: DocumentType) -> List[Dict[str, Any]]:
    """Извлекает все таблицы из документа с их данными."""
    tables_data = []
    
    for table_idx, table in enumerate(doc.tables):
        table_data = {
            "index": table_idx,
            "rows": len(table.rows),
            "cols": len(table.columns) if table.rows else 0,
            "data": [],
            "headers": [],
            "type": "unknown",  # Будет определено позже
            "context": "",  # Текст до/после таблицы для контекста
        }
        
        # Извлекаем данные таблицы
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            
            if row_idx == 0:
                # Первая строка - заголовки
                table_data["headers"] = row_data
            else:
                table_data["data"].append(row_data)
        
        # Определяем тип таблицы по заголовкам
        headers_text = " ".join(table_data["headers"]).lower()
        if any(word in headers_text for word in ["баланс", "balance", "технологические", "собственные"]):
            table_data["type"] = "energy_balance"
        elif any(word in headers_text for word in ["динамика", "динамик", "квартал", "quarter"]):
            table_data["type"] = "consumption_dynamics"
        elif any(word in headers_text for word in ["структура", "structure", "потребление"]):
            table_data["type"] = "consumption_structure"
        elif any(word in headers_text for word in ["мероприятия", "measures", "рекомендации", "экономия"]):
            table_data["type"] = "energy_measures"
        elif any(word in headers_text for word in ["оборудование", "equipment", "мощность", "power"]):
            table_data["type"] = "equipment"
        elif any(word in headers_text for word in ["потери", "losses", "трансформатор"]):
            table_data["type"] = "losses"
        elif any(word in headers_text for word in ["удельный", "specific", "расход", "consumption"]):
            table_data["type"] = "specific_consumption"
        else:
            table_data["type"] = "other"
        
        tables_data.append(table_data)
    
    return tables_data


def extract_sections_content(doc: DocumentType) -> List[Dict[str, Any]]:
    """Извлекает содержимое разделов документа."""
    sections = []
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Проверяем, является ли параграф заголовком
        if para.style.name.startswith('Heading'):
            # Сохраняем предыдущий раздел
            if current_section:
                sections.append(current_section)
            
            # Начинаем новый раздел
            level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 0
            current_section = {
                "title": text,
                "level": level,
                "style": para.style.name,
                "content": [],
                "has_tables": False,
                "has_images": False,
            }
        elif current_section:
            # Добавляем контент в текущий раздел
            current_section["content"].append({
                "text": text,
                "style": para.style.name,
            })
    
    # Добавляем последний раздел
    if current_section:
        sections.append(current_section)
    
    return sections


def analyze_content_style(sections: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Анализирует стиль написания разделов."""
    style_guide = {
        "introduction": {
            "typical_phrases": [],
            "structure": [],
            "detail_level": "medium",
        },
        "enterprise_info": {
            "typical_phrases": [],
            "structure": [],
            "detail_level": "high",
        },
        "energy_analysis": {
            "typical_phrases": [],
            "structure": [],
            "detail_level": "high",
        },
        "measures": {
            "typical_phrases": [],
            "structure": [],
            "detail_level": "high",
        },
    }
    
    # Извлекаем типичные фразы из каждого раздела
    for section in sections:
        title_lower = section["title"].lower()
        
        # Определяем тип раздела
        section_type = None
        if any(word in title_lower for word in ["введение", "вводная", "introduction"]):
            section_type = "introduction"
        elif any(word in title_lower for word in ["предприятие", "enterprise", "характеристика"]):
            section_type = "enterprise_info"
        elif any(word in title_lower for word in ["потребление", "consumption", "анализ", "analysis"]):
            section_type = "energy_analysis"
        elif any(word in title_lower for word in ["мероприятия", "measures", "рекомендации"]):
            section_type = "measures"
        
        if section_type:
            # Извлекаем типичные фразы
            for content_item in section["content"]:
                text = content_item["text"]
                # Ищем типичные конструкции
                if "анализ показал" in text.lower() or "analysis showed" in text.lower():
                    style_guide[section_type]["typical_phrases"].append("Анализ показал, что...")
                if "основными потребителями" in text.lower() or "main consumers" in text.lower():
                    style_guide[section_type]["typical_phrases"].append("Основными потребителями являются...")
                if "целесообразно реализовать" in text.lower() or "recommended to implement" in text.lower():
                    style_guide[section_type]["typical_phrases"].append("На предприятии целесообразно реализовать следующие мероприятия...")
    
    return style_guide


def main():
    """Основная функция анализа."""
    print("🔍 АНАЛИЗ ОБРАЗЦОВОГО WORD-ОТЧЁТА")
    print("=" * 80)
    
    if not REFERENCE_DOC_PATH.exists():
        print(f"❌ Файл не найден: {REFERENCE_DOC_PATH}")
        return 1
    
    print(f"📄 Открытие файла: {REFERENCE_DOC_PATH}")
    
    try:
        doc = Document(str(REFERENCE_DOC_PATH))
    except Exception as e:
        print(f"❌ Ошибка открытия файла: {e}")
        return 1
    
    # Анализируем структуру
    print("\n📋 Анализ структуры документа...")
    structure = analyze_document_structure(doc)
    print(f"  - Параграфов: {structure['total_paragraphs']}")
    print(f"  - Таблиц: {structure['total_tables']}")
    print(f"  - Секций: {structure['total_sections']}")
    print(f"  - Заголовков: {len(structure['headings'])}")
    
    # Извлекаем таблицы
    print("\n📊 Извлечение таблиц...")
    tables = extract_tables(doc)
    print(f"  - Найдено таблиц: {len(tables)}")
    
    # Группируем таблицы по типам
    tables_by_type = {}
    for table in tables:
        table_type = table["type"]
        if table_type not in tables_by_type:
            tables_by_type[table_type] = []
        tables_by_type[table_type].append(table)
    
    print("\n  Типы таблиц:")
    for table_type, type_tables in tables_by_type.items():
        print(f"    - {table_type}: {len(type_tables)} таблиц")
    
    # Извлекаем разделы
    print("\n📑 Извлечение разделов...")
    sections = extract_sections_content(doc)
    print(f"  - Найдено разделов: {len(sections)}")
    
    for section in sections[:10]:  # Показываем первые 10
        print(f"    - {section['title']} (уровень {section['level']}, параграфов: {len(section['content'])})")
    
    # Анализируем стиль
    print("\n✍️  Анализ стиля написания...")
    style_guide = analyze_content_style(sections)
    
    # Сохраняем результаты
    output_dir = PROJECT_ROOT / "data" / "reference_analysis"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    results = {
        "document_structure": structure,
        "tables": tables,
        "tables_by_type": {k: len(v) for k, v in tables_by_type.items()},
        "sections": sections,
        "style_guide": style_guide,
        "analysis_date": str(Path(__file__).stat().st_mtime),
    }
    
    output_file = output_dir / "reference_word_report_analysis.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"\n💾 Результаты сохранены: {output_file}")
    
    # Создаём контент-гайд
    print("\n📝 Создание контент-гайда...")
    create_content_guide(results, output_dir / "content_guide.md")
    
    print("\n✅ Анализ завершён!")
    return 0


def create_content_guide(analysis_results: Dict[str, Any], output_path: Path):
    """Создаёт контент-гайд на основе анализа."""
    guide = """# 📝 Контент-гайд для генерации Word-отчётов по ПКМ-690

**Источник:** Анализ образцового отчёта "МЕТИН ИРОДА ОТЧЕТ 1107.docx"  
**Дата анализа:** {analysis_date}

---

## 📊 Статистика документа

- **Параграфов:** {total_paragraphs}
- **Таблиц:** {total_tables}
- **Секций:** {total_sections}
- **Заголовков:** {total_headings}

---

## 📑 Структура разделов

{section_list}

---

## 📊 Типы таблиц

{table_types}

---

## ✍️ Стиль написания

{style_guide}

---

## 🔧 Рекомендации по использованию

1. **Таблицы:** Использовать структуру таблиц из образцового отчёта как эталон
2. **Стиль:** Следовать типичным фразам и структуре разделов
3. **Данные:** Все числовые показатели должны браться из `energy_passport_calculations.py`
4. **Согласованность:** Таблицы в Word должны совпадать с Excel-паспортом

""".format(
        analysis_date=analysis_results.get("analysis_date", "N/A"),
        total_paragraphs=analysis_results["document_structure"]["total_paragraphs"],
        total_tables=analysis_results["document_structure"]["total_tables"],
        total_sections=analysis_results["document_structure"]["total_sections"],
        total_headings=len(analysis_results["document_structure"]["headings"]),
        section_list="\n".join([
            f"- **{s['title']}** (уровень {s['level']}, {len(s['content'])} параграфов)"
            for s in analysis_results["sections"][:20]
        ]),
        table_types="\n".join([
            f"- **{t}:** {c} таблиц"
            for t, c in analysis_results["tables_by_type"].items()
        ]),
        style_guide="\n".join([
            f"### {section_type}\n- Типичные фразы: {', '.join(phrases) if phrases else 'Не найдено'}"
            for section_type, phrases in analysis_results["style_guide"].items()
        ]),
    )
    
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(guide)
    
    print(f"  ✅ Контент-гайд создан: {output_path}")


if __name__ == "__main__":
    sys.exit(main())

