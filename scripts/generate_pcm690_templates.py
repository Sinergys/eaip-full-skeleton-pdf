"""
Utility script to bootstrap PCM №690 document templates.

Creates:
- templates/pcm690/energy_passport_template.xlsx
- templates/pcm690/energy_audit_template.docx
- templates/pcm690/README.md (if missing)

Requirements: openpyxl, python-docx
"""

from pathlib import Path

from openpyxl import Workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.utils import get_column_letter
from openpyxl.workbook.defined_name import DefinedName
from docx import Document
from docx.shared import Inches

ROOT_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = ROOT_DIR / "templates" / "pcm690"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def _setup_summary_sheet(ws):
    ws.title = "Summary"
    ws["A1"] = "Энергопаспорт предприятия"
    ws["A2"] = "Предприятие:"
    ws["B2"] = "{{enterprise.name}}"
    ws["A3"] = "ИНН:"
    ws["B3"] = "{{enterprise.tax_id}}"
    ws["A4"] = "Период отчёта (с):"
    ws["B4"] = "{{period.start}}"
    ws["A5"] = "Период отчёта (по):"
    ws["B5"] = "{{period.end}}"
    ws["A7"] = "Ответственный за заполнение:"
    ws["B7"] = "{{responsible.full_name}}"
    ws["A8"] = "Контактный телефон:"
    ws["B8"] = "{{responsible.phone}}"
    ws.column_dimensions["A"].width = 32
    ws.column_dimensions["B"].width = 48


def _setup_energy_sheet(ws, resource_name):
    ws.title = resource_name
    headers = ["Месяц", "Факт, сум", "Факт, м3", "Норматив, сум", "Норматив, м3"]
    ws.append(["{{meta.year}}", "", "", "", ""])
    ws.append(headers)
    months = [
        "Январь",
        "Февраль",
        "Март",
        "Апрель",
        "Май",
        "Июнь",
        "Июль",
        "Август",
        "Сентябрь",
        "Октябрь",
        "Ноябрь",
        "Декабрь",
    ]
    for month in months:
        ws.append([month, "{{data.fact_sum}}", "{{data.fact_volume}}", "{{data.norm_sum}}", "{{data.norm_volume}}"])

    for idx in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(idx)].width = 18

    # Define named range for charts
    start_row = 3
    end_row = start_row + len(months) - 1
    named_range = DefinedName(
        f"{resource_name}_FactVolume",
        attr_text=f"'{resource_name}'!$C${start_row}:$C${end_row}",
    )

    workbook = ws.parent
    workbook.defined_names.add(named_range)

    chart = BarChart()
    chart.title = f"{resource_name}: Фактическое потребление (м3)"
    data = Reference(ws, min_col=3, min_row=start_row - 1, max_row=end_row)
    categories = Reference(ws, min_col=1, min_row=start_row, max_row=end_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(categories)
    chart.y_axis.title = "м3"
    chart.x_axis.title = "Месяц"
    chart.height = 9
    chart.width = 18
    ws.add_chart(chart, "G3")


def _setup_analytics_sheet(ws):
    ws.title = "Analytics"
    ws["A1"] = "Ключевые показатели"
    ws.append(["Показатель", "Значение", "Источник"])
    ws.append(["Годовое потребление газа, м3", "{{analytics.gas.total_volume}}", "Газ"])
    ws.append(["Годовое потребление электроэнергии, кВт·ч", "{{analytics.power.total_volume}}", "Electricity"])
    ws.append(["Отклонение от норматива, %", "{{analytics.total.delta_pct}}", "Aggregated"])
    ws.append(["Удельное потребление (газ), м3/продукция", "{{analytics.gas.specific_volume}}", "Расчёт"])
    ws.append(["Удельное потребление (электроэнергия), кВт·ч/продукция", "{{analytics.power.specific_volume}}", "Расчёт"])
    ws.column_dimensions["A"].width = 48
    ws.column_dimensions["B"].width = 24
    ws.column_dimensions["C"].width = 16


def create_excel_template():
    wb = Workbook()
    summary = wb.active
    _setup_summary_sheet(summary)
    _setup_energy_sheet(wb.create_sheet("Gas"), "Gas")
    _setup_energy_sheet(wb.create_sheet("Electricity"), "Electricity")
    _setup_analytics_sheet(wb.create_sheet("Analytics"))

    output_path = OUTPUT_DIR / "energy_passport_template.xlsx"
    wb.save(output_path)
    print(f"[OK] Excel template saved to {output_path}")


def create_word_template():
    doc = Document()
    doc.add_heading("Отчёт об энергоаудите", level=0)
    doc.add_paragraph("{{enterprise.full_name}}")
    doc.add_paragraph("ИНН: {{enterprise.tax_id}}")
    doc.add_paragraph("Адрес: {{enterprise.address}}")

    doc.add_heading("Содержание", level=1)
    doc.add_paragraph("TOC будет сформирован автоматически в MS Word (References → Table of Contents).")

    doc.add_heading("1. Вводная часть", level=1)
    doc.add_paragraph("Цель аудита: {{report.goal}}")
    doc.add_paragraph("Нормативная база: ПКМ №690, внутренние регламенты предприятия.")

    doc.add_heading("2. Общая характеристика предприятия", level=1)
    doc.add_paragraph("Основные процессы: {{enterprise.processes}}")
    doc.add_paragraph("Персонал: {{enterprise.headcount}}")

    doc.add_heading("3. Анализ энергопотребления", level=1)
    doc.add_paragraph("Сводные показатели: {{analytics.summary}}")
    doc.add_paragraph("График 1 — потребление газа (см. вложенную диаграмму).")
    doc.add_paragraph("График 2 — потребление электроэнергии (см. вложенную диаграмму).")
    doc.add_picture((OUTPUT_DIR / "placeholder_chart.png"), width=Inches(5)) if (OUTPUT_DIR / "placeholder_chart.png").exists() else None

    doc.add_heading("4. Выявленные проблемы и рекомендации", level=1)
    doc.add_paragraph("{{findings.summary}}")
    doc.add_paragraph("Рекомендуемые мероприятия: {{recommendations.list}}")

    doc.add_heading("5. План мероприятий", level=1)
    doc.add_paragraph("{{actions.timeline}}")

    doc.add_heading("Приложения", level=1)
    doc.add_paragraph("A. Детализированные таблицы потребления.")
    doc.add_paragraph("B. Фотофиксация объектов.")

    output_path = OUTPUT_DIR / "energy_audit_template.docx"
    doc.save(output_path)
    print(f"[OK] Word template saved to {output_path}")


def ensure_templates_readme():
    readme_path = OUTPUT_DIR / "README.md"
    if readme_path.exists():
        return
    readme_path.write_text(
        """# PCM №690 Templates

Generated via `scripts/generate_pcm690_templates.py`.

- `energy_passport_template.xlsx` — табличная структура с листами Summary/Gas/Electricity/Analytics и заготовками гистограмм.
- `energy_audit_template.docx` — текстовый отчёт с секциями и placeholder'ами.

Placeholders оформлены в формате `{{key}}` и должны заменяться backend-сервисом при формировании итоговых документов.
""",
        encoding="utf-8",
    )
    print(f"[OK] README created at {readme_path}")


def main():
    ensure_templates_readme()
    create_excel_template()
    create_word_template()


if __name__ == "__main__":
    main()

