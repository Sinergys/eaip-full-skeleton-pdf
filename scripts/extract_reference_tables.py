"""
Извлечение и нормализация таблиц из образцового Word-отчёта.

Парсит таблицы из "МЕТИН ИРОДА ОТЧЕТ 1107.docx" и приводит их к нормализованному виду
для использования как эталон и источник данных.
"""

import json
import sys
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.table import Table

PROJECT_ROOT = Path(__file__).resolve().parent.parent
REFERENCE_DOC_PATH = Path(r"C:\Users\DELL\Downloads\Telegram Desktop\МЕТИН ИРОДА ОТЧЕТ 1107.docx")
OUTPUT_DIR = PROJECT_ROOT / "data" / "reference_analysis" / "tables"
ANALYSIS_JSON = PROJECT_ROOT / "data" / "reference_analysis" / "reference_word_report_analysis.json"


def load_analysis_results() -> Dict[str, Any]:
    """Загружает результаты предыдущего анализа."""
    if not ANALYSIS_JSON.exists():
        return {}
    
    with open(ANALYSIS_JSON, "r", encoding="utf-8") as f:
        return json.load(f)


def extract_table_data(table: Table, table_index: int) -> Dict[str, Any]:
    """
    Извлекает данные из таблицы Word.
    
    Args:
        table: Объект таблицы из python-docx
        table_index: Индекс таблицы в документе
    
    Returns:
        Словарь с данными таблицы
    """
    data = {
        "index": table_index,
        "rows": len(table.rows),
        "cols": len(table.columns) if table.rows else 0,
        "headers": [],
        "data": [],
        "raw_text": [],
    }
    
    # Извлекаем все данные
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            row_data.append(cell_text)
        
        data["raw_text"].append(row_data)
        
        if row_idx == 0:
            # Первая строка - заголовки
            data["headers"] = row_data
        else:
            data["data"].append(row_data)
    
    return data


def classify_table(table_data: Dict[str, Any]) -> str:
    """
    Классифицирует таблицу по типу на основе заголовков и содержимого.
    
    Returns:
        Тип таблицы: "equipment", "specific_consumption", "consumption_structure",
                     "losses", "measures", "other"
    """
    headers_text = " ".join(table_data["headers"]).lower()
    data_text = " ".join([" ".join(row) for row in table_data["data"][:5]]).lower()
    combined_text = headers_text + " " + data_text
    
    # Определяем тип по ключевым словам
    if any(word in combined_text for word in ["мероприятия", "рекомендации", "экономия", "окупаемость", "capex", "payback"]):
        return "measures"
    elif any(word in combined_text for word in ["оборудование", "equipment", "мощность", "power", "квт", "kw"]):
        return "equipment"
    elif any(word in combined_text for word in ["удельный", "specific", "расход", "consumption", "единицу", "единица"]):
        return "specific_consumption"
    elif any(word in combined_text for word in ["баланс", "balance", "технологические", "собственные", "хоз-бытовые"]):
        return "consumption_structure"
    elif any(word in combined_text for word in ["потери", "losses", "трансформатор", "transformer"]):
        return "losses"
    else:
        return "other"


def normalize_equipment_table(table_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Нормализует таблицу оборудования."""
    normalized = []
    headers = [h.lower() for h in table_data["headers"]]
    
    # Маппинг возможных названий колонок
    column_mapping = {
        "name": ["наименование", "название", "оборудование", "name", "equipment"],
        "section": ["раздел", "цех", "участок", "section", "workshop"],
        "power_kw": ["мощность", "квт", "kw", "power", "мощность квт"],
        "count": ["количество", "кол-во", "шт", "count", "quantity", "qty"],
        "vfd": ["чрп", "частотный", "vfd", "частотно-регулируемый"],
        "group": ["группа", "категория", "group", "category"],
    }
    
    # Находим индексы колонок
    column_indices = {}
    for field, keywords in column_mapping.items():
        for idx, header in enumerate(headers):
            if any(kw in header for kw in keywords):
                column_indices[field] = idx
                break
    
    # Обрабатываем данные
    for row in table_data["data"]:
        if not any(row):  # Пропускаем пустые строки
            continue
        
        item = {
            "name": row[column_indices.get("name", 0)] if column_indices.get("name") is not None and column_indices["name"] < len(row) else "",
            "section": row[column_indices.get("section", 1)] if column_indices.get("section") is not None and column_indices["section"] < len(row) else "",
            "power_kw": _parse_float(row[column_indices.get("power_kw", 2)]) if column_indices.get("power_kw") is not None and column_indices["power_kw"] < len(row) else 0.0,
            "count": _parse_float(row[column_indices.get("count", 3)]) if column_indices.get("count") is not None and column_indices["count"] < len(row) else 1.0,
            "vfd": _parse_bool(row[column_indices.get("vfd", 4)]) if column_indices.get("vfd") is not None and column_indices["vfd"] < len(row) else False,
            "group": row[column_indices.get("group", 5)] if column_indices.get("group") is not None and column_indices["group"] < len(row) else "",
        }
        
        normalized.append(item)
    
    return normalized


def normalize_measures_table(table_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Нормализует таблицу мероприятий."""
    normalized = []
    headers = [h.lower() for h in table_data["headers"]]
    
    # Проверяем, является ли это таблицей в формате "Показатель" / "Значение"
    is_key_value_format = len(headers) == 2 and any(
        word in headers[0] for word in ["показатель", "эффект", "название", "name", "indicator"]
    ) and any(
        word in headers[1] for word in ["значение", "примерно", "value", "значение"]
    )
    
    if is_key_value_format:
        # Обрабатываем формат "Показатель" / "Значение"
        measure = {
            "id": 1,
            "name": "",
            "essence": "",
            "capex": 0.0,
            "saving_kwh": 0.0,
            "saving_money": 0.0,
            "payback_years": 0.0,
            "priority": "medium",
            "raw_data": {},  # Сохраняем сырые данные для ручной обработки
        }
        
        for row in table_data["data"]:
            if len(row) < 2:
                continue
            
            key = row[0].lower()
            value = row[1]
            
            # Извлекаем данные по ключевым словам
            if any(word in key for word in ["мероприятие", "название", "наименование", "эффект"]):
                measure["name"] = row[0]
                measure["essence"] = value
            elif any(word in key for word in ["стоимость", "затраты", "инвестиции", "capex", "модернизация"]):
                measure["capex"] = _parse_float(value)
                measure["raw_data"]["capex_text"] = value
            elif any(word in key for word in ["экономия", "квт", "kwh", "электроэнергия"]) and "деньги" not in key:
                measure["saving_kwh"] = _parse_float(value)
                measure["raw_data"]["saving_kwh_text"] = value
            elif any(word in key for word in ["экономия", "деньги", "сум", "руб"]) and "деньги" in key or "сум" in value.lower():
                measure["saving_money"] = _parse_float(value)
                measure["raw_data"]["saving_money_text"] = value
            elif any(word in key for word in ["окупаемость", "payback", "срок"]):
                measure["payback_years"] = _parse_float(value)
                measure["raw_data"]["payback_text"] = value
            
            measure["raw_data"][row[0]] = value
        
        if measure["name"] or measure.get("raw_data"):
            normalized.append(measure)
    else:
        # Стандартный формат таблицы с колонками
        column_mapping = {
            "id": ["№", "номер", "id", "n", "#"],
            "name": ["наименование", "название", "мероприятие", "name", "measure"],
            "essence": ["суть", "описание", "описание мероприятия", "essence", "description"],
            "capex": ["капитальные", "затраты", "инвестиции", "capex", "investment", "стоимость", "руб", "сум"],
            "saving_kwh": ["экономия", "квт·ч", "kwh", "электроэнергия", "saving", "экономия квт"],
            "saving_money": ["экономия", "руб", "сум", "деньги", "saving_money", "экономия денег"],
            "payback_years": ["окупаемость", "лет", "год", "payback", "срок окупаемости"],
            "priority": ["приоритет", "важность", "priority", "важность"],
        }
        
        # Находим индексы колонок
        column_indices = {}
        for field, keywords in column_mapping.items():
            for idx, header in enumerate(headers):
                if any(kw in header for kw in keywords):
                    if field not in column_indices:  # Берем первое совпадение
                        column_indices[field] = idx
        
        # Обрабатываем данные
        for row_idx, row in enumerate(table_data["data"]):
            if not any(row):  # Пропускаем пустые строки
                continue
            
            # Пытаемся извлечь значения
            measure = {
                "id": _parse_int(row[column_indices.get("id", 0)]) if column_indices.get("id") is not None and column_indices["id"] < len(row) else row_idx + 1,
                "name": row[column_indices.get("name", 1)] if column_indices.get("name") is not None and column_indices["name"] < len(row) else "",
                "essence": row[column_indices.get("essence", 2)] if column_indices.get("essence") is not None and column_indices["essence"] < len(row) else "",
                "capex": _parse_float(row[column_indices.get("capex", 3)]) if column_indices.get("capex") is not None and column_indices["capex"] < len(row) else 0.0,
                "saving_kwh": _parse_float(row[column_indices.get("saving_kwh", 4)]) if column_indices.get("saving_kwh") is not None and column_indices["saving_kwh"] < len(row) else 0.0,
                "saving_money": _parse_float(row[column_indices.get("saving_money", 5)]) if column_indices.get("saving_money") is not None and column_indices["saving_money"] < len(row) else 0.0,
                "payback_years": _parse_float(row[column_indices.get("payback_years", 6)]) if column_indices.get("payback_years") is not None and column_indices["payback_years"] < len(row) else 0.0,
                "priority": row[column_indices.get("priority", 7)] if column_indices.get("priority") is not None and column_indices["priority"] < len(row) else "medium",
            }
            
            # Если название пустое, пропускаем
            if not measure["name"]:
                continue
            
            normalized.append(measure)
    
    return normalized


def normalize_specific_consumption_table(table_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Нормализует таблицу удельного расхода."""
    normalized = []
    headers = [h.lower() for h in table_data["headers"]]
    
    column_mapping = {
        "product": ["продукция", "product", "изделие", "наименование"],
        "period": ["период", "квартал", "год", "period", "quarter", "year"],
        "energy_type": ["энергия", "ресурс", "energy", "resource", "тип"],
        "value": ["значение", "расход", "value", "consumption", "квт·ч", "kwh"],
        "unit": ["единица", "unit", "ед. изм"],
    }
    
    column_indices = {}
    for field, keywords in column_mapping.items():
        for idx, header in enumerate(headers):
            if any(kw in header for kw in keywords):
                column_indices[field] = idx
                break
    
    for row in table_data["data"]:
        if not any(row):
            continue
        
        item = {
            "product": row[column_indices.get("product", 0)] if column_indices.get("product") is not None and column_indices["product"] < len(row) else "",
            "period": row[column_indices.get("period", 1)] if column_indices.get("period") is not None and column_indices["period"] < len(row) else "",
            "energy_type": row[column_indices.get("energy_type", 2)] if column_indices.get("energy_type") is not None and column_indices["energy_type"] < len(row) else "electricity",
            "value": _parse_float(row[column_indices.get("value", 3)]) if column_indices.get("value") is not None and column_indices["value"] < len(row) else 0.0,
            "unit": row[column_indices.get("unit", 4)] if column_indices.get("unit") is not None and column_indices["unit"] < len(row) else "кВт·ч/кг",
        }
        
        normalized.append(item)
    
    return normalized


def normalize_consumption_structure_table(table_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Нормализует таблицу структуры потребления."""
    normalized = []
    headers = [h.lower() for h in table_data["headers"]]
    
    column_mapping = {
        "period": ["период", "квартал", "год", "period", "quarter"],
        "technological": ["технологические", "технология", "technological"],
        "own_needs": ["собственные", "собственные нужды", "own_needs"],
        "production": ["производственные", "production"],
        "household": ["хоз-бытовые", "хозяйственные", "household"],
        "total": ["итого", "total", "всего"],
    }
    
    column_indices = {}
    for field, keywords in column_mapping.items():
        for idx, header in enumerate(headers):
            if any(kw in header for kw in keywords):
                column_indices[field] = idx
                break
    
    for row in table_data["data"]:
        if not any(row):
            continue
        
        item = {
            "period": row[column_indices.get("period", 0)] if column_indices.get("period") is not None and column_indices["period"] < len(row) else "",
            "technological": _parse_float(row[column_indices.get("technological", 1)]) if column_indices.get("technological") is not None and column_indices["technological"] < len(row) else 0.0,
            "own_needs": _parse_float(row[column_indices.get("own_needs", 2)]) if column_indices.get("own_needs") is not None and column_indices["own_needs"] < len(row) else 0.0,
            "production": _parse_float(row[column_indices.get("production", 3)]) if column_indices.get("production") is not None and column_indices["production"] < len(row) else 0.0,
            "household": _parse_float(row[column_indices.get("household", 4)]) if column_indices.get("household") is not None and column_indices["household"] < len(row) else 0.0,
            "total": _parse_float(row[column_indices.get("total", 5)]) if column_indices.get("total") is not None and column_indices["total"] < len(row) else 0.0,
        }
        
        normalized.append(item)
    
    return normalized


def normalize_losses_table(table_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Нормализует таблицу потерь."""
    normalized = []
    headers = [h.lower() for h in table_data["headers"]]
    
    column_mapping = {
        "transformer": ["трансформатор", "transformer", "тп", "подстанция"],
        "power_kva": ["мощность", "ква", "kva", "power"],
        "loss_active": ["потери активные", "активные", "loss_active", "квт·ч"],
        "loss_reactive": ["потери реактивные", "реактивные", "loss_reactive", "квар·ч"],
        "percentage": ["процент", "процент потерь", "percentage", "%"],
    }
    
    column_indices = {}
    for field, keywords in column_mapping.items():
        for idx, header in enumerate(headers):
            if any(kw in header for kw in keywords):
                column_indices[field] = idx
                break
    
    for row in table_data["data"]:
        if not any(row):
            continue
        
        item = {
            "transformer": row[column_indices.get("transformer", 0)] if column_indices.get("transformer") is not None and column_indices["transformer"] < len(row) else "",
            "power_kva": _parse_float(row[column_indices.get("power_kva", 1)]) if column_indices.get("power_kva") is not None and column_indices["power_kva"] < len(row) else 0.0,
            "loss_active": _parse_float(row[column_indices.get("loss_active", 2)]) if column_indices.get("loss_active") is not None and column_indices["loss_active"] < len(row) else 0.0,
            "loss_reactive": _parse_float(row[column_indices.get("loss_reactive", 3)]) if column_indices.get("loss_reactive") is not None and column_indices["loss_reactive"] < len(row) else 0.0,
            "percentage": _parse_float(row[column_indices.get("percentage", 4)]) if column_indices.get("percentage") is not None and column_indices["percentage"] < len(row) else 0.0,
        }
        
        normalized.append(item)
    
    return normalized


def normalize_other_table(table_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Нормализует прочие таблицы (сохраняет как есть с метаданными)."""
    return {
        "headers": table_data["headers"],
        "data": table_data["data"],
        "raw_text": table_data["raw_text"],
    }


def _parse_float(value: str) -> float:
    """Парсит строку в float, обрабатывая запятые и пробелы."""
    if not value:
        return 0.0
    
    # Убираем пробелы и заменяем запятую на точку
    cleaned = value.replace(" ", "").replace(",", ".")
    
    try:
        return float(cleaned)
    except (ValueError, TypeError):
        return 0.0


def _parse_int(value: str) -> int:
    """Парсит строку в int."""
    if not value:
        return 0
    
    try:
        return int(float(value.replace(",", ".")))
    except (ValueError, TypeError):
        return 0


def _parse_bool(value: str) -> bool:
    """Парсит строку в bool."""
    if not value:
        return False
    
    value_lower = value.lower()
    return any(word in value_lower for word in ["да", "yes", "есть", "есть", "vfd", "чрп", "частотный", "1", "true"])


def extract_all_tables(doc_path: Path) -> Dict[str, List[Dict[str, Any]]]:
    """
    Извлекает и нормализует все таблицы из документа.
    
    Returns:
        Словарь с нормализованными таблицами по типам
    """
    print(f"📄 Открытие документа: {doc_path}")
    
    if not doc_path.exists():
        raise FileNotFoundError(f"Файл не найден: {doc_path}")
    
    doc = Document(str(doc_path))
    
    # Группируем таблицы по типам
    tables_by_type = {
        "equipment": [],
        "specific_consumption": [],
        "consumption_structure": [],
        "losses": [],
        "measures": [],
        "other": [],
    }
    
    print(f"📊 Извлечение {len(doc.tables)} таблиц...")
    
    for table_idx, table in enumerate(doc.tables):
        # Извлекаем сырые данные
        table_data = extract_table_data(table, table_idx)
        
        # Классифицируем
        table_type = classify_table(table_data)
        
        # Нормализуем в зависимости от типа
        if table_type == "equipment":
            normalized = normalize_equipment_table(table_data)
        elif table_type == "measures":
            normalized = normalize_measures_table(table_data)
        elif table_type == "specific_consumption":
            normalized = normalize_specific_consumption_table(table_data)
        elif table_type == "consumption_structure":
            normalized = normalize_consumption_structure_table(table_data)
        elif table_type == "losses":
            normalized = normalize_losses_table(table_data)
        else:
            normalized = normalize_other_table(table_data)
        
        tables_by_type[table_type].append({
            "table_index": table_idx,
            "headers": table_data["headers"],
            "normalized_data": normalized,
            "raw_data": table_data,
        })
        
        print(f"  ✅ Таблица {table_idx + 1}: тип={table_type}, строк={len(table_data['data'])}")
    
    return tables_by_type


def save_tables(tables_by_type: Dict[str, List[Dict[str, Any]]], output_dir: Path):
    """Сохраняет нормализованные таблицы в JSON файлы."""
    output_dir.mkdir(parents=True, exist_ok=True)
    
    for table_type, tables in tables_by_type.items():
        if not tables:
            continue
        
        output_file = output_dir / f"tables_{table_type}.json"
        
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump({
                "table_type": table_type,
                "tables_count": len(tables),
                "tables": tables,
            }, f, ensure_ascii=False, indent=2)
        
        print(f"💾 Сохранено: {output_file} ({len(tables)} таблиц)")


def main():
    """Основная функция."""
    print("🔍 ИЗВЛЕЧЕНИЕ И НОРМАЛИЗАЦИЯ ТАБЛИЦ ИЗ ОБРАЗЦОВОГО ОТЧЁТА")
    print("=" * 80)
    
    # Извлекаем таблицы
    tables_by_type = extract_all_tables(REFERENCE_DOC_PATH)
    
    # Сохраняем
    save_tables(tables_by_type, OUTPUT_DIR)
    
    # Статистика
    print("\n📊 Статистика:")
    for table_type, tables in tables_by_type.items():
        if tables:
            total_items = sum(len(t["normalized_data"]) if isinstance(t["normalized_data"], list) else 1 for t in tables)
            print(f"  - {table_type}: {len(tables)} таблиц, {total_items} записей")
    
    print("\n✅ Извлечение завершено!")
    return 0


if __name__ == "__main__":
    sys.exit(main())

