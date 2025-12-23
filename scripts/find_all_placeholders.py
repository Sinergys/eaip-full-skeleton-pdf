"""
Скрипт для поиска всех placeholder'ов в проекте
Находит TEMPLATE_ константы и {{...}} placeholder'ы
"""

import re
from pathlib import Path
from typing import List, Tuple
from openpyxl import load_workbook
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent

def find_template_constants(file_path: Path) -> List[Tuple[int, str, str]]:
    """Поиск TEMPLATE_ констант в Python файлах"""
    results = []
    try:
        content = file_path.read_text(encoding="utf-8")
        for line_num, line in enumerate(content.splitlines(), 1):
            # Поиск TEMPLATE_ констант
            for match in re.finditer(r'TEMPLATE_([A-Z_]+)', line):
                results.append((line_num, match.group(0), "TEMPLATE_CONSTANT"))
    except Exception:
        pass
    return results

def find_curly_placeholders(file_path: Path) -> List[Tuple[int, str, str]]:
    """Поиск {{...}} placeholder'ов в текстовых файлах"""
    results = []
    try:
        content = file_path.read_text(encoding="utf-8")
        for line_num, line in enumerate(content.splitlines(), 1):
            # Поиск {{...}} placeholder'ов
            for match in re.finditer(r'\{\{([^}]+)\}\}', line):
                placeholder = match.group(0)
                key = match.group(1).strip()
                results.append((line_num, placeholder, key))
    except Exception:
        pass
    return results

def find_excel_placeholders(file_path: Path) -> List[Tuple[str, int, int, str]]:
    """Поиск placeholder'ов в Excel файлах"""
    results = []
    try:
        wb = load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row_idx, row in enumerate(ws.iter_rows(), 1):
                for col_idx, cell in enumerate(row, 1):
                    if cell.value and isinstance(cell.value, str):
                        # Поиск {{...}} placeholder'ов
                        for match in re.finditer(r'\{\{([^}]+)\}\}', cell.value):
                            placeholder = match.group(0)
                            key = match.group(1).strip()
                            results.append((sheet_name, row_idx, col_idx, key))
    except Exception:
        pass
    return results

def find_word_placeholders(file_path: Path) -> List[Tuple[str, str]]:
    """Поиск placeholder'ов в Word файлах"""
    results = []
    try:
        doc = Document(file_path)
        # Поиск в параграфах
        for para_idx, para in enumerate(doc.paragraphs, 1):
            if '{{' in para.text:
                for match in re.finditer(r'\{\{([^}]+)\}\}', para.text):
                    key = match.group(1).strip()
                    results.append((f"paragraph_{para_idx}", key))
        
        # Поиск в таблицах
        for table_idx, table in enumerate(doc.tables, 1):
            for row_idx, row in enumerate(table.rows, 1):
                for col_idx, cell in enumerate(row.cells, 1):
                    if '{{' in cell.text:
                        for match in re.finditer(r'\{\{([^}]+)\}\}', cell.text):
                            key = match.group(1).strip()
                            results.append((f"table_{table_idx}_row_{row_idx}_col_{col_idx}", key))
    except Exception:
        pass
    return results

def get_data_source(key: str) -> str:
    """Определение источника данных для placeholder'а"""
    key_lower = key.lower()
    
    if 'enterprise' in key_lower or 'company' in key_lower:
        return "enterprise_data (БД: enterprises)"
    elif 'period' in key_lower or 'date' in key_lower or 'year' in key_lower:
        return "period_data (метаданные)"
    elif 'responsible' in key_lower or 'director' in key_lower:
        return "enterprise_data (БД: enterprises)"
    elif 'analytics' in key_lower:
        return "analytics_data (расчет из aggregated JSON)"
    elif 'data.fact' in key_lower or 'data.norm' in key_lower:
        return "aggregated_data (JSON: resources)"
    elif 'meta' in key_lower:
        return "meta_data (метаданные периода)"
    elif 'energy' in key_lower or 'saving' in key_lower:
        return "energy_data (расчет из мероприятий)"
    elif 'gas' in key_lower:
        return "aggregated_data.resources.gas"
    elif 'power' in key_lower or 'electricity' in key_lower:
        return "aggregated_data.resources.electricity"
    else:
        return "unknown (требует уточнения)"

def get_data_type(key: str) -> str:
    """Определение типа данных для placeholder'а"""
    key_lower = key.lower()
    
    if 'name' in key_lower or 'title' in key_lower:
        return "string"
    elif 'id' in key_lower or 'inn' in key_lower:
        return "string"
    elif 'date' in key_lower or 'start' in key_lower or 'end' in key_lower:
        return "date"
    elif 'volume' in key_lower or 'sum' in key_lower or 'cost' in key_lower:
        return "number"
    elif 'pct' in key_lower or 'percent' in key_lower or 'efficiency' in key_lower:
        return "float"
    elif 'phone' in key_lower or 'address' in key_lower:
        return "string"
    else:
        return "mixed"

def main():
    """Основная функция"""
    print("=" * 100)
    print("ПОИСК ВСЕХ PLACEHOLDER'ОВ В ПРОЕКТЕ")
    print("=" * 100)
    
    all_placeholders = []
    
    # Поиск в Python файлах
    print("\n🔍 Поиск в Python файлах...")
    for py_file in PROJECT_ROOT.rglob("*.py"):
        if "venv" in str(py_file) or "__pycache__" in str(py_file):
            continue
        
        # TEMPLATE_ константы
        template_consts = find_template_constants(py_file)
        for line_num, const, _ in template_consts:
            all_placeholders.append({
                "file": str(py_file.relative_to(PROJECT_ROOT)),
                "line": line_num,
                "placeholder": const,
                "type": "TEMPLATE_CONSTANT",
                "data_type": "constant",
                "source": "code_constant"
            })
        
        # {{...}} placeholder'ы
        curly_placeholders = find_curly_placeholders(py_file)
        for line_num, placeholder, key in curly_placeholders:
            all_placeholders.append({
                "file": str(py_file.relative_to(PROJECT_ROOT)),
                "line": line_num,
                "placeholder": placeholder,
                "type": "CURLY_BRACES",
                "data_type": get_data_type(key),
                "source": get_data_source(key)
            })
    
    # Поиск в Excel файлах
    print("🔍 Поиск в Excel файлах...")
    for xlsx_file in PROJECT_ROOT.rglob("*.xlsx"):
        if "venv" in str(xlsx_file) or "__pycache__" in str(xlsx_file):
            continue
        
        excel_placeholders = find_excel_placeholders(xlsx_file)
        for sheet_name, row, col, key in excel_placeholders:
            all_placeholders.append({
                "file": str(xlsx_file.relative_to(PROJECT_ROOT)),
                "line": f"{sheet_name}:{row}:{col}",
                "placeholder": f"{{{{{key}}}}}",
                "type": "EXCEL_CELL",
                "data_type": get_data_type(key),
                "source": get_data_source(key)
            })
    
    # Поиск в Word файлах
    print("🔍 Поиск в Word файлах...")
    for docx_file in PROJECT_ROOT.rglob("*.docx"):
        if "venv" in str(docx_file) or "__pycache__" in str(docx_file):
            continue
        
        word_placeholders = find_word_placeholders(docx_file)
        for location, key in word_placeholders:
            all_placeholders.append({
                "file": str(docx_file.relative_to(PROJECT_ROOT)),
                "line": location,
                "placeholder": f"{{{{{key}}}}}",
                "type": "WORD_DOC",
                "data_type": get_data_type(key),
                "source": get_data_source(key)
            })
    
    # Вывод результатов
    print(f"\n✅ Найдено placeholder'ов: {len(all_placeholders)}")
    print("\n" + "=" * 100)
    print("ТАБЛИЦА PLACEHOLDER'ОВ")
    print("=" * 100)
    
    # Группировка по файлам
    by_file = {}
    for p in all_placeholders:
        file = p["file"]
        if file not in by_file:
            by_file[file] = []
        by_file[file].append(p)
    
    # Вывод таблицы
    print(f"\n{'Файл':<50} {'Строка/Ячейка':<20} {'Placeholder':<40} {'Тип данных':<15} {'Источник':<40}")
    print("-" * 165)
    
    for file, placeholders in sorted(by_file.items()):
        for p in placeholders:
            print(f"{p['file']:<50} {str(p['line']):<20} {p['placeholder']:<40} {p['data_type']:<15} {p['source']:<40}")
    
    # Сохранение в JSON
    import json
    output_file = PROJECT_ROOT / "data" / "aggregated" / "all_placeholders.json"
    output_file.parent.mkdir(parents=True, exist_ok=True)
    
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(all_placeholders, f, ensure_ascii=False, indent=2)
    
    print(f"\n💾 Результаты сохранены в: {output_file}")
    
    # Статистика
    print("\n" + "=" * 100)
    print("СТАТИСТИКА")
    print("=" * 100)
    
    by_type = {}
    by_source = {}
    for p in all_placeholders:
        by_type[p["type"]] = by_type.get(p["type"], 0) + 1
        by_source[p["source"]] = by_source.get(p["source"], 0) + 1
    
    print("\nПо типам:")
    for type_name, count in sorted(by_type.items()):
        print(f"  {type_name}: {count}")
    
    print("\nПо источникам:")
    for source, count in sorted(by_source.items()):
        print(f"  {source}: {count}")
    
    return all_placeholders

if __name__ == "__main__":
    main()
